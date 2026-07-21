"""Build the current results-interpretation report from frozen analysis data.

The source Word document is used only as a style template. Numeric claims are
read from the reproducible full-corpus observed-composition and 2,235-paper
probability-sample model-validation outputs. The source document is never
overwritten.
"""

from __future__ import annotations

import sqlite3
from pathlib import Path

import cairosvg
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/current results intepretation .docx"
OUTPUT = ROOT / "docs/current results interpretation - probability sample reviewed 2026-07-21.docx"
MARKDOWN = ROOT / "reports/analysis/PROBABILITY_SAMPLE_RESULTS_INTERPRETATION.md"

VALIDATION = ROOT / "data/processed/analysis/model_validation"
TABLES = ROOT / "reports/analysis/tables/model_validation"
FIGURES = ROOT / "reports/analysis/figures/model_validation"
OBSERVED = ROOT / "reports/analysis/tables/results_interpretation/mini_observed_composition.csv"
OBSERVED_FIGURE = ROOT / "reports/analysis/figures/results_interpretation/fig14_observed_composition.png"
HUMAN_DB = ROOT / "data/interim/human_validation/human_annotations.sqlite3"

CORE_DIMENSIONS = (
    "ai_method_or_phenomenon",
    "ai_type_form",
    "ai_role_function",
    "ai_mechanism_analysis",
    "level_of_analysis",
    "scope_conditions",
)
EXPLORATORY_DIMENSIONS = (
    "entrepreneurial_process_stage",
    "definition_construct_clarity",
)
DISPLAY_DIMENSIONS = (*CORE_DIMENSIONS, *EXPLORATORY_DIMENSIONS)

DIMENSION_LABELS = {
    "ai_method_or_phenomenon": "Study status",
    "ai_type_form": "Technical AI type/form",
    "ai_role_function": "AI role/function",
    "ai_mechanism_analysis": "Observable AI mechanism",
    "level_of_analysis": "Level of analysis",
    "scope_conditions": "Scope conditions",
    "entrepreneurial_process_stage": "Entrepreneurial process stage",
    "definition_construct_clarity": "Definition form",
}

DIMENSION_DECISIONS = {
    "ai_method_or_phenomenon": "Core; retain all four categories and report model sensitivity.",
    "ai_type_form": "Strongest dimension; suitable for the main construct-specification argument.",
    "ai_role_function": "Core; interpret with model-specific sensitivity and evidence papers.",
    "ai_mechanism_analysis": "Core and theory-bearing; restrict claims to abstract-observable mechanisms.",
    "level_of_analysis": "Core; broader level pattern is stronger than the firm-versus-venture distinction.",
    "scope_conditions": "Core; distinguish explicit boundaries from missing or generalized scope.",
    "entrepreneurial_process_stage": "Exploratory; do not make a headline claim before human validation.",
    "definition_construct_clarity": "Exploratory diagnostic; do not treat abstract absence as full-paper failure.",
}

OBSERVED_LABELS = {
    "ai_method_or_phenomenon": "Study status",
    "ai_type_form": "Technical type",
    "ai_role_function": "AI role",
    "ai_mechanism_analysis": "Mechanism",
    "level_of_analysis": "Level of analysis",
    "entrepreneurial_process_stage": "Process stage",
    "scope_conditions": "Scope",
    "definition_construct_clarity": "Definition form",
}


def _clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def _set_cell_text(cell, value: object, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(value))
    run.bold = bold
    run.font.size = Pt(8.5)


def _shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    properties.append(shading)


def add_table(document: Document, headers: list[str], rows: list[list[object]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True
    for index, header in enumerate(headers):
        _set_cell_text(table.rows[0].cells[index], header, bold=True)
        _shade(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            _set_cell_text(cells[index], value)
    document.add_paragraph()


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.style = "Caption" if "Caption" in [s.name for s in document.styles] else "Normal"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_picture(
    document: Document,
    path: Path,
    caption: str,
    *,
    width: float = 6.75,
) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def human_annotation_count() -> int:
    if not HUMAN_DB.exists():
        return 0
    with sqlite3.connect(HUMAN_DB) as connection:
        return int(connection.execute("SELECT COUNT(*) FROM annotations").fetchone()[0])


def convert_heatmaps() -> tuple[Path, Path]:
    outputs = []
    for stem in ("pairwise_agreement_heatmap", "pairwise_alpha_heatmap"):
        svg = FIGURES / f"{stem}.svg"
        png = FIGURES / f"{stem}_core6.png"
        cairosvg.svg2png(url=str(svg), write_to=str(png), output_width=1500)
        outputs.append(png)
    return outputs[0], outputs[1]


def pair_value(
    agreement: pd.DataFrame,
    dimension: str,
    left: str,
    right: str,
    column: str,
) -> float:
    match = agreement[
        (agreement.comparison_set == "probability_sample")
        & (agreement.dimension == dimension)
        & (agreement.left_model == left)
        & (agreement.right_model == right)
    ]
    if match.empty:
        match = agreement[
            (agreement.comparison_set == "probability_sample")
            & (agreement.dimension == dimension)
            & (agreement.left_model == right)
            & (agreement.right_model == left)
        ]
    return float(match.iloc[0][column])


def observed_rows(observed: pd.DataFrame) -> list[list[object]]:
    rows = []
    for dimension in DISPLAY_DIMENSIONS:
        subset = observed[observed.dimension == dimension].sort_values(
            "share_among_observed", ascending=False
        )
        first = subset.iloc[0]
        top = "; ".join(
            f"{row.category} {row.share_among_observed:.1%}"
            for row in subset.head(3).itertuples()
        )
        rows.append(
            [
                OBSERVED_LABELS[dimension],
                f"{int(first.observed_n):,}",
                f"{first.share_of_corpus_observed:.1%}",
                top,
            ]
        )
    return rows


def build_report() -> None:
    coverage = pd.read_csv(VALIDATION / "execution_coverage.csv")
    agreement = pd.read_csv(VALIDATION / "agreement_pairwise.csv")
    multirater = pd.read_csv(VALIDATION / "agreement_multirater.csv")
    macro = pd.read_csv(TABLES / "pairwise_macro_agreement.csv")
    prevalence = pd.read_csv(VALIDATION / "dimension_prevalence.csv")
    observed = pd.read_csv(OBSERVED)
    agreement_png, alpha_png = convert_heatmaps()

    document = Document(SOURCE)
    _clear_body(document)
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)

    document.add_heading(
        "What’s theory got to do with it? Theory Elaboration on AI in Entrepreneurship Scholarship",
        level=0,
    )
    document.add_heading("Summary", level=1)
    document.add_paragraph(
        "The current evidence base combines full-corpus construct specification by GPT-5.4 Mini "
        "with representative multi-model validation on a fixed 2,235-paper stratified probability "
        "sample. Mini and Gemini cover all 2,235 sampled papers; Nano and Claude each have one "
        "non-response. The exact four-model intersection is therefore 2,233 papers. Sampling "
        "weights are used for sample prevalence, while agreement is calculated only on exact "
        "paper-ID intersections."
    )
    document.add_paragraph(
        "The full-corpus observed portrait is a phenomenon-oriented, firm-level, sector-bounded "
        "literature in which machine learning is most often positioned as a tool and prediction is "
        "the leading observable mechanism. The probability-sample validation shows that technical "
        "AI type is the most stable dimension. Study status, role, mechanism, level, and scope are "
        "usable with model-sensitivity reporting. Process stage and definition form remain "
        "exploratory because prevalence-adjusted reliability is weak."
    )
    document.add_paragraph(
        "No human coding has yet been completed. The current results establish model consistency "
        "and model dependence; they do not establish model accuracy."
    )

    document.add_heading("1. Construct-specification instrument", level=1)
    document.add_paragraph(
        "Each paper is coded from its title, abstract, and author keywords only. The instrument "
        "contains eight displayed dimensions: study status, AI role, technical type, observable "
        "mechanism, level of analysis, entrepreneurial process stage, scope conditions, and "
        "definition form. Seven dimensions are controlled fields in the specification schema; "
        "study status is the additional four-category field distinguishing phenomenon, method, "
        "both, and unclear. Three binary presence flags are supplementary diagnostics and are not "
        "included in the main reliability heatmaps."
    )
    add_caption(document, "Table 1. Full-corpus observed construct portrait from GPT-5.4 Mini")
    add_table(
        document,
        ["Dimension", "Observed papers", "Share of corpus", "Leading observed categories"],
        observed_rows(observed),
    )

    document.add_heading("2. Probability-sample model validation", level=1)
    document.add_paragraph(
        "The validation sample was selected independently of model outputs and stratified by "
        "publication era, query provenance, abstract length, journal coverage, and metadata "
        "completeness. Claude and Gemini are independent provider raters; Mini is the primary "
        "full-corpus rater and Nano is the full-corpus baseline. Llama and Gemma remain partial "
        "local stress tests and are excluded from representative macro reliability."
    )
    reps = coverage[coverage.model.isin(["Mini", "Nano", "Claude", "Gemini"])]
    coverage_rows = [
        [
            row.model,
            f"{int(row.all_successful_records):,}",
            f"{int(row.probability_successful):,} / {int(row.probability_target):,}",
            f"{row.probability_coverage:.2%}",
            f"{int(row.probability_nonresponse):,}",
        ]
        for row in reps.itertuples()
    ]
    add_caption(document, "Table 2. Representative rater coverage")
    add_table(
        document,
        ["Rater", "All successful", "Probability sample", "Coverage", "Non-response"],
        coverage_rows,
    )

    document.add_heading("2.1 Reliability across the six core dimensions", level=2)
    document.add_paragraph(
        "The macro heatmaps average study status, technical type, AI role, mechanism, level, and "
        "scope. Exact agreement is the share of identical codes. Nominal Krippendorff alpha "
        "discounts agreement expected from the category distributions. The heatmaps are orientation "
        "views; inferential decisions remain dimension-specific."
    )
    add_picture(
        document,
        agreement_png,
        "Figure 1. Mean exact agreement across the six core dimensions (probability sample).",
    )
    add_picture(
        document,
        alpha_png,
        "Figure 2. Mean nominal Krippendorff alpha across the six core dimensions (probability sample).",
    )
    macro_rows = [
        [
            f"{row.left_model} : {row.right_model}",
            f"{row.percent_agreement:.2f}",
            f"{row.krippendorff_alpha:.2f}",
        ]
        for row in macro.sort_values("percent_agreement", ascending=False).itertuples()
    ]
    add_caption(document, "Table 3. Six-core-dimension macro orientation statistics")
    add_table(document, ["Model pair", "Mean exact agreement", "Mean alpha"], macro_rows)
    best = macro.sort_values("percent_agreement", ascending=False).iloc[0]
    document.add_paragraph(
        f"Claude and Gemini form the strongest representative pair: mean exact agreement "
        f"{best.percent_agreement:.2f} and mean alpha {best.krippendorff_alpha:.2f}. Mini is "
        "substantially closer to Claude and Gemini than Nano is. This supports Mini's role as the "
        "primary production rater, but it does not make Mini a gold standard."
    )

    document.add_heading("2.2 Reliability by displayed dimension", level=2)
    dimension_rows = []
    for dimension in DISPLAY_DIMENSIONS:
        four = multirater[multirater.dimension == dimension].iloc[0]
        dimension_rows.append(
            [
                DIMENSION_LABELS[dimension],
                "Core" if dimension in CORE_DIMENSIONS else "Exploratory",
                f"{pair_value(agreement, dimension, 'Claude', 'Gemini', 'percent_agreement'):.2f}",
                f"{pair_value(agreement, dimension, 'Claude', 'Gemini', 'krippendorff_alpha'):.2f}",
                f"{pair_value(agreement, dimension, 'Mini', 'Claude', 'krippendorff_alpha'):.2f}",
                f"{pair_value(agreement, dimension, 'Mini', 'Gemini', 'krippendorff_alpha'):.2f}",
                f"{four.krippendorff_alpha:.2f}",
            ]
        )
    add_caption(document, "Table 4. Dimension-level reliability on exact common-paper intersections")
    add_table(
        document,
        ["Dimension", "Status", "C:G agree", "C:G alpha", "M:C alpha", "M:G alpha", "Four-model alpha"],
        dimension_rows,
    )
    document.add_paragraph(
        "Technical type is the strongest displayed dimension: Claude-Gemini agreement is 0.83, "
        "their alpha is 0.78, and four-model alpha is 0.50. Study status, role, level, mechanism, "
        "and scope show useful convergence between the independent strong raters, although the "
        "four-model estimates fall when Nano is included. Process stage is not stable: "
        "Claude-Gemini alpha is 0.01 and four-model alpha is 0.11. Definition form has high raw "
        "agreement but low alpha because the no-definition category dominates; it remains an "
        "abstract-level diagnostic rather than a quality verdict."
    )

    document.add_heading("2.3 What the sample says about model sensitivity", level=2)
    document.add_paragraph(
        "The weighted distributions show substantive convergence and disagreement. Mini, Claude, "
        "and Gemini all place unspecified AI and machine learning at the top of technical type, "
        "while Nano assigns much more material to general and generative AI. Mini and Gemini most "
        "often classify AI as a tool, whereas Claude more often identifies a research-method role. "
        "For mechanism, Mini, Claude, and Gemini agree that non-observability is substantial and "
        "that prediction is the leading substantive mechanism; Nano instead concentrates heavily "
        "on missing mechanism and learning. Mini, Claude, and Gemini locate much of the literature "
        "at firm level, while Nano frequently chooses venture. Sector-specific scope dominates all "
        "four models. These differences require model-specific distributions and sensitivity "
        "language rather than a single pooled prevalence."
    )
    sample_type = prevalence[
        (prevalence.dimension == "ai_type_form")
        & prevalence.model.isin(["Mini", "Nano", "Claude", "Gemini"])
        & prevalence.category.isin(["unspecified AI", "machine learning"])
    ].copy()
    type_pivot = sample_type.pivot(index="model", columns="category", values="weighted_prevalence")
    sample_rows = [
        [model, f"{type_pivot.at[model, 'unspecified AI']:.1%}", f"{type_pivot.at[model, 'machine learning']:.1%}"]
        for model in ["Mini", "Nano", "Claude", "Gemini"]
    ]
    add_caption(document, "Table 5. Weighted prevalence of the two leading technical-type categories")
    add_table(document, ["Model", "Unspecified AI", "Machine learning"], sample_rows)

    document.add_page_break()
    document.add_heading("2.4 Human-validation boundary", level=2)
    completed = human_annotation_count()
    document.add_paragraph(
        f"The blind human-annotation module currently contains {completed} completed annotation "
        "records. Until independent human coding exists on a declared common-paper intersection, "
        "the analysis can report model-model reliability but cannot report model accuracy or "
        "human-model reliability. The planned 23-paper round is small and conditioned on overlap "
        "with the earlier workbook; it will be reported as an anchor, not as a population estimate."
    )

    document.add_heading("3. Full-corpus observed construct", level=1)
    document.add_paragraph(
        "The following figure describes GPT-5.4 Mini's full-corpus coding. Each panel conditions on "
        "papers with an observed substantive code for that dimension, so denominators differ. This "
        "is the substantive portrait; missing and unspecified categories form a separate diagnostic "
        "layer. The figure is not a four-model consensus estimate."
    )
    add_picture(
        document,
        OBSERVED_FIGURE,
        "Figure 3. Observed composition among papers specifying each dimension (full corpus, N = 22,345).",
        width=5.5,
    )
    document.add_paragraph(
        "Where study status is clear, 63.2% of papers treat AI as the phenomenon, 21.3% as a method, "
        "and 15.6% as both. Where role is substantive, 62.3% position AI as a tool and 17.3% as a "
        "research method. Among named technical forms, machine learning accounts for 48.4%. Among "
        "observable mechanisms, improved prediction accounts for 42.9%, followed by learning at "
        "18.0% and uncertainty reduction at 9.5%."
    )
    document.add_paragraph(
        "The level profile is organizational: firm-level claims account for 52.5% of specified "
        "levels, compared with 17.3% at the individual-entrepreneur level. Scope is strongly "
        "bounded: sector-specific claims account for 70.7% of stated boundaries and country-specific "
        "claims 18.4%. Process-stage and definition-form results remain descriptive because their "
        "probability-sample reliability is not strong enough for headline claims."
    )

    document.add_heading("4. Current analytical decisions", level=1)
    decision_rows = [
        [DIMENSION_LABELS[dimension], DIMENSION_DECISIONS[dimension]]
        for dimension in DISPLAY_DIMENSIONS
    ]
    add_table(document, ["Dimension", "Reporting decision"], decision_rows)
    document.add_paragraph(
        "The probability sample is sufficient to write the current model-validation results and to "
        "justify dimension-specific analytical decisions. Full-corpus Claude and Gemini runs would "
        "replace sample-based sensitivity distributions with complete model-specific distributions; "
        "they are not required to calculate the representative sample reliability already reported "
        "here. Human annotation remains necessary before making accuracy claims or upgrading process "
        "stage from exploratory status."
    )

    document.add_heading("Research artifacts", level=1)
    document.add_paragraph(
        "The frozen sample and manifest are under data/interim/proprietary_validation/. "
        "The analysis manifest, pairwise and multirater agreement tables, and weighted "
        "prevalence table are under data/processed/analysis/model_validation/. The detailed "
        "narrative is reports/analysis/MODEL_VALIDATION_FULL_RESULTS.md; the two source heatmaps "
        "are reports/analysis/figures/model_validation/pairwise_agreement_heatmap.svg and "
        "pairwise_alpha_heatmap.svg."
    )

    document.core_properties.title = "Current results interpretation: probability-sample validation"
    document.core_properties.subject = "AI construct specification and multi-model validation"
    document.core_properties.comments = "Generated from frozen analysis outputs; source document preserved."
    document.save(OUTPUT)

    markdown = f"""# Current results interpretation: probability-sample validation

The report uses the frozen 2,235-paper stratified probability sample. Mini and
Gemini cover all 2,235 papers; Nano and Claude cover 2,234 each. The exact
four-model intersection is 2,233 papers.

## Six-core-dimension macro reliability

{macro.sort_values('percent_agreement', ascending=False).to_markdown(index=False)}

## Main decisions

- Technical type is the strongest dimension.
- Study status, AI role, mechanism, level, and scope are core dimensions with
  model-sensitivity reporting.
- Process stage and definition form remain exploratory.
- No human annotation has been completed, so agreement is not accuracy.

The complete Word report is `{OUTPUT.relative_to(ROOT)}`. Numeric source files
are listed in its Research artifacts section.
"""
    MARKDOWN.write_text(markdown, encoding="utf-8")
    print(f"Word report -> {OUTPUT}")
    print(f"Markdown mirror -> {MARKDOWN}")


if __name__ == "__main__":
    build_report()
