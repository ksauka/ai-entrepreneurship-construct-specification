"""Apply the verified agreement decomposition to the KS manuscript and appendix.

The two DOCX files contain researcher edits that are not regenerated from the
manuscript builders. This updater therefore changes only the identified
methodology paragraphs and Appendix A6 tables, and fails if the expected
document structure is not present.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
MODEL_TABLES = ROOT / "reports/analysis/tables/model_validation"
IRR_SUMMARY = MODEL_TABLES / "full_corpus_pairwise_irr_core_summary.csv"
IRR_DIMENSIONS = MODEL_TABLES / "full_corpus_pairwise_irr_dimensions.csv"
IRR_CONSENSUS = MODEL_TABLES / "full_corpus_dimension_consensus.csv"

DIMENSION_ORDER = {
    "ai_method_or_phenomenon": 0,
    "ai_type_form": 1,
    "ai_role_function": 2,
    "ai_mechanism_analysis": 3,
    "level_of_analysis": 4,
    "scope_conditions": 5,
    "entrepreneurial_process_stage": 6,
    "definition_construct_clarity": 7,
}

RESEARCH_DESIGN_PARAGRAPHS = [
    (
        "This study uses theory elaboration to examine how artificial intelligence (AI) is specified across "
        "entrepreneurship research and how these specifications compare with the wider business and management "
        "literature. The research design combines systematic corpus construction, large-scale construct-"
        "specification coding, horizontal and vertical contrasting, structuring, topic modeling, and systematic "
        "close reading, and the resulting analytical states are operationalized in an interactive evidence "
        "platform. The Preferred Reporting Items for Systematic Reviews and Meta-Analyses (PRISMA) 2020 framework "
        "guided the search, screening, eligibility assessment, and inclusion decisions (Page et al., 2021). Four "
        "Scopus queries produced the broad business and management corpus of 22,345 papers and the Core "
        "entrepreneurship, Additional entrepreneurship, Combined entrepreneurship, and Financial Times 50 (FT50) "
        "analytical populations."
    ),
    (
        "Each paper was coded from its title, abstract, and author keywords by large language models working from "
        "a fixed construct-specification instrument of eight dimensions, and the models used, the coding procedure, "
        "and the cross-model agreement analysis are reported in Sections 3.4 and 3.5. Study status distinguishes AI "
        "as the substantive phenomenon, a research method, both, or unclear, while the remaining seven dimensions "
        "record technical AI type, AI role, observable mechanism, level of analysis, entrepreneurial process stage, "
        "scope conditions, and definition clarity. Full text, index terms, citation data, and topic assignments were "
        "excluded from the coding evidence, hence a missing value denotes non-observability in the title, abstract, "
        "and author keywords rather than absence from the full paper. Every assigned value was linked to supporting "
        "text, an evidence status of stated, inferred, or absent, and a dimension-specific confidence value, and one "
        "primary value was retained per dimension. Study status, technical AI type, AI role, mechanism, level of "
        "analysis, and scope conditions form the six core dimensions used in the main construct-specification and "
        "cross-model agreement analyses. Process stage and definition clarity remain part of the instrument and are "
        "reported throughout the study, but their interpretation is exploratory because cross-model agreement is "
        "substantially weaker for both dimensions than for the six core dimensions, for reasons set out in Section "
        "3.5 and Appendix A6."
    ),
    (
        "Following Fisher and Aguinis (2017), construct specification establishes what AI represents in each paper, "
        "horizontal contrasting compares how the same construct dimensions vary across business domains, vertical "
        "contrasting examines whether the theoretical meaning of AI changes across levels of analysis, and "
        "structuring identifies recurring relations among the eight dimensions. Table 1 links each theory-elaboration "
        "tactic to its analytical question, population, and output. These analyses describe observed patterns in a "
        "cross-sectional research corpus and do not establish temporal or causal sequences."
    ),
    (
        "Topic modeling was run after construct-specification coding and was kept analytically separate so that the "
        "discovered topics could not influence the instrument. The topics organize the corpus into navigable research "
        "conversations and are used as reading sites rather than as theoretical categories, while systematic close "
        "reading supported researcher-led comparison across these conversations to identify recurring theoretical "
        "insights for entrepreneurship research. The reading examined how relations among technical type, AI role, "
        "mechanism, level of analysis, process stage, and scope conditions changed across topic areas and business "
        "domains, and how these patterns informed the construct-clarification framework and the entrepreneurship "
        "implications developed in the paper."
    ),
    (
        "The interactive platform operationalizes the same theory-elaboration design by linking construct "
        "distributions, domain and level contrasts, recurring relations, topic areas, and theoretical interpretations "
        "to the papers and coding evidence on which they are based. Search strings, corpus-construction rules, coding "
        "specifications, and evidence boundaries are reported in Supplementary Appendix A1; platform reproduction "
        "procedures in Appendix A2; cross-model agreement and validation in Appendix A6; topic-model configuration "
        "and selection in Appendix A8; and systematic close-reading procedures in Appendix A9."
    ),
]

THEORY_ELABORATION_TABLE = [
    [
        "Construct specification",
        "What does each study mean by AI?",
        "Core, Additional, and Combined entrepreneurship",
        "Eight-dimensional construct portrait and AI type by role",
    ],
    [
        "Horizontal contrasting",
        "Does the same dimension vary across business domains?",
        "Full-corpus domains; FT50 restriction",
        "Within-domain matrix and percentage-point contrasts",
    ],
    [
        "Vertical contrasting",
        "Does AI's theoretical meaning change across levels?",
        "Combined entrepreneurship",
        "Study status, role, type, mechanism, stage, and scope by level",
    ],
    [
        "Structuring",
        "Which theoretically interpretable relations recur?",
        "Combined entrepreneurship",
        "Pairwise relations and selected configurations with evidence papers",
    ],
]

AMERICAN_ENGLISH = {
    "organisational": "organizational",
    "Organisational": "Organizational",
    "organisation": "organization",
    "Organisation": "Organization",
    "topic-modelling": "topic-modeling",
    "Topic-Modelling": "Topic-Modeling",
    "modelling": "modeling",
    "Modelling": "Modeling",
    "optimisation": "optimization",
    "Optimisation": "Optimization",
    "organised": "organized",
    "Organised": "Organized",
    "operationalises": "operationalizes",
    "Operationalises": "Operationalizes",
    "conceptualised": "conceptualized",
    "Conceptualised": "Conceptualized",
    "summarise": "summarize",
    "Summarise": "Summarize",
    "analysed": "analyzed",
    "Analysed": "Analyzed",
    "normalised": "normalized",
    "Normalised": "Normalized",
    "artefact": "artifact",
    "Artefact": "Artifact",
    "labelled": "labeled",
    "Labelled": "Labeled",
    "towards": "toward",
    "Towards": "Toward",
    "emphasises": "emphasizes",
    "Emphasises": "Emphasizes",
    "emphasise": "emphasize",
    "Emphasise": "Emphasize",
    "judgement": "judgment",
    "Judgement": "Judgment",
    "specialisation": "specialization",
    "Specialisation": "Specialization",
    "recognisable": "recognizable",
    "Recognisable": "Recognizable",
}


def format_alpha(value: float) -> str:
    text = f"{value:.2f}"
    if text.startswith("-0"):
        return f"-{text[2:]}"
    return text[1:] if text.startswith("0") else text


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def find_paragraph_any(document: Document, prefixes: tuple[str, ...]) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefixes)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning with one of {prefixes!r}; found {len(matches)}"
        )
    return matches[0]


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def replace_paragraph_runs(
    paragraph: Paragraph,
    pieces: list[tuple[str, bool]],
) -> None:
    paragraph.clear()
    for text, italic in pieces:
        run = paragraph.add_run(text)
        run.italic = italic


def americanize_paragraph(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        updated = run.text
        for source, target in AMERICAN_ENGLISH.items():
            updated = updated.replace(source, target)
        if updated != run.text:
            run.text = updated


def americanize_document(document: Document, *, stop_at_references: bool = False) -> None:
    for paragraph in document.paragraphs:
        if stop_at_references and paragraph.text.strip() == "References":
            break
        americanize_paragraph(paragraph)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    americanize_paragraph(paragraph)


def insert_paragraph_before(reference: Paragraph, text: str, style: str | None = None) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addprevious(element)
    paragraph = Paragraph(element, reference._parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def find_table_by_first_header(document: Document, header: str):
    matches = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells and table.rows[0].cells[0].text.strip() == header
    ]
    if len(matches) > 1:
        raise RuntimeError(f"Expected at most one table beginning with {header!r}; found {len(matches)}")
    return matches[0] if matches else None


def replace_research_design(document: Document) -> None:
    heading = find_paragraph(document, "3.1 Research design")
    next_heading = find_paragraph(document, "3.2 ")
    paragraphs = document.paragraphs
    start = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is heading._p)
    end = next(index for index, paragraph in enumerate(paragraphs) if paragraph._p is next_heading._p)
    section = paragraphs[start + 1 : end]
    caption = next(
        (paragraph for paragraph in section if paragraph.text.strip().startswith("Table 1.")),
        None,
    )
    for paragraph in section:
        if caption is None or paragraph._p is not caption._p:
            remove_paragraph(paragraph)

    if caption is None:
        caption = insert_paragraph_before(
            next_heading,
            "Table 1. Theory-elaboration design and analytical questions",
            "Normal",
        )
    else:
        replace_paragraph(caption, "Table 1. Theory-elaboration design and analytical questions")

    for text in RESEARCH_DESIGN_PARAGRAPHS:
        insert_paragraph_before(caption, text, "Normal")

    table = find_table_by_first_header(document, "Tactic")
    if table is None:
        table = document.add_table(rows=1, cols=4)
        table.style = "Normal Table"
        next_heading._p.addprevious(table._tbl)
    replace_table_contents(
        table,
        ["Tactic", "Question", "Primary population", "Main output"],
        THEORY_ELABORATION_TABLE,
        8.0,
        [1.15, 1.75, 1.65, 1.95],
    )


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell(cell, text: str, *, bold: bool = False, font_size: float = 6.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(font_size)
    if bold:
        shade_cell(cell)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def replace_table_contents(
    table,
    headers: list[str],
    rows: list[list[str]],
    font_size: float,
    widths: list[float] | None = None,
) -> None:
    while len(table.columns) < len(headers):
        table.add_column(Inches(0.55))
    if len(table.columns) != len(headers):
        raise RuntimeError(
            f"Cannot safely reduce table from {len(table.columns)} to {len(headers)} columns"
        )
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, font_size=font_size)
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            set_cell(
                table.rows[row_index].cells[column_index],
                value,
                font_size=font_size,
            )
        prevent_row_split(table.rows[row_index])
    if widths is not None:
        if len(widths) != len(headers):
            raise ValueError("Table width count must match the header count")
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for column_index, width in enumerate(widths):
            table.columns[column_index].width = Inches(width)
            for row in table.rows:
                row.cells[column_index].width = Inches(width)


def update_manuscript() -> None:
    document = Document(MANUSCRIPT)
    replace_research_design(document)

    instrument = find_paragraph(document, "The instrument contains eight coded fields")
    replace_paragraph(
        instrument,
        "The instrument contains eight coded fields: the four-category study-status field, which records whether "
        "AI is the substantive phenomenon, the research method, both, or unclear, and seven construct-"
        "specification dimensions covering technical AI type, AI role, observable mechanism, level of analysis, "
        "entrepreneurial process stage, scope conditions, and definition clarity. Six dimensions form the core "
        "construct-specification and model-agreement record: study status, technical AI type, AI role, mechanism, "
        "level, and scope. Process stage and definition clarity remain exploratory dimensions for the distinct "
        "reliability reasons reported in Section 3.5 and Appendix A6. Exploratory does not mean excluded, invalid, "
        "or unavailable; both dimensions retain paper counts, distributions, evidence papers, and complete "
        "agreement estimates, but they are not included in the six-dimension summary averages and conclusions "
        "based on them are interpreted cautiously.",
    )

    boundary = find_paragraph(document, "Mechanism is treated as substantive only")
    replace_paragraph(
        boundary,
        "Mechanism is treated as substantive only when the available evidence states what AI changes or enables. "
        "Process stage remains useful for examining where AI is positioned in an entrepreneurial or organisational "
        "process, while definition clarity records whether a definitional signal is observable and, where present, "
        "the form that signal takes in the title, abstract, or author keywords. Definition clarity cannot establish "
        "whether the full paper defines AI adequately. Both exploratory dimensions remain part of the instrument "
        "and platform, but findings based on them are labelled exploratory.",
    )

    validation = find_paragraph(document, "Generative Pre-trained Transformer (GPT)-5.4 Mini")
    replace_paragraph(
        validation,
        "Generative Pre-trained Transformer (GPT)-5.4 Mini was selected as the primary population coder, GPT-4.1 "
        "Nano as the baseline and sensitivity coder, and Claude Sonnet 5 and Gemini 3.1 Pro Preview as independent "
        "cross-provider coders. The four proprietary coders supplied eight dimension-level estimates on one common "
        "21,930-paper intersection. Overall agreement heatmaps and pair summaries average the six core dimensions "
        "only; the individual estimates for all eight dimensions remain the analytical record. Claude-Gemini "
        "produced the highest mean convergence across the six core dimensions (71.92% exact agreement; nominal "
        "α = .59), followed by Mini-Gemini (65.72%; α = .51) and Mini-Claude (65.02%; α = .50). The designated "
        "primary coder was therefore less convergent with either independent capable model than those two models "
        "were with each other; the primary designation is justified by the preregistered sequence and direct "
        "Gemini re-estimation rather than by convergence.",
    )
    paragraphs = document.paragraphs
    validation_index = next(
        index for index, paragraph in enumerate(paragraphs) if paragraph._p is validation._p
    )
    decomposition = next(
        (
            paragraph
            for paragraph in paragraphs
            if paragraph.text.strip().startswith("Agreement was decomposed because")
        ),
        None,
    )
    detail = next(
        (
            paragraph
            for paragraph in paragraphs
            if paragraph.text.strip().startswith(
                ("For Claude-Gemini, conditional alpha", "For Claude-Gemini, conditional α")
            )
        ),
        None,
    )
    if decomposition is None or detail is None:
        blanks = [p for p in paragraphs[validation_index + 1 :] if not p.text.strip()][:2]
        if len(blanks) != 2:
            raise RuntimeError("Expected two blank validation paragraphs before topic modelling")
        decomposition, detail = blanks
    replace_paragraph(
        decomposition,
        "Agreement was decomposed because a full-category comparison combines two decisions: whether the "
        "available evidence supports an observed category and, if so, which category applies. Full-category "
        "agreement therefore retains unobserved values as categories; observability agreement collapses ratings "
        "to observed versus unobserved; and conditional-category agreement compares categories only where both "
        "models observed evidence. The final layer excludes disagreements about observability and must be read "
        "alongside the observability result. Applied across all eight dimensions, this decomposition distinguishes "
        "disagreement about presence from disagreement about category rather than assuming one common failure mode.",
    )
    replace_paragraph_runs(
        detail,
        [
            (
                "For Claude-Gemini, conditional α across the six core dimensions ranges from .54 to .81. "
                "Mechanism has the thinnest jointly observed core base (6,676 papers, compared with 12,819 for "
                "the next smallest core dimension, technical type), but retains 74.07% conditional exact "
                "agreement and α = .63. Process stage shows a different pattern: observability α is .001, "
                "whereas conditional agreement among 3,298 jointly observed papers is 57.13% and α = .46. "
                "Definition clarity remains weaker at both signal detection (observability α = .17) and "
                "conditional classification among 1,252 jointly observed papers (56.87%; α = .26). On the "
                "separate four-model unanimity base, 11,448 of the 11,497 unanimous definition classifications are ",
                False,
            ),
            ("no definition", True),
            (
                " codes, while 6,019 of the 6,176 unanimous mechanism classifications are ",
                False,
            ),
            ("mechanism missing", True),
            (
                " codes. These results establish coding convergence rather than accuracy. Complete model coverage, "
                "full-category, observability, and conditional-category estimates, probability-sample checks, and "
                "three-model and four-model convergence counts are reported in Supplementary Tables A6.1-A6.5.",
                False,
            ),
        ],
    )

    contribution = find_paragraph(document, "The construct contribution is not a new universal definition")
    replace_paragraph(
        contribution,
        "The construct contribution is not a new universal definition of AI. It is an eight-dimensional framework "
        "for establishing the theoretical identity of an AI-related claim. Six dimensions form the core record: "
        "study status distinguishes AI as a phenomenon, method, or both; technical AI type states what system is "
        "named; AI role states what theoretical work AI performs; mechanism states what it changes or enables; "
        "level states where the relation operates; and scope states the boundary or organisational condition. "
        "Entrepreneurial process stage and definition clarity remain part of the framework as exploratory "
        "dimensions. Their exploratory status reflects distinct cross-model measurement limits: process-stage "
        "disagreement is concentrated in whether a stage is observable, whereas definition clarity remains weak "
        "at both signal detection and classification of observed definitional form. Neither dimension is excluded "
        "from paper-level inspection, distributions, comparisons, or downloads. The dimensions show whether "
        "studies are theoretically comparable or merely share the same technical language.",
    )

    status_result = find_paragraph(document, "Among the 1,497 entrepreneurship papers")
    if "135 papers" not in status_result.text:
        replace_paragraph(
            status_result,
            status_result.text.rstrip()
            + " The remaining 135 papers have an unclear study status; they remain in unconditioned full views "
            "but are not assigned to the phenomenon, method, or both profiles.",
        )

    horizontal = find_table_by_first_header(document, "Domain")
    if horizontal is None:
        raise RuntimeError("Could not find the horizontal-contrast results table")
    source = pd.read_csv(ROOT / "reports/analysis/tables/contrasting/horizontal_domain_contrast_full_corpus.csv")
    source_headers = [cell.text.strip() for cell in horizontal.rows[0].cells]
    source_index = {header: index for index, header in enumerate(source_headers)}
    required_headers = {
        "Domain",
        "Dimension",
        "Category",
        "Observed denominator",
        "Within-domain share",
        "Difference from full corpus",
    }
    if not required_headers.issubset(source_index):
        raise RuntimeError(
            "Horizontal-contrast table is missing required columns: "
            + ", ".join(sorted(required_headers - set(source_index)))
        )
    horizontal_rows: list[list[str]] = []
    for row in horizontal.rows[1:]:
        domain = row.cells[source_index["Domain"]].text.strip()
        dimension = row.cells[source_index["Dimension"]].text.strip()
        category = row.cells[source_index["Category"]].text.strip()
        denominator = row.cells[source_index["Observed denominator"]].text.strip()
        within_share = row.cells[source_index["Within-domain share"]].text.strip()
        difference = row.cells[source_index["Difference from full corpus"]].text.strip()
        match = source[
            source["domain_label"].eq(domain)
            & source["dimension_label"].eq(dimension)
            & source["category"].eq(category)
            & source["distribution"].eq("observed")
        ]
        if len(match) != 1:
            raise RuntimeError(
                f"Expected one horizontal source row for {domain!r}, {dimension!r}, {category!r}; found {len(match)}"
            )
        item = match.iloc[0]
        baseline = float(item["share"]) - float(item["percentage_point_difference"]) / 100
        horizontal_rows.append(
            [domain, dimension, category, denominator, f"{baseline:.1%}", within_share, difference]
        )
    replace_table_contents(
        horizontal,
        [
            "Domain",
            "Dimension",
            "Category",
            "Observed denominator",
            "Full-corpus baseline",
            "Within-domain share",
            "Difference from full corpus",
        ],
        horizontal_rows,
        6.7,
        [1.0, 0.85, 1.25, 0.75, 0.75, 0.75, 0.95],
    )

    structuring = find_paragraph(document, "The recurring-relation results are more coder-sensitive")
    if "thinnest jointly observed core base" not in structuring.text:
        replace_paragraph(
            structuring,
            structuring.text.rstrip()
            + " Mechanism also has the thinnest jointly observed base among the six core dimensions, so these "
            "frequency comparisons must be read with the conditional agreement estimates in Appendix A6.",
        )

    insight = find_paragraph(document, "Bottleneck relocation is the central entrepreneurship insight")
    replace_paragraph(
        insight,
        "Bottleneck relocation is the central entrepreneurship insight generated by the configuration analysis "
        "and systematic close reading. The dominant relations pair AI with informational mechanisms: tools improve "
        "prediction, support learning, reduce uncertainty, and expand search, while AI-supported judgment is located "
        "closer to individual action. More information does not settle the entrepreneurial decision. Chalmers et al. "
        "(2021) connect cheaper solution generation to a greater need for evaluation and selection; Ramoglou et al. "
        "(2026) describe opportunity search as adjudication among machine-generated possibilities; Rady et al. (2026) "
        "show the difficulty of separating plausible opportunities from hallucinated outputs; and de Véricourt and "
        "Gurkan (2026) show that superior prediction does not resolve verification and reliance. Read together, these "
        "patterns support the interpretation that expanded prediction, search, and generation can relocate the "
        "binding constraint toward plausibility judgment, selective commitment, calibrated reliance, and responsibility.",
    )

    agency_result = find_paragraph(document, "Agency remains an open frontier")
    replace_paragraph(
        agency_result,
        "Agency remains an open frontier. Nineteen of the 1,414 entrepreneurship papers with an observed AI-role "
        "code position AI as an actor or agent. The close-reading papers range from AI augmenting entrepreneurial "
        "cognition to AI as a teammate, relational nonhuman actor, or co-agent (Shepherd & Majchrzak, 2022; Murtinu "
        "& De Massis, 2025; Al-Bashrawi et al., 2026; Spurrier et al., 2025). The title, abstract, and keyword evidence "
        "can identify these actor-like positions, but it cannot establish how initiative, authority, decision rights, "
        "or responsibility are distributed in the underlying interaction. These cases therefore identify a "
        "theoretical boundary for full-text and claim-level research rather than support a settled claim about "
        "autonomous AI agency.",
    )

    discussion = find_paragraph_any(
        document,
        (
            "The construct-clarification analysis reveals",
            "The construct-clarification analysis and systematic close reading generate",
        ),
    )
    replace_paragraph(
        discussion,
        "The construct-clarification analysis and systematic close reading generate a bottleneck-relocation "
        "interpretation: AI-related entrepreneurship claims repeatedly shift the constraint rather than simply "
        "remove it. AI can expand search, generate options, improve prediction, and reduce some information costs, "
        "but the papers linking tools to judgment, uncertainty, and experimentation indicate that the entrepreneurial "
        "problem can move toward evaluating possibilities, assessing plausibility, calibrating reliance, selecting "
        "what deserves commitment, and retaining responsibility for action. In this interpretation, information "
        "abundance does not remove uncertainty; it changes where entrepreneurial judgment is required.",
    )

    agency_discussion = find_paragraph(document, "Agency becomes theoretically important")
    if "19" not in agency_discussion.text:
        replace_paragraph(
            agency_discussion,
            "Agency becomes theoretically important in the 19 observed actor/agent cases, where AI is positioned "
            "as an actor-like participant, teammate, co-agent, or system that alters judgment and interaction. The "
            "present instrument can identify these positions but cannot determine whether the human retains "
            "initiative, decision authority, authorship, or responsibility. Agency is therefore not part of the "
            "coded construct-clarification framework. It is a frontier problem generated by the framework, requiring "
            "full-text analysis and designs that observe interaction, decision rights, intervention, and "
            "responsibility directly. Treating it in this way preserves the theoretical question without presenting "
            "abstract-level inference as measurement.",
        )

    limitations = find_paragraph(document, "The evidence is restricted to titles")
    replace_paragraph(
        limitations,
        limitations.text.replace(
            "blind human validation of the current dimensions remains incomplete",
            "no blind human validation was available for the current eight-dimensional instrument",
        ),
    )

    americanize_document(document, stop_at_references=True)

    output = MANUSCRIPT.with_suffix(".tmp.docx")
    document.save(output)
    Document(output)
    output.replace(MANUSCRIPT)


def update_supplement() -> None:
    document = Document(SUPPLEMENT)
    intro_note = find_paragraph(document, "Note. All eight dimensions were coded")
    replace_paragraph(
        intro_note,
        "Note. All eight dimensions were coded for every successfully processed paper and remain available for "
        "filtering, comparison, evidence inspection, and download. Study status, technical AI type, AI role, "
        "mechanism, level of analysis, and scope conditions form the six core dimensions used in the main "
        "construct-specification and cross-model agreement analyses. Process stage and definition clarity remain "
        "part of the instrument but are interpreted as exploratory because cross-model agreement is substantially "
        "weaker for both, for distinct reasons decomposed in Appendix A6. Exploratory does not mean excluded, "
        "invalid, or unavailable. Agency is not a ninth coded dimension because agency allocation requires claim-"
        "level, full-text, or interaction-level evidence.",
    )

    conditioned = find_paragraph(document, "The Combined entrepreneurship population contains 1,497 papers")
    replace_paragraph(
        conditioned,
        "The Combined entrepreneurship population contains 1,497 papers with a clear study status: 824 treat AI "
        "as the substantive phenomenon, 385 as a research method, and 288 as both. The remaining 135 papers have "
        "an unclear study status. They remain in the unconditioned full view but are not assigned to these three "
        "profiles. The tables below report the observed denominator separately for every outcome dimension because "
        "observability differs across titles, abstracts, and author keywords.",
    )

    a6_intro = find_paragraph(document, "Construct-specification distributions retain each model's")
    replace_paragraph(
        a6_intro,
        "Construct-specification distributions retain each model's actual successful-paper denominator. Cross-"
        "model agreement uses a separate rule: every pair is restricted to the same 21,930-paper intersection "
        "shared by GPT-5.4 Mini, GPT-4.1 Nano, Claude Sonnet 5, and Gemini 3.1 Pro Preview. Exact agreement and "
        "nominal Krippendorff's α are reported for all eight dimensions. Because full-category agreement "
        "combines the observability decision with the category decision, Table A6.3 additionally reports binary "
        "observability agreement and conditional-category agreement where both models observed evidence. "
        "Arithmetic means across the six core dimensions remain orientation summaries rather than omnibus "
        "reliability coefficients. Llama 3.2 3B Instruct and Gemma 4 31B are partial local stress-test raters and "
        "are not included in this balanced four-model matrix.",
    )
    a6_status = find_paragraph(document, "All eight dimensions remain part of the analytical instrument")
    replace_paragraph(
        a6_status,
        "All eight dimensions remain part of the analytical instrument. Study status, technical AI type, AI role, "
        "mechanism, level of analysis, and scope conditions form the six core dimensions used in the main analyses. "
        "Core does not mean that these dimensions behave uniformly or achieve perfect reliability. Process stage "
        "and definition clarity remain exploratory dimensions for different empirical reasons: process-stage "
        "weakness is concentrated in agreement about whether a stage is observable, while definition clarity "
        "remains weak at both signal detection and classification of the observed definitional form. All eight "
        "dimension-level results remain reported, and the overall heatmaps exclude only the two exploratory fields.",
    )

    summary = pd.read_csv(IRR_SUMMARY).sort_values(
        ["mean_krippendorff_alpha", "mean_exact_agreement"], ascending=False
    )
    summary_rows = [
        [
            row.model_pair,
            f"{int(row.balanced_common_papers):,}",
            f"{row.mean_exact_agreement:.2%}",
            format_alpha(row.mean_krippendorff_alpha),
        ]
        for row in summary.itertuples()
    ]
    replace_table_contents(
        document.tables[18],
        ["Model pair", "Balanced papers", "Mean exact agreement", "Mean nominal α"],
        summary_rows,
        7.2,
    )

    dimensions = pd.read_csv(IRR_DIMENSIONS)
    current_pair_order = list(
        dict.fromkeys(row.cells[0].text.strip() for row in document.tables[19].rows[1:])
    )
    dimensions["pair_order"] = dimensions["model_pair"].map(
        {value: index for index, value in enumerate(current_pair_order)}
    )
    dimensions["dimension_order"] = dimensions["dimension"].map(DIMENSION_ORDER)
    dimensions = dimensions.sort_values(["pair_order", "dimension_order"])
    dimension_rows = [
        [
            row.model_pair,
            row.dimension_label,
            row.classification,
            f"{int(row.comparable_papers):,}",
            f"{row.exact_agreement:.2%}",
            format_alpha(row.krippendorff_alpha),
            f"{row.observability_exact_agreement:.2%}",
            format_alpha(row.observability_krippendorff_alpha),
            f"{int(row.jointly_observed_papers):,}",
            f"{row.observed_category_exact_agreement:.2%}",
            format_alpha(row.observed_category_krippendorff_alpha),
        ]
        for row in dimensions.itertuples()
    ]
    replace_table_contents(
        document.tables[19],
        [
            "Model pair",
            "Dimension",
            "Use",
            "Full papers",
            "Full exact",
            "Full α",
            "Observability exact",
            "Observability α",
            "Both observed",
            "Conditional exact",
            "Conditional α",
        ],
        dimension_rows,
        5.4,
        [1.45, 1.05, 0.48, 0.60, 0.58, 0.50, 0.66, 0.60, 0.62, 0.66, 0.60],
    )
    a63_note = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith("Note. Full-category agreement")
        ),
        None,
    )
    if a63_note is None:
        a63_note = next(
            paragraph
            for index, paragraph in enumerate(document.paragraphs)
            if index > 55 and not paragraph.text.strip()
        )
    replace_paragraph(
        a63_note,
        "Note. Full-category agreement retains each dimension's unobserved value as a category. Observability "
        "agreement collapses ratings to observed versus unobserved. Conditional-category agreement uses only "
        "papers for which both models assigned an observed category, so it excludes disagreement about "
        "observability and must be read alongside the observability result. Across Claude-Gemini, conditional "
        "α for the six core dimensions ranges from .54 to .81. Mechanism has the thinnest jointly observed "
        "core base (6,676 papers; technical type is next at 12,819), but retains 74.07% conditional exact agreement "
        "and α = .63. Process stage has observability α = .001 and conditional α = .46; definition "
        "clarity has observability α = .17 and conditional α = .26.",
    )

    consensus = pd.read_csv(IRR_CONSENSUS)
    consensus_rows = [
        [
            row.dimension_label,
            row.classification,
            (
                f"{int(row.preferred_trio_agreement_papers):,} "
                f"({row.preferred_trio_agreement_share:.2%}); unobserved "
                f"{int(row.preferred_trio_unobserved_agreement_papers):,}; observed "
                f"{int(row.preferred_trio_observed_agreement_papers):,}"
            ),
            (
                f"{int(row.all_four_agreement_papers):,} "
                f"({row.all_four_agreement_share:.2%}); unobserved "
                f"{int(row.all_four_unobserved_agreement_papers):,}; observed "
                f"{int(row.all_four_observed_agreement_papers):,}"
            ),
        ]
        for row in consensus.itertuples()
    ]
    replace_table_contents(
        document.tables[20],
        [
            "Dimension",
            "Use",
            "Mini + Claude + Gemini: total; unobserved; observed",
            "All four models: total; unobserved; observed",
        ],
        consensus_rows,
        6.7,
    )
    convergence = find_paragraph(document, "The platform also supports paper-level convergence filters")
    replace_paragraph(
        convergence,
        "The platform also supports paper-level convergence filters. General agreement requires any two or more "
        "models to assign the exact selected value. The three-model criterion requires agreement among GPT-5.4 "
        "Mini, Claude Sonnet 5, and Gemini 3.1 Pro Preview, while the stricter column requires all four models. "
        "Table A6.4 partitions each total into unobserved-category and observed-category agreement so that a high "
        "unanimity count cannot be mistaken for agreement about substantive form. These counts identify high-"
        "convergence papers for evidence inspection and do not replace the primary codes with a consensus "
        "classification.",
    )
    probability = find_paragraph(document, "A prospective 2,235-paper probability sample")
    unanimity_note = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith("Note. Table A6.4 changes")
        ),
        None,
    )
    if unanimity_note is None:
        unanimity_note = insert_paragraph_before(probability, "")
    replace_paragraph(
        unanimity_note,
        "Note. Table A6.4 changes from pairwise Claude-Gemini statistics to three-model and four-model unanimity. "
        "On the four-model base, 11,448 of the 11,497 unanimous definition-clarity classifications are no-"
        "definition codes and only 49 are unanimous observed categories. Mechanism unanimity is similarly "
        "absence-dominated: 6,019 of 6,176 unanimous classifications are mechanism-missing codes and 157 are "
        "unanimous observed mechanisms. Mechanism nevertheless retains materially stronger conditional pairwise "
        "agreement where both Claude and Gemini observe a mechanism, as reported in Table A6.3.",
    )

    americanize_document(document)

    output = SUPPLEMENT.with_suffix(".tmp.docx")
    document.save(output)
    Document(output)
    output.replace(SUPPLEMENT)


def main() -> None:
    update_manuscript()
    update_supplement()
    print(f"Updated {MANUSCRIPT}")
    print(f"Updated {SUPPLEMENT}")


if __name__ == "__main__":
    main()
