"""Apply the audited ASJC domain aggregation to the KS manuscript files.

This updater changes only the analytical-population table, its explanatory
text, the selected horizontal-contrast table and figure, and a new traceability
appendix. It reads the frozen domain manifest and horizontal-contrast table so
the DOCX values cannot drift from the platform artifacts.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from aecsp.api.graph_service import GraphService


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
MANIFEST = (
    ROOT
    / "data/processed/analysis/theory_elaboration/domains/business_domain_manifest.json"
)
HORIZONTAL = (
    ROOT
    / "reports/analysis/tables/contrasting/horizontal_domain_contrast_full_corpus.csv"
)
HORIZONTAL_FIGURE = (
    ROOT / "reports/analysis/figures/contrasting/horizontal_role_by_domain_with_ent.png"
)
RESIDUAL_SOURCES = (
    ROOT
    / "data/processed/analysis/theory_elaboration/domains/business_domain_residual_source_titles.csv"
)
PRIMARY_MODEL = "gpt-5.4-mini-2026-03-17"

ANALYTICAL_POPULATIONS = [
    ["Full corpus", "22,345", "558 source-title labels; 640 title-ISSN pairs"],
    ["Leading entrepreneurship journals", "646", "15 source-title labels; all papers retained"],
    ["Additional entrepreneurship", "986", "13 source-title labels; all papers retained"],
    ["Combined entrepreneurship", "1,632", "28 source-title labels; exact Leading-Additional union"],
    ["FT50 restriction", "438", "37 source-title labels; all papers retained"],
]

SELECTED_HORIZONTAL = [
    ("Management Science and Operations Research", "AI role", "AI as tool"),
    ("Marketing", "Mechanism", "transforms stakeholder interaction"),
    ("Management of Technology and Innovation", "Technical type", "generative AI"),
    ("Organization studies", "Mechanism", "alters judgment"),
    ("Finance", "Technical type", "machine learning"),
    ("Environmental and sustainability", "Mechanism", "improves prediction"),
    ("Leading entrepreneurship journals", "AI role", "AI as research method"),
    ("Additional entrepreneurship", "AI role", "AI as firm capability"),
]


def find_paragraph(document: Document, prefix: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def find_paragraph_any(document: Document, prefixes: tuple[str, ...]):
    matches = [
        p for p in document.paragraphs if p.text.strip().startswith(prefixes)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning with {prefixes!r}; found {len(matches)}"
        )
    return matches[0]


def find_table(document: Document, first_header: str):
    matches = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text.strip() == first_header
    ]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one table headed {first_header!r}; found {len(matches)}")
    return matches[0]


def replace_paragraph(paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell(cell, value: object, *, bold: bool = False, font_size: float = 7.0) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(value))
    run.bold = bold
    run.font.size = Pt(font_size)
    if bold:
        shade_cell(cell)


def replace_table(table, headers, rows, *, font_size=7.0, widths=None) -> None:
    while len(table.columns) < len(headers):
        table.add_column(Inches(0.55))
    if len(table.columns) != len(headers):
        raise RuntimeError(
            f"Cannot safely change {len(table.columns)} columns to {len(headers)}"
        )
    while len(table.rows) < len(rows) + 1:
        table.add_row()
    while len(table.rows) > len(rows) + 1:
        table._tbl.remove(table.rows[-1]._tr)
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, font_size=font_size)
    for row_index, values in enumerate(rows, start=1):
        for column_index, value in enumerate(values):
            set_cell(table.rows[row_index].cells[column_index], value, font_size=font_size)
        properties = table.rows[row_index]._tr.get_or_add_trPr()
        if properties.find(qn("w:cantSplit")) is None:
            properties.append(OxmlElement("w:cantSplit"))
    if widths:
        table.autofit = False
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for column_index, width in enumerate(widths):
            table.columns[column_index].width = Inches(width)
            for row in table.rows:
                row.cells[column_index].width = Inches(width)


def append_table(document: Document, caption: str, headers, rows, widths=None) -> None:
    document.add_paragraph(caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    replace_table(table, headers, rows, font_size=6.7, widths=widths)


def population_rows(manifest: dict) -> list[list[str]]:
    rows = list(ANALYTICAL_POPULATIONS)
    for item in manifest["domains"].values():
        rows.append(
            [
                item["label"],
                f"{int(item['papers']):,}",
                f"{int(item['represented_source_titles']):,} source-title labels",
            ]
        )
    return rows


def selected_horizontal_rows(service: GraphService, horizontal: pd.DataFrame):
    baselines = {}
    coverage = {}
    for dimension_id in ("ai_role", "mechanism", "technical_type"):
        result = service.theory_horizontal_contrast(
            PRIMARY_MODEL, dimension_id, "observed"
        )
        baselines[dimension_id] = result["baseline"]["denominator"]
        coverage[dimension_id] = result["baseline_domain_coverage"]
    dimension_ids = {
        "AI role": "ai_role",
        "Mechanism": "mechanism",
        "Technical type": "technical_type",
    }
    rows = []
    for domain, dimension, category in SELECTED_HORIZONTAL:
        match = horizontal[
            horizontal["domain_label"].eq(domain)
            & horizontal["dimension_label"].eq(dimension)
            & horizontal["category"].eq(category)
            & horizontal["distribution"].eq("observed")
        ]
        if len(match) != 1:
            raise RuntimeError(
                f"Expected one horizontal row for {domain}, {dimension}, {category}; found {len(match)}"
            )
        item = match.iloc[0]
        baseline_share = float(item["share"]) - float(
            item["percentage_point_difference"]
        ) / 100
        dimension_id = dimension_ids[dimension]
        rows.append(
            [
                domain,
                dimension,
                category,
                f"{int(item['denominator']):,}",
                f"{int(baselines[dimension_id]):,}",
                f"{baseline_share:.1%}",
                f"{float(item['share']):.1%}",
                f"{float(item['percentage_point_difference']):+.1f} pp",
            ]
        )
    return rows, coverage


def replace_horizontal_figure(document: Document) -> None:
    caption = find_paragraph(document, "Figure 5.")
    paragraphs = document.paragraphs
    index = next(i for i, paragraph in enumerate(paragraphs) if paragraph._p is caption._p)
    if index == 0:
        raise RuntimeError("Figure 5 has no preceding image paragraph")
    image_paragraph = paragraphs[index - 1]
    image_paragraph.clear()
    image_paragraph.add_run().add_picture(str(HORIZONTAL_FIGURE), width=Inches(6.45))


def remove_existing_a11(document: Document) -> None:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith("A11. ASJC Domain Aggregation")
    ]
    if not matches:
        return
    start = matches[0]._p
    body = document._body._body
    deleting = False
    for child in list(body):
        if child is start:
            deleting = True
        if deleting and child.tag != qn("w:sectPr"):
            body.remove(child)


def save_checked(document: Document, path: Path) -> None:
    temporary = path.with_suffix(".tmp.docx")
    document.save(temporary)
    Document(temporary)
    temporary.replace(path)


def update_manuscript(manifest: dict, horizontal_rows: list[list[str]]) -> None:
    document = Document(MANUSCRIPT)
    table3 = find_table(document, "Population or domain")
    replace_table(
        table3,
        ["Population or domain", "Papers retained", "Source-title coverage in this corpus"],
        population_rows(manifest),
        font_size=7.2,
        widths=[2.35, 0.85, 3.15],
    )
    validation = manifest["validation"]
    domain_count = int(validation["business_domain_count"])
    replace_paragraph(
        find_paragraph_any(
            document,
            (
                "Business-domain assignments were applied",
                "Business-domain assignments were rebuilt",
            ),
        ),
        "Business-domain assignments were rebuilt from the completed official Scopus All Science Journal "
        "Classification (ASJC) data for journals already represented in the corpus; no papers were retrieved "
        "or added to fill a domain. The explicit code-to-domain aggregation assigns "
        f"{validation['unique_assigned_papers']:,} of 22,345 papers ({validation['assigned_papers_percent']:.2f}%) "
        f"to at least one of {domain_count} selected ASJC-derived business domains, with "
        f"{validation['multi_domain_papers']:,} papers "
        "belonging to more than one domain. The remaining "
        f"{validation['residual_papers']:,} papers ({validation['residual_papers_percent']:.2f}%) all retain "
        "official ASJC codes but fall outside the selected domain rows. They remain in the full-corpus "
        "comparison baseline. The Leading, Additional, Combined entrepreneurship, and FT50 populations are "
        "registered journal-based analytical views; domain counts are therefore overlapping and not additive. "
        "The complete aggregation, residual audit, and baseline rule are reported in Supplementary Appendix A11.",
    )
    table8 = find_table(document, "Domain")
    replace_table(
        table8,
        [
            "Domain",
            "Dimension",
            "Category",
            "Domain observed n",
            "Full-corpus observed n",
            "Full-corpus share",
            "Within-domain share",
            "Difference from full corpus",
        ],
        horizontal_rows,
        font_size=5.8,
        widths=[1.0, 0.7, 1.05, 0.65, 0.75, 0.67, 0.72, 0.82],
    )
    replace_horizontal_figure(document)
    save_checked(document, MANUSCRIPT)


def update_supplement(
    manifest: dict,
    horizontal_rows: list[list[str]],
    baseline_coverage: dict,
) -> None:
    document = Document(SUPPLEMENT)
    table = find_table(document, "Population or domain")
    replace_table(
        table,
        ["Population or domain", "Papers retained", "Source-title coverage in this corpus"],
        population_rows(manifest),
        font_size=7.2,
        widths=[2.35, 0.85, 3.15],
    )
    validation = manifest["validation"]
    domain_count = int(validation["business_domain_count"])
    registered_group_count = int(validation["all_registered_group_count"])
    replace_paragraph(
        find_paragraph_any(
            document,
            ("The broad corpus is not an entrepreneurship-only dataset",),
        ),
        "The broad corpus is not an entrepreneurship-only dataset. Leading entrepreneurship journals contain 646 papers, "
        "Additional entrepreneurship contains 986 papers, and their exact union contains 1,632 papers. FT50 "
        "contains 438 papers and is used as a robustness and boundary restriction. The completed official Scopus "
        f"ASJC assignments place {validation['unique_assigned_papers']:,} papers in at least one of the "
        f"{domain_count} selected business domains. Because membership is multi-label, "
        f"{validation['multi_domain_papers']:,} papers occur in more than one domain and the row counts are not additive. "
        f"The {validation['residual_papers']:,}-paper residual lies outside the {domain_count} analytical domain rows but is not unclassified by Scopus; "
        "every residual paper retains official ASJC codes and remains in the full-corpus baseline.",
    )
    remove_existing_a11(document)
    document.add_paragraph(
        "A11. ASJC Domain Aggregation, Residual Coverage, and Comparison Baseline",
        style="Heading 1",
    )
    document.add_paragraph(
        "The business-domain analysis uses an explicit multi-label aggregation of the official Scopus ASJC "
        "classifications attached to each represented source. A paper inherits every selected domain whose "
        "registered code occurs among its source classifications. Broad codes that do not identify one selected "
        "domain are not forced into an arbitrary primary category. All primary horizontal domain rows are derived "
        "from official ASJC codes; no reviewed source-title overlay is included. No paper text, topic label, or "
        "model classification is used to assign a business domain."
    )
    mapping_rows = []
    for item in manifest["domains"].values():
        codes = "; ".join(
            f"{code} {label}" for code, label in item.get("asjc_codes", {}).items()
        )
        if not codes:
            raise RuntimeError(f"Active domain {item['label']} has no official ASJC code")
        mapping_rows.append(
            [item["label"], item["mapping_mode"].replace("_", " "), codes, item["rationale"]]
        )
    append_table(
        document,
        "Table A11.1. Explicit ASJC-code-to-business-domain aggregation",
        ["Domain", "Assignment mode", "Official ASJC code(s)", "Rationale"],
        mapping_rows,
        [1.15, 0.9, 2.6, 2.0],
    )
    coverage_rows = [
        [item["label"], f"{item['papers']:,}", f"{item['represented_source_titles']:,}", item["mapping_mode"].replace("_", " ")]
        for item in manifest["domains"].values()
    ]
    append_table(
        document,
        "Table A11.2. Corrected business-domain coverage in the retained corpus",
        ["Business domain", "Unique papers", "Represented source titles", "Assignment mode"],
        coverage_rows,
        [2.0, 1.0, 1.25, 1.55],
    )
    audit_rows = [
        ["Full retained corpus", "22,345", "100.00%", "Baseline includes every retained paper"],
        [f"Inside at least one of the {domain_count} selected business domains", f"{validation['unique_assigned_papers']:,}", f"{validation['assigned_papers_percent']:.2f}%", "Multi-label union; rows are not summed"],
        [f"Outside the {domain_count} selected business domains", f"{validation['residual_papers']:,}", f"{validation['residual_papers_percent']:.2f}%", "Official ASJC present; retained in baseline"],
        [f"Inside all {registered_group_count} registered groups", f"{validation['all_registered_group_unique_papers']:,}", f"{100 - validation['all_registered_group_residual_percent']:.2f}%", f"{domain_count} business domains plus Leading, Additional, and FT50"],
        [f"Outside all {registered_group_count} registered groups", f"{validation['all_registered_group_residual_papers']:,}", f"{validation['all_registered_group_residual_percent']:.2f}%", "Official ASJC present; retained in baseline"],
    ]
    append_table(
        document,
        "Table A11.3. Residual and baseline audit",
        ["Coverage set", "Unique papers", "Share of corpus", "Interpretation"],
        audit_rows,
        [2.3, 0.9, 0.9, 2.3],
    )
    document.add_paragraph(
        "Note. The full-corpus horizontal baseline is the complete eligible denominator for the selected coding "
        "model, study-status filter, and full or observed view. It is not restricted to papers assigned to one of "
        "the displayed domain rows. In an observed view, the residual is recalculated within the substantive-code "
        "denominator for that dimension."
    )
    append_table(
        document,
        "Table A11.4. Corrected selected horizontal contrasts with explicit baseline denominators",
        [
            "Domain",
            "Dimension",
            "Category",
            "Domain observed n",
            "Full-corpus observed n",
            "Full-corpus share",
            "Within-domain share",
            "Difference",
        ],
        horizontal_rows,
        [0.95, 0.68, 1.0, 0.62, 0.72, 0.65, 0.72, 0.65],
    )
    dimension_labels = {
        "ai_role": "AI role",
        "mechanism": "Mechanism",
        "technical_type": "Technical type",
    }
    observed_residual_rows = []
    for dimension_id, coverage in baseline_coverage.items():
        observed_residual_rows.append(
            [
                dimension_labels[dimension_id],
                f"{coverage['baseline_papers']:,}",
                f"{coverage['inside_selected_business_domains']:,}",
                f"{coverage['outside_selected_business_domains']:,}",
                f"{coverage['outside_selected_business_domains_percent']:.2f}%",
            ]
        )
    append_table(
        document,
        "Table A11.5. Domain residual within the observed baseline used in Table A11.4",
        ["Dimension", "Observed baseline", f"Inside {domain_count} domains", f"Outside {domain_count} domains", "Outside share"],
        observed_residual_rows,
        [1.3, 1.1, 1.15, 1.15, 0.95],
    )
    residual_sources = pd.read_csv(RESIDUAL_SOURCES, dtype=str, keep_default_na=False)
    residual_source_rows = [
        [
            row["Source title"],
            f"{int(row['papers']):,}",
            row["asjc_codes"],
            row["asjc_descriptions"],
        ]
        for _, row in residual_sources.head(15).iterrows()
    ]
    append_table(
        document,
        f"Table A11.6. Largest represented source titles outside the {domain_count} selected business domains",
        ["Source title", "Papers", "Official ASJC codes", "Official ASJC descriptions"],
        residual_source_rows,
        [2.35, 0.65, 1.1, 2.55],
    )
    outputs = manifest["outputs"]
    artifact_rows = [
        [manifest["aggregation"]["config_path"], manifest["aggregation"]["config_sha256"][:16], "Code-to-domain specification"],
        [outputs["paper_assignments"], outputs["paper_assignments_sha256"][:16], "Paper-to-domain assignments"],
        [outputs["represented_sources"], outputs["represented_sources_sha256"][:16], "Represented-source audit"],
        [outputs["residual_papers"], outputs["residual_papers_sha256"][:16], "Paper-level residual ledger"],
        [outputs["residual_sources"], outputs["residual_sources_sha256"][:16], "Residual source-title audit"],
    ]
    append_table(
        document,
        "Table A11.7. Frozen domain-analysis artifacts",
        ["Artifact", "Abbreviated SHA-256", "Purpose"],
        artifact_rows,
        [3.35, 1.35, 2.0],
    )
    save_checked(document, SUPPLEMENT)


def main() -> None:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    horizontal = pd.read_csv(HORIZONTAL)
    service = GraphService()
    rows, coverage = selected_horizontal_rows(service, horizontal)
    update_manuscript(manifest, rows)
    update_supplement(manifest, rows, coverage)
    print(f"Updated {MANUSCRIPT}")
    print(f"Updated {SUPPLEMENT}")


if __name__ == "__main__":
    main()
