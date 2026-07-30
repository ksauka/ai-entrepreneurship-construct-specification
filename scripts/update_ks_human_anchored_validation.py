"""Add completed human-anchored model validation to the KS supplement.

This is a surgical update of the author-edited supplementary DOCX. It replaces
the obsolete statement that human validation was unavailable, inserts a
complete audit table for every available model, and leaves all other appendix
material untouched.
"""

from __future__ import annotations

import hashlib
import json
import os
import shutil
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph

from aecsp.analytics.agreement import (
    krippendorff_alpha_nominal,
    pairwise_percent_agreement,
)
from aecsp.api.graph_service import IRR_UNOBSERVED_VALUES
from aecsp.human_annotation import (
    CORE_COLUMNS,
    DISPLAY_COLUMNS,
    HUMAN_DIMENSIONS,
    HumanAnnotationStore,
)
from aecsp.specification.analysis_columns import enrich_for_analysis
from aecsp.specification.llm_coder import PROTOCOL_ID, cache_key


ROOT = Path(__file__).resolve().parents[1]
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
BACKUP = (
    ROOT
    / "docs/ETP supplementary material july2026 ks.before-human-anchored-validation.docx"
)
OUTPUT_DIR = ROOT / "reports/analysis/tables/model_validation"
SUMMARY_CSV = OUTPUT_DIR / "human_anchored_model_validation_summary.csv"
DIMENSIONS_CSV = OUTPUT_DIR / "human_anchored_model_validation_dimensions.csv"
CONSENSUS_CSV = OUTPUT_DIR / "human_anchored_model_validation_consensus.csv"
MANIFEST_JSON = OUTPUT_DIR / "human_anchored_model_validation_manifest.json"

MODEL_ORDER = (
    "gpt-5.4-mini-2026-03-17",
    "gpt-4.1-nano-2025-04-14",
    "claude-sonnet-5",
    "gemini-3.1-pro-preview",
    "llama3.2",
    "gemma4:31b",
)
MODEL_LABELS = {
    "gpt-5.4-mini-2026-03-17": "GPT-5.4 Mini",
    "gpt-4.1-nano-2025-04-14": "GPT-4.1 Nano",
    "claude-sonnet-5": "Claude Sonnet 5",
    "gemini-3.1-pro-preview": "Gemini 3.1 Pro Preview",
    "llama3.2": "Llama 3.2",
    "gemma4:31b": "Gemma 4 31B",
}
MODEL_STATUS = {
    "gpt-5.4-mini-2026-03-17": "Complete human-sample intersection",
    "gpt-4.1-nano-2025-04-14": "Complete human-sample intersection",
    "claude-sonnet-5": "Complete human-sample intersection",
    "gemini-3.1-pro-preview": "Complete human-sample intersection",
    "llama3.2": "Diagnostic partial intersection",
    "gemma4:31b": "Diagnostic selective intersection",
}
DIMENSION_LABELS = {
    item["analysis_column"]: item["label"] for item in HUMAN_DIMENSIONS
}
DIMENSION_USE = {
    column: ("Core" if column in CORE_COLUMNS else "Exploratory")
    for column in DISPLAY_COLUMNS
}


def alpha_text(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    text = f"{float(value):.2f}"
    if text.startswith("-0"):
        return f"-{text[2:]}"
    return text[1:] if text.startswith("0") else text


def percent_text(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "n/a"
    return f"{float(value):.2%}"


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def insert_paragraph_before(
    reference: Paragraph,
    text: str,
    *,
    bold: bool = False,
    color: str | None = None,
    size: float | None = None,
    centered: bool = False,
    keep_with_next: bool = False,
) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addprevious(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = "Normal"
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.bold = bold
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return paragraph


def remove_between(start, end) -> None:
    """Remove XML siblings from start through the element before end."""

    current = start
    parent = current.getparent()
    while current is not None and current is not end:
        following = current.getnext()
        parent.remove(current)
        current = following
    if current is None:
        raise RuntimeError("The validation block did not terminate at Appendix A7")


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell(cell, text: object, *, bold: bool = False, size: float = 6.0) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if bold:
        shade_cell(cell)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)


def insert_table_before(
    document: Document,
    reference: Paragraph,
    headers: list[str],
    rows: list[list[object]],
    *,
    font_size: float,
    widths: list[float] | None = None,
) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, size=font_size)
    repeat_table_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell(row.cells[index], value, size=font_size)
        prevent_row_split(row)
    if widths is not None:
        table.autofit = False
        for index, width in enumerate(widths):
            table.columns[index].width = Inches(width)
            for row in table.rows:
                row.cells[index].width = Inches(width)
    reference._p.addprevious(table._tbl)


def update_research_file_manifest(document: Document) -> None:
    matches = [
        table
        for table in document.tables
        if len(table.rows) > 0
        and len(table.rows[0].cells) == 4
        and table.rows[0].cells[0].text.strip() == "File"
        and table.rows[0].cells[1].text.strip() == "Rows represented"
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one research-file manifest table; found {len(matches)}"
        )
    table = matches[0]
    additions = {
        SUMMARY_CSV.name: (
            "6",
            sha256(SUMMARY_CSV)[:16] + "…",
            "Human-model six-core agreement summaries for every available model",
        ),
        DIMENSIONS_CSV.name: (
            "48",
            sha256(DIMENSIONS_CSV)[:16] + "…",
            "Human-model full, evidence-presence, and conditional agreement by dimension",
        ),
        CONSENSUS_CSV.name: (
            "8",
            sha256(CONSENSUS_CSV)[:16] + "…",
            "Human confirmation of Mini-Claude-Gemini and four-model convergence cells",
        ),
    }
    existing = {
        row.cells[0].text.strip(): row
        for row in table.rows[1:]
        if row.cells[0].text.strip()
    }
    for filename, values in additions.items():
        row = existing.get(filename) or table.add_row()
        for index, value in enumerate((filename, *values)):
            set_cell(row.cells[index], value, size=6.0)
        prevent_row_split(row)
    repeat_table_header(table.rows[0])


def load_gemma_frame(store: HumanAnnotationStore) -> pd.DataFrame:
    sample_ids = set(store._sample()["paper_id"].astype(str))
    cache_dir = ROOT / "data/interim/spec_cache" / PROTOCOL_ID / "gemma4_31b"
    rows: list[dict] = []
    for paper_id in sorted(sample_ids):
        path = cache_dir / cache_key(paper_id)
        if not path.exists():
            continue
        try:
            payload = json.loads(path.read_text(encoding="utf-8"))
        except (OSError, json.JSONDecodeError):
            continue
        if str(payload.get("paper_id", "")).strip() != paper_id:
            continue
        if str(payload.get("coding_protocol", "")).strip() != PROTOCOL_ID:
            continue
        rows.append(payload)
    if not rows:
        return pd.DataFrame(columns=["paper_id", *DISPLAY_COLUMNS])
    frame = enrich_for_analysis(pd.DataFrame(rows).fillna(""))
    for column in DISPLAY_COLUMNS:
        if column not in frame.columns:
            frame[column] = ""
    return frame[["paper_id", *DISPLAY_COLUMNS]].drop_duplicates(
        "paper_id", keep="first"
    )


def metric_rows(
    human: pd.DataFrame,
    model: pd.DataFrame,
    model_id: str,
) -> list[dict[str, object]]:
    merged = human.merge(
        model,
        on="paper_id",
        how="inner",
        suffixes=("__human", "__model"),
        validate="one_to_one",
    ).sort_values("paper_id")
    rows: list[dict[str, object]] = []
    for column in DISPLAY_COLUMNS:
        human_values = (
            merged[f"{column}__human"].fillna("").astype(str).str.strip()
        )
        model_values = (
            merged[f"{column}__model"].fillna("").astype(str).str.strip()
        )
        full = pairwise_percent_agreement(
            human_values.tolist(), model_values.tolist()
        )
        full_alpha = krippendorff_alpha_nominal(
            [list(pair) for pair in zip(human_values, model_values)]
        )
        excluded = IRR_UNOBSERVED_VALUES.get(column, frozenset())
        human_observed = human_values.ne("") & ~human_values.isin(excluded)
        model_observed = model_values.ne("") & ~model_values.isin(excluded)
        human_presence = human_observed.map(
            {True: "observed", False: "unobserved"}
        )
        model_presence = model_observed.map(
            {True: "observed", False: "unobserved"}
        )
        presence = pairwise_percent_agreement(
            human_presence.tolist(), model_presence.tolist()
        )
        presence_alpha = krippendorff_alpha_nominal(
            [list(pair) for pair in zip(human_presence, model_presence)]
        )
        both = human_observed & model_observed
        conditional = pairwise_percent_agreement(
            human_values.loc[both].tolist(), model_values.loc[both].tolist()
        )
        conditional_alpha = krippendorff_alpha_nominal(
            [
                list(pair)
                for pair in zip(
                    human_values.loc[both], model_values.loc[both]
                )
            ]
        )
        rows.append(
            {
                "model_id": model_id,
                "model": MODEL_LABELS[model_id],
                "dimension": column,
                "dimension_label": DIMENSION_LABELS[column],
                "classification": DIMENSION_USE[column],
                "comparable_papers": full.comparable,
                "agreements": full.agreements,
                "exact_agreement": full.percent_agreement,
                "krippendorff_alpha": full_alpha,
                "observability_exact_agreement": presence.percent_agreement,
                "observability_krippendorff_alpha": presence_alpha,
                "jointly_observed_papers": conditional.comparable,
                "observed_category_exact_agreement": (
                    conditional.percent_agreement
                ),
                "observed_category_krippendorff_alpha": conditional_alpha,
            }
        )
    return rows


def consensus_rows(
    human: pd.DataFrame,
    frames: dict[str, pd.DataFrame],
) -> list[dict[str, object]]:
    preferred = (
        "gpt-5.4-mini-2026-03-17",
        "claude-sonnet-5",
        "gemini-3.1-pro-preview",
    )
    all_four = (
        "gpt-5.4-mini-2026-03-17",
        "gpt-4.1-nano-2025-04-14",
        "claude-sonnet-5",
        "gemini-3.1-pro-preview",
    )
    combined = human.rename(
        columns={column: f"human__{column}" for column in DISPLAY_COLUMNS}
    )
    for model_id in all_four:
        selected = frames[model_id].rename(
            columns={
                column: f"{model_id}__{column}" for column in DISPLAY_COLUMNS
            }
        )
        combined = combined.merge(
            selected, on="paper_id", how="inner", validate="one_to_one"
        )
    if len(combined) != 23:
        raise RuntimeError(
            f"Expected 23 papers on the complete-model intersection; found {len(combined)}"
        )

    rows = []
    for column in DISPLAY_COLUMNS:
        trio_values = [
            combined[f"{model_id}__{column}"].fillna("").astype(str).str.strip()
            for model_id in preferred
        ]
        four_values = [
            combined[f"{model_id}__{column}"].fillna("").astype(str).str.strip()
            for model_id in all_four
        ]
        human_values = (
            combined[f"human__{column}"].fillna("").astype(str).str.strip()
        )
        trio_unanimous = trio_values[0].eq(trio_values[1]) & trio_values[0].eq(
            trio_values[2]
        )
        four_unanimous = (
            four_values[0].eq(four_values[1])
            & four_values[0].eq(four_values[2])
            & four_values[0].eq(four_values[3])
        )
        trio_matches = trio_unanimous & trio_values[0].eq(human_values)
        four_matches = four_unanimous & four_values[0].eq(human_values)
        rows.append(
            {
                "dimension": column,
                "dimension_label": DIMENSION_LABELS[column],
                "classification": DIMENSION_USE[column],
                "preferred_trio_unanimous_cells": int(trio_unanimous.sum()),
                "preferred_trio_human_matches": int(trio_matches.sum()),
                "all_four_unanimous_cells": int(four_unanimous.sum()),
                "all_four_human_matches": int(four_matches.sum()),
            }
        )
    return rows


def build_outputs() -> tuple[pd.DataFrame, pd.DataFrame, pd.DataFrame, dict]:
    store = HumanAnnotationStore(ROOT)
    progress = store.progress()
    completed = [
        row
        for row in progress["annotators"]
        if row["completed_papers"] == progress["sample_papers"]
    ]
    if len(completed) != 1:
        raise RuntimeError(
            "Expected exactly one completed 23-paper human annotation record"
        )
    annotator_id = completed[0]["annotator_id"]
    human = store._human_frame(annotator_id)
    frames = store._model_frames()
    frames["gemma4:31b"] = load_gemma_frame(store)

    missing = set(MODEL_ORDER) - set(frames)
    if missing:
        raise RuntimeError(f"Required model frames are unavailable: {sorted(missing)}")

    dimension_records: list[dict[str, object]] = []
    for model_id in MODEL_ORDER:
        dimension_records.extend(metric_rows(human, frames[model_id], model_id))
    dimensions = pd.DataFrame(dimension_records)
    dimensions["model_order"] = dimensions["model_id"].map(
        {model: index for index, model in enumerate(MODEL_ORDER)}
    )
    dimensions["dimension_order"] = dimensions["dimension"].map(
        {column: index for index, column in enumerate(DISPLAY_COLUMNS)}
    )
    dimensions = dimensions.sort_values(
        ["model_order", "dimension_order"]
    ).drop(columns=["model_order", "dimension_order"])

    summary_rows = []
    for model_id in MODEL_ORDER:
        current = dimensions[
            (dimensions["model_id"] == model_id)
            & (dimensions["classification"] == "Core")
        ]
        summary_rows.append(
            {
                "model_id": model_id,
                "model": MODEL_LABELS[model_id],
                "human_common_papers": int(current["comparable_papers"].iloc[0]),
                "core_agreements": int(current["agreements"].sum()),
                "core_comparisons": int(current["comparable_papers"].sum()),
                "mean_exact_agreement_six_core": float(
                    current["exact_agreement"].mean()
                ),
                "mean_krippendorff_alpha_six_core": float(
                    current["krippendorff_alpha"].mean()
                ),
                "analytical_status": MODEL_STATUS[model_id],
            }
        )
    summary = pd.DataFrame(summary_rows)
    consensus = pd.DataFrame(consensus_rows(human, frames))

    # Guard against silently publishing a different analytical state.
    expected = {
        "gpt-5.4-mini-2026-03-17": (23, 91, 0.6594, 0.5178),
        "gpt-4.1-nano-2025-04-14": (23, 57, 0.4130, 0.1403),
        "claude-sonnet-5": (23, 98, 0.7101, 0.5928),
        "gemini-3.1-pro-preview": (23, 101, 0.7319, 0.6235),
        "llama3.2": (22, 45, 0.3409, -0.0486),
        "gemma4:31b": (6, 15, 0.4167, 0.1433),
    }
    for row in summary.itertuples():
        papers, agreements, exact, alpha = expected[row.model_id]
        if row.human_common_papers != papers or row.core_agreements != agreements:
            raise RuntimeError(f"Unexpected human-validation counts for {row.model}")
        if abs(row.mean_exact_agreement_six_core - exact) > 0.0001:
            raise RuntimeError(f"Unexpected mean exact agreement for {row.model}")
        if abs(row.mean_krippendorff_alpha_six_core - alpha) > 0.0001:
            raise RuntimeError(f"Unexpected mean alpha for {row.model}")

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    summary.to_csv(SUMMARY_CSV, index=False, encoding="utf-8-sig")
    dimensions.to_csv(DIMENSIONS_CSV, index=False, encoding="utf-8-sig")
    consensus.to_csv(CONSENSUS_CSV, index=False, encoding="utf-8-sig")
    manifest = {
        "analysis": "human-anchored model validation",
        "human_raters": 1,
        "human_completed_papers": 23,
        "human_completion_timestamp": completed[0]["updated_at"],
        "evidence_boundary": "title, abstract, and author keywords",
        "agreement_statistics": (
            "exact agreement and pairwise nominal Krippendorff's alpha"
        ),
        "complete_model_intersection": [
            MODEL_LABELS[model] for model in MODEL_ORDER[:4]
        ],
        "diagnostic_partial_models": [
            MODEL_LABELS[model] for model in MODEL_ORDER[4:]
        ],
        "files": {},
    }
    for path in (SUMMARY_CSV, DIMENSIONS_CSV, CONSENSUS_CSV):
        manifest["files"][path.name] = {
            "rows": int(len(pd.read_csv(path))),
            "sha256": sha256(path),
        }
    MANIFEST_JSON.write_text(
        json.dumps(manifest, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )
    return summary, dimensions, consensus, manifest


def update_document(
    summary: pd.DataFrame,
    dimensions: pd.DataFrame,
    consensus: pd.DataFrame,
) -> None:
    document = Document(SUPPLEMENT)
    a7 = find_paragraph(document, "A7. Research-File Manifest")

    existing_heading = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip() == "Human-anchored model validation"
    ]
    if existing_heading:
        if len(existing_heading) != 1:
            raise RuntimeError("Multiple human-validation sections found")
        remove_between(existing_heading[0]._p, a7._p)
    else:
        obsolete = find_paragraph(
            document, "Note. No blind human validation was available"
        )
        remove_between(obsolete._p, a7._p)

    heading = insert_paragraph_before(
        a7,
        "Human-anchored model validation",
        bold=True,
        color="0F4761",
        size=12.0,
        keep_with_next=True,
    )
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    heading.paragraph_format.space_before = Pt(9)
    heading.paragraph_format.space_after = Pt(4)

    insert_paragraph_before(
        a7,
        "Blind human annotation was completed for all 23 papers that occurred "
        "naturally in both the 2,235-paper probability sample and the "
        "systematic close-reading set. One researcher independently applied the "
        "same eight-dimensional instrument using only titles, abstracts, and "
        "author keywords, with model outputs withheld. The resulting estimates "
        "are human-anchored convergence checks on a small probability-sample "
        "overlap. They do not make the human classifications ground truth, "
        "estimate full-text validity, or support population-level accuracy claims.",
    )
    insert_paragraph_before(
        a7,
        "Table A6.6. Human-model agreement across the six core dimensions",
        bold=True,
        size=9.0,
        centered=True,
        keep_with_next=True,
    )
    summary_rows = [
        [
            row.model,
            f"{int(row.human_common_papers):,}",
            f"{int(row.core_agreements):,}/{int(row.core_comparisons):,}",
            percent_text(row.mean_exact_agreement_six_core),
            alpha_text(row.mean_krippendorff_alpha_six_core),
            row.analytical_status,
        ]
        for row in summary.itertuples()
    ]
    insert_table_before(
        document,
        a7,
        [
            "Model",
            "Human-common papers",
            "Exact core cells",
            "Mean exact",
            "Mean nominal α",
            "Analytical use",
        ],
        summary_rows,
        font_size=6.4,
        widths=[1.35, 0.82, 0.75, 0.65, 0.66, 1.70],
    )
    insert_paragraph_before(
        a7,
        "Note. Means are arithmetic orientation summaries across AI positioning, "
        "technical type, AI role, mechanism, level, and scope on each model's "
        "exact paper intersection with the human rater. Gemini, Claude, Mini, "
        "and Nano each cover all 23 papers. Llama and Gemma are reported only as "
        "partial diagnostics because their intersections contain 22 and 6 "
        "papers, respectively.",
    )
    insert_paragraph_before(
        a7,
        "Among the four complete models, Gemini showed the strongest average "
        "human convergence across the six core dimensions (73.19% exact; "
        "mean α = .62), followed by Claude (71.01%; α = .59), Mini (65.94%; "
        "α = .52), and Nano (41.30%; α = .14). Agreement varied materially by "
        "dimension, so these means are not treated as omnibus reliability "
        "coefficients and the dimension-level estimates remain the analytical record.",
    )

    insert_paragraph_before(
        a7,
        "Table A6.7. Complete human-model agreement by model and dimension",
        bold=True,
        size=9.0,
        centered=True,
        keep_with_next=True,
    )
    dimension_rows = [
        [
            row.model,
            row.dimension_label,
            row.classification,
            f"{int(row.comparable_papers):,}",
            percent_text(row.exact_agreement),
            alpha_text(row.krippendorff_alpha),
            percent_text(row.observability_exact_agreement),
            alpha_text(row.observability_krippendorff_alpha),
            f"{int(row.jointly_observed_papers):,}",
            percent_text(row.observed_category_exact_agreement),
            alpha_text(row.observed_category_krippendorff_alpha),
        ]
        for row in dimensions.itertuples()
    ]
    insert_table_before(
        document,
        a7,
        [
            "Model",
            "Dimension",
            "Use",
            "Papers",
            "Full exact",
            "Full α",
            "Evidence-presence exact",
            "Evidence-presence α",
            "Both observed",
            "Conditional exact",
            "Conditional α",
        ],
        dimension_rows,
        font_size=5.1,
        widths=[1.02, 1.07, 0.48, 0.50, 0.58, 0.48, 0.64, 0.55, 0.64, 0.70, 0.58],
    )
    insert_paragraph_before(
        a7,
        "Note. Full agreement treats each unobserved value as a substantive "
        "category in the agreement calculation. Presence agreement reduces each "
        "rating to evidence observed versus not observed. Conditional agreement "
        "uses only papers for which both the human and model found evidence and "
        "must therefore be interpreted with the displayed jointly observed "
        "count. Every α is pairwise nominal Krippendorff's α. The Llama and "
        "Gemma rows are diagnostic partial intersections and are not directly "
        "comparable with the 23-paper complete-model rows.",
    )
    insert_paragraph_before(
        a7,
        "The decomposition qualifies several raw percentages. Human-Gemini "
        "mechanism agreement was 73.91% with α = .56; where both identified a "
        "mechanism, exact category agreement was 83.33% with α = .81 (n = 6). "
        "Claude produced the same conditional mechanism result, while Mini's "
        "full mechanism agreement was 47.83% but rose to 75.00% with α = .72 "
        "among four jointly observed cases. Nano's technical-type result shows "
        "the opposite distinction: full agreement was 39.13%, yet category "
        "agreement reached 88.89% with α = .81 among the nine papers where both "
        "raters found a technical type, indicating that disagreement centered on "
        "whether the evidence justified a type. Definition clarity illustrates "
        "absence inflation: the human rater coded no definition in 20 of 23 "
        "papers, and Gemini's 86.96% exact agreement corresponded to α = -.03 "
        "with no jointly observed definitional category.",
    )

    a68_caption = insert_paragraph_before(
        a7,
        "Table A6.8. Human confirmation of exact model-convergence cells",
        bold=True,
        size=9.0,
        centered=True,
        keep_with_next=True,
    )
    a68_caption.paragraph_format.page_break_before = True
    consensus_rows_display = []
    for row in consensus.itertuples():
        trio_share = (
            row.preferred_trio_human_matches
            / row.preferred_trio_unanimous_cells
            if row.preferred_trio_unanimous_cells
            else None
        )
        four_share = (
            row.all_four_human_matches / row.all_four_unanimous_cells
            if row.all_four_unanimous_cells
            else None
        )
        consensus_rows_display.append(
            [
                row.dimension_label,
                row.classification,
                f"{row.preferred_trio_unanimous_cells}",
                (
                    f"{row.preferred_trio_human_matches}/"
                    f"{row.preferred_trio_unanimous_cells} "
                    f"({percent_text(trio_share)})"
                ),
                f"{row.all_four_unanimous_cells}",
                (
                    f"{row.all_four_human_matches}/"
                    f"{row.all_four_unanimous_cells} "
                    f"({percent_text(four_share)})"
                ),
            ]
        )
    core = consensus[consensus["classification"] == "Core"]
    all_dimensions = consensus
    for label, frame in (
        ("Six core dimensions", core),
        ("All eight dimensions", all_dimensions),
    ):
        trio_total = int(frame["preferred_trio_unanimous_cells"].sum())
        trio_match = int(frame["preferred_trio_human_matches"].sum())
        four_total = int(frame["all_four_unanimous_cells"].sum())
        four_match = int(frame["all_four_human_matches"].sum())
        consensus_rows_display.append(
            [
                label,
                "Summary",
                f"{trio_total}",
                f"{trio_match}/{trio_total} ({trio_match / trio_total:.2%})",
                f"{four_total}",
                f"{four_match}/{four_total} ({four_match / four_total:.2%})",
            ]
        )
    insert_table_before(
        document,
        a7,
        [
            "Dimension",
            "Use",
            "Mini-Claude-Gemini unanimous",
            "Human matches",
            "All four unanimous",
            "Human matches",
        ],
        consensus_rows_display,
        font_size=6.2,
        widths=[1.35, 0.58, 1.02, 1.12, 0.88, 1.12],
    )
    insert_paragraph_before(
        a7,
        "Note. On the six core dimensions, Mini, Claude, and Gemini were "
        "unanimous in 85 paper-dimension cells, of which 74 matched the human "
        "classification (87.06%). All four complete models were unanimous in 42 "
        "core cells, of which 37 matched the human classification (88.10%). "
        "This supports use of the Mini-Claude-Gemini convergence set as a "
        "high-convergence evidence filter. It does not replace Mini as the "
        "prespecified primary coding record and does not convert model consensus "
        "into a consensus-coded dataset.",
    )
    insert_paragraph_before(
        a7,
        "The human anchor therefore strengthens the interpretation of the "
        "complete-model reliability analysis while preserving its limits. One "
        "human rater and 23 papers are sufficient to identify dimension-specific "
        "convergence, absence inflation, and diagnostic model failure, but not to "
        "establish criterion validity for the 22,345-paper corpus. Human-human "
        "reliability for the eight-dimensional instrument remains unavailable. "
        "Qwen 3.5 27B and Gemini 2.5 Pro had no usable ratings on this human "
        "sample and are consequently absent from the tables.",
    )
    update_research_file_manifest(document)

    if not BACKUP.exists():
        shutil.copy2(SUPPLEMENT, BACKUP)
    with tempfile.NamedTemporaryFile(
        prefix=SUPPLEMENT.stem + ".",
        suffix=".docx",
        dir=SUPPLEMENT.parent,
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        os.replace(temporary, SUPPLEMENT)
    finally:
        temporary.unlink(missing_ok=True)


def main() -> None:
    if not SUPPLEMENT.exists():
        raise FileNotFoundError(SUPPLEMENT)
    summary, dimensions, consensus, _ = build_outputs()
    update_document(summary, dimensions, consensus)
    print(f"Updated {SUPPLEMENT}")
    print(f"Wrote {SUMMARY_CSV}")
    print(f"Wrote {DIMENSIONS_CSV}")
    print(f"Wrote {CONSENSUS_CSV}")
    print(f"Wrote {MANIFEST_JSON}")


if __name__ == "__main__":
    main()
