"""Insert the structured close-reading results and Appendix A9 exhibits.

This surgically updates the current author-edited manuscript and supplement. It
replaces only Results Section 4.5 and the material following Table A9.1 up to
Appendix A10. It does not rebuild either DOCX.
"""

from __future__ import annotations

import os
from pathlib import Path
import tempfile

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
SPECIFICATION_FIGURE = (
    ROOT
    / "reports/analysis/figures/specification/specification_close_reading.png"
)
COVERAGE_FIGURE = (
    ROOT
    / "reports/analysis/figures/contrasting/close_reading_topic_vos_coverage.png"
)


SECTION_45 = [
    (
        "The 136-paper structured close-reading set was used to determine whether "
        "the recurrent construct specifications identified in the aggregate "
        "analysis formed coherent theoretical explanations when the papers were "
        "read together, and to identify the conditions, boundaries, and "
        "counterexamples surrounding those explanations. The set contains 51 "
        "papers from Leading entrepreneurship journals, 73 from Additional "
        "entrepreneurship journals, and 12 cross-domain contrast cases. It "
        "supports theoretical interpretation rather than prevalence estimation."
    ),
    (
        "The specification profile confirms that the reading set "
        "captures substantively different uses of AI. AI positioning was "
        "observable in 134 papers: 85 treated AI as the phenomenon, 26 as a "
        "research method, and 23 as both. AI role was observable in 132 papers, "
        "led by tool (n = 50), research method (n = 27), context (n = 25), and "
        "firm capability (n = 20). In contrast, an observable mechanism was "
        "identified in 79 papers and a definitional signal in 48. Reading the "
        "papers together therefore reinforced the aggregate finding that an "
        "identifiable AI role does not necessarily provide the causal mechanism "
        "or definitional clarity required for theoretical accumulation."
    ),
    (
        "Bottleneck relocation emerged as the central entrepreneurship insight. "
        "The recurrent configurations associate AI with prediction, learning, "
        "uncertainty reduction, search, and judgment, but the close reading "
        "shows that these mechanisms do not simply remove entrepreneurial "
        "constraints. Chalmers et al. (2021) connect cheaper solution generation "
        "to a greater need for evaluation and selection. Ramoglou et al. (2026) "
        "describe opportunity search as adjudication among machine-generated "
        "possibilities, while Rady et al. (2026) show the difficulty of separating "
        "plausible opportunities from hallucinated outputs. De Véricourt and "
        "Gurkan (2026) similarly show that improved prediction does not resolve "
        "verification and reliance. Across these papers, expanded prediction, "
        "search, and generation relocate the binding constraint toward "
        "plausibility assessment, calibrated reliance, selective commitment, and "
        "responsibility for action."
    ),
    (
        "Organizational embedding emerged as the condition under which firms can "
        "manage this relocated bottleneck. Papers connecting AI capability with "
        "learning locate value in routines, knowledge bases, data, skills, "
        "absorptive capacity, and governance rather than in access to an isolated "
        "technical tool (De Fano et al., 2025; Shore et al., 2024; Abbas et al., "
        "2026; Grashof & Kopka, 2023). Studies of small and medium-sized "
        "enterprises identify the same condition through infrastructure, "
        "readiness, trust, skills, and governance limitations (Schwaeke et al., "
        "2025; Ledesma Chaves et al., 2026; Metzger et al., 2025). AI-related "
        "value therefore depends on how the technology is embedded in "
        "organizational arrangements for learning, evaluation, and accountability."
    ),
    (
        "Domain context bounds the mechanism through which bottleneck relocation "
        "occurs. Entrepreneurship papers describe AI as an innovation "
        "intermediary, an ecosystem force, and a collection of technical forms "
        "with different organizational consequences (Just, 2024; Hunt & Kurdoglu, "
        "2025; Chalmers et al., 2026). The cross-domain cases show that loan-text "
        "signal extraction, lending-bias detection, and national "
        "information-processing institutions make different mechanism claims "
        "rather than applying one universal mechanism in different settings "
        "(Netzer et al., 2019; Fu et al., 2021; Yoon et al., 2025). Domain "
        "variation therefore determines whether bottleneck relocation assumes a "
        "predictive, learning, interactional, or judgment-centered form."
    ),
    (
        "Agency remains a frontier insight. Six papers in the reading set assign "
        "AI an actor-like role, while other papers describe AI as augmenting "
        "entrepreneurial cognition, acting as a teammate, or "
        "participating as a relational nonhuman actor or co-agent (Shepherd & "
        "Majchrzak, 2022; Murtinu & De Massis, 2025; Al-Bashrawi et al., 2026; "
        "Spurrier et al., 2025). The available evidence identifies these positions "
        "but cannot establish how initiative, authority, decision rights, and "
        "responsibility are distributed between human and artificial actors. "
        "Agency is therefore retained as a boundary requiring full-text and "
        "claim-level investigation rather than as a settled conclusion."
    ),
    (
        "The structured close reading consequently produced one central insight, "
        "bottleneck relocation; one organizational condition, organizational "
        "embedding; one domain-mechanism boundary; and one unresolved agency "
        "frontier. These interpretations were supported across the current topic "
        "assignments and bounded by counterexamples rather than inferred from "
        "isolated papers. Independent human allocation of 14 papers to the "
        "interpretive families produced 14/14 agreement (Cohen’s κ = 1.00; "
        "Supplementary Appendix A9), supporting the consistency of the "
        "interpretation while not converting the reading set into a prevalence "
        "sample."
    ),
]


A9_RESULTS_ROWS = [
    (
        "Bottleneck relocation",
        "Tool × prediction: 157 papers (Leading 76; Additional 81); tool × "
        "judgment: 62 (32; 30); tool × uncertainty reduction: 66 (28; 38)",
        "Prediction, generation, and search expand information but relocate the "
        "constraint toward evaluation, plausibility, calibrated reliance, "
        "selective commitment, and responsibility (Chalmers et al., 2021; "
        "Ramoglou et al., 2026; Rady et al., 2026; de Véricourt & Gurkan, 2026)",
        "Central entrepreneurship insight; not a universal performance claim",
    ),
    (
        "Organizational embedding",
        "Firm capability × learning: 57 papers (Leading 6; Additional 51)",
        "AI-related value depends on routines, data, skills, absorptive capacity, "
        "learning, and governance rather than access to an isolated tool (De Fano "
        "et al., 2025; Shore et al., 2024; Abbas et al., 2026; Grashof & Kopka, "
        "2023)",
        "Condition for managing the relocated bottleneck",
    ),
    (
        "Domain-mechanism boundary",
        "Context × stakeholder interaction: 23 papers (Leading 4; Additional 19), "
        "supplemented by 12 cross-domain contrast cases",
        "The same AI label carries predictive, learning, interactional, and "
        "judgment-centered mechanisms in different domains (Just, 2024; Hunt & "
        "Kurdoglu, 2025; Netzer et al., 2019; Fu et al., 2021; Yoon et al., 2025)",
        "Boundary condition, not a separate competing insight",
    ),
    (
        "Agency frontier",
        "Actor/agent appears in 19 of 1,414 Combined-entrepreneurship papers with "
        "an observed role and in 6 of 132 close-reading papers with an observed role",
        "Actor-like, teammate, and co-agent descriptions raise unresolved questions "
        "about initiative, authority, decision rights, and responsibility "
        "(Shepherd & Majchrzak, 2022; Murtinu & De Massis, 2025; Al-Bashrawi et "
        "al., 2026; Spurrier et al., 2025)",
        "Open frontier requiring full-text, claim-level analysis",
    ),
]


def save_atomic(document: Document, path: Path) -> None:
    with tempfile.NamedTemporaryFile(
        prefix=path.stem + ".",
        suffix=".docx",
        dir=path.parent,
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def paragraph_starting(document: Document, prefixes: tuple[str, ...]):
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if any(paragraph.text.strip().startswith(prefix) for prefix in prefixes)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning one of {prefixes}; found {len(matches)}"
        )
    return matches[0]


def set_paragraph_text(paragraph, text: str) -> None:
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = text
    else:
        paragraph.add_run(text)


def insert_paragraph_before(reference, text: str = "", style: str = "Normal"):
    paragraph = reference.insert_paragraph_before(text)
    paragraph.style = style
    return paragraph


def remove_between(start_element, end_element) -> None:
    current = start_element.getnext()
    while current is not None and current is not end_element:
        following = current.getnext()
        current.getparent().remove(current)
        current = following


def update_manuscript() -> None:
    document = Document(MANUSCRIPT)
    heading = paragraph_starting(
        document,
        (
            "4.5 Entrepreneurship interpretation",
            "4.5 Structured close-reading results",
        ),
    )
    discussion = paragraph_starting(document, ("Discussion",))
    if heading._p.getparent() is not discussion._p.getparent():
        raise RuntimeError("Section 4.5 and Discussion are not document siblings")
    remove_between(heading._p, discussion._p)
    set_paragraph_text(heading, "4.5 Structured close-reading results")
    heading.style = "Heading 2"
    for text in SECTION_45:
        insert_paragraph_before(discussion, text)
    save_atomic(document, MANUSCRIPT)


def style_table(table, column_widths: tuple[float, ...] | None = None) -> None:
    table.style = "Table Grid"
    table.autofit = False
    if column_widths:
        if len(column_widths) != len(table.columns):
            raise ValueError("One width is required for every table column")
        for column, width in zip(table.columns, column_widths):
            column.width = Inches(width)
    for row_number, row in enumerate(table.rows):
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for column_number, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if column_widths:
                width = column_widths[column_number]
                cell.width = Inches(width)
                tc_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if tc_width is not None:
                    tc_width.set(qn("w:w"), str(int(width * 1440)))
                    tc_width.set(qn("w:type"), "dxa")
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)
                    if row_number == 0:
                        run.bold = True


def add_caption(reference, text: str, *, page_break_before: bool = False):
    paragraph = insert_paragraph_before(reference, text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break_before
    return paragraph


def add_figure(reference, path: Path, width: float) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = insert_paragraph_before(reference)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def update_supplement() -> None:
    document = Document(SUPPLEMENT)
    a10 = paragraph_starting(document, ("A10.",))
    target_tables = [
        table
        for table in document.tables
        if table.cell(0, 0).text.strip() == "Step"
        and any(
            row.cells[0].text.strip() == "Current reading base"
            for row in table.rows
        )
    ]
    if len(target_tables) != 1:
        raise RuntimeError(
            f"Expected one Appendix A9.1 procedure table; found {len(target_tables)}"
        )
    table_a91 = target_tables[0]
    if table_a91._tbl.getparent() is not a10._p.getparent():
        raise RuntimeError("Appendix A9.1 and A10 are not document siblings")
    remove_between(table_a91._tbl, a10._p)

    insert_paragraph_before(
        a10,
        "The structured close-reading results are reported through two linked "
        "exhibits. Figure A9.1 describes the coded construct profile of the "
        "136-paper set, while Table A9.2 traces recurrent analytical patterns to "
        "their theory-elaboration interpretations and boundaries. Figure A9.2 "
        "reports the topic and VOS network-position checks used to assess the "
        "coverage of the 124 entrepreneurship papers.",
    )
    add_caption(
        a10,
        "Figure A9.1. Construct-specification profile of the 136-paper structured "
        "close-reading set",
        page_break_before=True,
    )
    add_figure(a10, SPECIFICATION_FIGURE, 6.4)
    insert_paragraph_before(
        a10,
        "Note. Each panel reports the distribution within papers having an "
        "observed value for that dimension. Observed denominators are 134 for AI "
        "positioning, 132 for AI role, 77 for technical type, 79 for mechanism, "
        "134 for level, 113 for process stage, 105 for scope, and 48 for "
        "definition clarity. The percentages characterize the structured "
        "close-reading set and do not estimate prevalence in Combined "
        "entrepreneurship.",
    )

    add_caption(
        a10,
        "Table A9.2. Analytical patterns and structured close-reading results",
        page_break_before=True,
    )
    results_table = document.add_table(rows=1, cols=4)
    headers = (
        "Structured result",
        "Recurring analytical evidence",
        "Close-reading interpretation",
        "Status in theory elaboration",
    )
    for cell, text in zip(results_table.rows[0].cells, headers):
        cell.text = text
    for values in A9_RESULTS_ROWS:
        row = results_table.add_row()
        for cell, text in zip(row.cells, values):
            cell.text = text
    style_table(results_table, (1.15, 1.65, 2.35, 1.35))
    a10._p.addprevious(results_table._tbl)
    insert_paragraph_before(
        a10,
        "Note. Counts derive from the primary coding record. Structured close "
        "reading establishes theoretical meaning, conditions, boundaries, and "
        "counterexamples; it does not estimate prevalence. Independent allocation "
        "of 14 papers to the interpretive families produced 14/14 agreement "
        "(Cohen’s κ = 1.00).",
    )

    add_caption(
        a10,
        "Figure A9.2. Topic and bibliometric-network coverage of the "
        "entrepreneurship close-reading papers",
        page_break_before=True,
    )
    add_figure(a10, COVERAGE_FIGURE, 6.4)
    insert_paragraph_before(
        a10,
        "Note. All 51 Leading and 73 Additional entrepreneurship papers appear "
        "in their respective VOSviewer document maps. Twenty-five Leading and 37 "
        "Additional papers fall in the highest total-link-strength quartile. "
        "Topic and network positions demonstrate coverage across the represented "
        "research conversations; they were not treated as statistical sampling "
        "weights or as the historical selection rule.",
    )
    save_atomic(document, SUPPLEMENT)


def main() -> None:
    update_manuscript()
    update_supplement()
    print(f"Updated {MANUSCRIPT.relative_to(ROOT)}")
    print(f"Updated {SUPPLEMENT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
