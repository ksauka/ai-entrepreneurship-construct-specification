"""Expand the platform-methodology section in the researcher-edited manuscript."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
BACKUP = ROOT / "docs/ETP draft - July2026ks.before-platform-expansion.docx"


PLATFORM_PARAGRAPHS = [
    (
        "The interactive platform was developed as an operational component of the methodology rather than "
        "as a presentation layer added after the analysis. It reads the same versioned corpus, construct-"
        "specification, domain, topic, and graph tables used to produce the manuscript results. Within each "
        "workspace, the current dataset scope and analytical selections are displayed explicitly. Changing a "
        "population, coding model, conditioning value, denominator view, matrix axis, or support threshold "
        "causes the corresponding distributions, matrices, counts, and evidence lists to be recalculated from "
        "the shared source tables. The platform therefore represents a set of reproducible analytical states "
        "rather than a collection of static figures. Reports, tables, figures, and release files can be generated "
        "or downloaded for the active state, which reduces the risk that a displayed result, its denominator, "
        "and its supporting papers become separated."
    ),
    (
        "The Analytics Dashboard provides corpus context. A dataset-scope selection is applied to publication "
        "and citation indicators, annual and cumulative publication output, productive journals and authors, "
        "highly cited papers, and keyword evolution. Users can change the publication window and plotted metric, "
        "compare author and Scopus index keywords, inspect leading, growing, or declining terms, and add any "
        "observed keyword to the graph. Selecting a publication year or keyword opens an evidence panel rather "
        "than only displaying an aggregate value. The panel identifies the contributing papers and provides "
        "bibliographic metadata and links to the corresponding Scopus records. This workspace was used to check "
        "population sizes, publication trends, keyword periods, and the papers underlying descriptive claims; "
        "the growth results themselves were treated as corpus context rather than evidence of theoretical "
        "accumulation."
    ),
    (
        "The Construct Specification workspace implements the marginal and conditioned analyses. Researchers "
        "can select the dataset scope, coding model, conditioning dimension and exact value, and either the full, "
        "observed, or comparative denominator view. The eight dimension distributions and paper counts update "
        "together, while a nested matrix permits any two different dimensions to be crossed using paper counts, "
        "row percentages, column percentages, or shares of the analyzed matrix. Clicking a bar or matrix cell "
        "opens the papers supporting that exact selection. For each paper, the evidence panel reports the title, "
        "abstract, author and index keywords, source, year, dataset membership, topic, assigned construct codes, "
        "supporting evidence, evidence status, confidence, mechanism logic, named theories, full-text-review "
        "flags, citation information, and source-record link. This makes it possible to move from a percentage "
        "to the admissible paper-level evidence without searching the repository."
    ),
    (
        "The same workspace makes model sensitivity and convergence inspectable. Users can compare aggregate "
        "leaders, cells conditioned on AI positioning, Leading-Additional directions, role-level locations, and "
        "recurring relations under alternative complete coding models. The inter-rater view reports exact "
        "agreement and nominal Krippendorff alpha on a balanced common-paper cohort and separates full-category, "
        "observability, and conditional-category agreement. Evidence lists can then be restricted to all "
        "supporting papers, papers receiving the same selected code from at least two models, higher agreement "
        "thresholds, or the preferred three-model convergence set comprising Mini, Claude, and Gemini. In the "
        "paper-level panel, the platform names the models that agree and displays every available model assignment. "
        "The three-model set was used as a high-convergence evidence or 'sweet-spot' filter when inspecting "
        "important cells; it was not used to replace the primary coding record with a majority classification, "
        "and convergence was not interpreted as proof that the shared assignment was correct."
    ),
    (
        "The Construct Contrasting workspace implements the remaining theory-elaboration tactics through three "
        "tabs. Horizontal contrasting applies the selected dimension to the registered business domains and "
        "entrepreneurship journal populations, with the full corpus or FT50 restriction providing the comparison "
        "base. Vertical contrasting allows level of analysis to be crossed with AI positioning, role, technical "
        "type, mechanism, process stage, or scope, with either matrix axis and cell metric controlled by the user. "
        "Structuring exposes recurring role-mechanism, role-level, mechanism-level, role-scope, and related "
        "configurations under a selectable minimum-support threshold. A selected cell opens an interpretation "
        "panel stating what is being contrasted, the relevant numerator and denominator, the baseline difference, "
        "and the supporting papers. The same agreement filters and paper-level model comparisons available in "
        "Construct Specification remain available here. These tabs were used to explore alternative matrix views, "
        "check that retained contrasts were not denominator artifacts, locate supporting and contrasting cases, "
        "and download the analysis states reported in Sections 4.2-4.4."
    ),
    (
        "The Topic Review workspace supports researcher interpretation of the separate, data-specific topic "
        "models. For each topic-model scope, it displays topic prevalence and allows a reviewer to inspect the "
        "automatic label, defining terms, centroid-nearest papers, and fitted papers before entering a humanized "
        "label, review status, reviewer name, and notes. Topic identity remains anchored to its scope and topic "
        "number, so changing the displayed label does not alter the fitted paper assignments. Approved labels can "
        "be applied to rebuilt prevalence figures, graph files, and downloadable topic tables only after the review "
        "requirements are met. This separation prevents provisional machine-generated labels from entering the "
        "manuscript or graph as if they were validated interpretations, while allowing researchers to perform the "
        "humanization and obtain the revised artifacts without editing repository files."
    ),
    (
        "Human participation is also implemented through the Human Annotation workspace, reached from Construct "
        "Specification. Annotators choose their own unique identifier, receive the same blinded paper order, and "
        "code the 23-paper probability-sample overlap independently from titles, abstracts, and author keywords, "
        "without seeing model outputs. The page provides the complete model coding request, permitted categories, "
        "evidence rules, and structured response schema so that reviewers can inspect the deductively constructed "
        "instrument rather than rely only on its description. Each annotator's decisions are saved separately and "
        "can be resumed or downloaded. Reliability is recalculated on the largest exact intersection of papers "
        "completed by the selected human and model raters, so unfinished annotations reduce the common-paper base "
        "rather than being treated as disagreements. This workspace permits additional human checks without "
        "requiring a person to code the complete 22,345-paper corpus."
    ),
    (
        "The Knowledge Graph workspace provides a relational route through the same evidence. It retrieves a "
        "bounded, dataset-specific seed from Neo4j rather than loading the complete graph into the browser. Users "
        "can search publications, authors, journals, topics, keywords, and construct codes; select node and "
        "relationship types; restrict seed publications to one observed specification value; focus on a node's "
        "direct neighbors; or expand the graph one relationship step at a time. Selecting a publication displays "
        "its bibliographic record and connections, while selecting a topic or construct node provides the connected "
        "papers. Read-only advanced graph queries are available where the graph-reader permission is enabled. The "
        "graph was used to organize navigation among papers, topics, journals, authors, and specification codes, "
        "not to generate an additional statistical sample or replace the tabular analyses."
    ),
    (
        "Finally, the Assistant tab provides guided retrieval for recurring interpretive questions, including "
        "which papers assign a selected role, technical type, or mechanism to AI; which papers leave the mechanism "
        "unspecified; and where the same technical type receives different roles. It returns records from the "
        "selected dataset scope rather than generating unsupported answers. Across all workspaces, clicking "
        "evidence therefore retains the chain from an aggregate pattern to its denominator, model assignment, "
        "supporting text, bibliographic record, and source link. The platform was used to explore the dynamic "
        "specification and contrasting results, organize close reading and topic interpretation, inspect model "
        "robustness and convergence, and preserve a reproducible audit trail. Detailed interface operations, "
        "research-file manifests, checksums, and reproduction procedures are reported in Supplementary Tables "
        "A2.1 and A7.1."
    ),
]


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def insert_paragraph_after(
    reference: Paragraph,
    text: str,
    *,
    style,
    font_name: str,
) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addnext(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = style
    run = paragraph.add_run(text)
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    return paragraph


def remove_between(start: Paragraph, stop: Paragraph) -> None:
    element = start._p.getnext()
    while element is not stop._p:
        if element is None:
            raise RuntimeError("Reached document end before the platform section stop")
        next_element = element.getnext()
        element.getparent().remove(element)
        element = next_element


def validate(document: Document) -> None:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = [
        "The interactive platform was developed as an operational component",
        "The Analytics Dashboard provides corpus context",
        "The Construct Specification workspace implements",
        "preferred three-model convergence set comprising Mini, Claude, and Gemini",
        "The Construct Contrasting workspace implements",
        "The Topic Review workspace supports researcher interpretation",
        "Human participation is also implemented through the Human Annotation workspace",
        "The Knowledge Graph workspace provides a relational route",
        "Finally, the Assistant tab provides guided retrieval",
        "Results",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Missing platform-methodology content: {missing}")
    old = "The interactive platform operationalizes the same analytical states used in the paper."
    if old in text:
        raise RuntimeError("The short platform paragraph remains in the manuscript")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(MANUSCRIPT, BACKUP)

    document = Document(MANUSCRIPT)
    if any(
        paragraph.text.startswith(PLATFORM_PARAGRAPHS[0][:60])
        for paragraph in document.paragraphs
    ):
        validate(document)
        print(f"Already expanded: {MANUSCRIPT}")
        return

    heading = find_paragraph(document, "3.7 Platform implementation")
    results = find_paragraph(document, "Results")
    existing = heading._p.getnext()
    if existing is results._p or not existing.tag.endswith("}p"):
        raise RuntimeError("Could not identify the existing Section 3.7 body paragraph")
    existing_paragraph = Paragraph(existing, heading._parent)
    style = existing_paragraph.style
    font_name = next(
        (run.font.name for run in existing_paragraph.runs if run.font.name),
        "Times New Roman",
    )

    remove_between(heading, results)
    last = heading
    for text in PLATFORM_PARAGRAPHS:
        last = insert_paragraph_after(
            last,
            text,
            style=style,
            font_name=font_name,
        )
    validate(document)

    temporary = MANUSCRIPT.with_suffix(".platform-expansion.tmp.docx")
    document.save(temporary)
    reopened = Document(temporary)
    validate(reopened)
    temporary.replace(MANUSCRIPT)
    print(f"Updated {MANUSCRIPT}")
    print(f"Backup  {BACKUP}")


if __name__ == "__main__":
    main()
