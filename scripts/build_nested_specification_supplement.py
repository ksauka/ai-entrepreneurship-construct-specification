"""Build the current supplementary appendix for nested specification, IRR, and topics.

The July supplementary methods document is used only as a style reference. Its
obsolete 367-paper content is removed. The rebuilt appendix reports the frozen
22,345-paper analytical contract, detailed study-status-conditioned results,
population boundary checks, and the inventory of every pairwise dimension
matrix exposed by the platform. It also reports the final four-model coverage,
balanced reliability, exact consensus diagnostics, and the complete topic-model
selection decision.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_full_theory_elaboration_manuscript import (
    british_spelling,
    format_table,
    set_cell,
    shade,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = ROOT / "docs/Supplementary_Methods  July 2026 version.docx"
OUTPUT = (
    ROOT
    / "docs/ETP supplementary methods and nested specification results - current evidence 2026-07-22.docx"
)
MARKDOWN = ROOT / "reports/analysis/ETP_SUPPLEMENTARY_NESTED_ANALYSIS_CURRENT_EVIDENCE.md"
TABLES = ROOT / "reports/analysis/tables/contrasting"
FIGURES = ROOT / "reports/analysis/figures/contrasting"

STATUS_SOURCE = TABLES / "study_status_conditioned_specification.csv"
CORE_ADDITIONAL = TABLES / "nested_status_core_additional_contrasts.csv"
PAIR_INVENTORY = TABLES / "nested_dimension_pair_inventory.csv"
MANIFEST = TABLES / "contrasting_manifest.json"
FIG_ALL = FIGURES / "nested_status_all_dimensions.png"
FIG_OBSERVABILITY = FIGURES / "nested_status_observability.png"
MODEL_TABLES = ROOT / "reports/analysis/tables/model_validation"
IRR_COVERAGE = MODEL_TABLES / "full_corpus_model_coverage.csv"
IRR_SUMMARY = MODEL_TABLES / "full_corpus_pairwise_irr_core_summary.csv"
IRR_DIMENSIONS = MODEL_TABLES / "full_corpus_pairwise_irr_dimensions.csv"
IRR_CONSENSUS = MODEL_TABLES / "full_corpus_dimension_consensus.csv"
IRR_MANIFEST = MODEL_TABLES / "full_corpus_model_irr_manifest.json"
TOPIC_OPTIMISATION = ROOT / "data/processed/topics/optimization"
TOPIC_SELECTION = TOPIC_OPTIMISATION / "topic_selection_review.json"
TOPIC_FINAL_MANIFEST = ROOT / "data/processed/topics/final_run_manifest.json"
TOPIC_JOINT_GRID = TOPIC_OPTIMISATION / "full_corpus_query_1_joint_grid.csv"
TOPIC_GRID_FIGURE = TOPIC_OPTIMISATION / "grid_search_review_overview.png"

STATUS_ORDER = ("phenomenon", "method", "both")
STATUS_LABELS = {"phenomenon": "Phenomenon", "method": "Method", "both": "Both"}
DIMENSION_ORDER = (
    ("ai_role", "AI role"),
    ("technical_type", "Technical type"),
    ("mechanism", "Mechanism"),
    ("level", "Level"),
    ("process_stage", "Process stage"),
    ("scope", "Scope"),
    ("definition", "Definition"),
)


def format_alpha(value: float) -> str:
    """Format a bounded coefficient without a leading zero (APA style)."""

    text = f"{value:.2f}"
    if text.startswith("-0"):
        return f"-{text[2:]}"
    if text.startswith("0"):
        return text[1:]
    return text


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if "Caption" in {style.name for style in document.styles}:
        paragraph.style = "Caption"


def add_table(
    document: Document,
    caption: str,
    headers: list[str],
    rows: list[list[object]],
    *,
    font_size: float = 8.0,
) -> None:
    add_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    if "Table Grid" in {style.name for style in document.styles if style.type == 3}:
        table.style = "Table Grid"
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True)
        shade(table.rows[0].cells[index], "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run(str(value))
            run.font.size = Pt(font_size)
    format_table(table, prevent_splitting=True)
    document.add_paragraph()


def add_picture(document: Document, path: Path, caption: str, width: float) -> None:
    if not path.exists():
        raise FileNotFoundError(path)
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run().add_picture(str(path), width=Inches(width))
    add_caption(document, caption)


def add_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run("Note. ")
    run.bold = True
    paragraph.add_run(text)


def status_table() -> pd.DataFrame:
    table = pd.read_csv(STATUS_SOURCE, keep_default_na=False)
    for column in ("share", "papers", "denominator", "control_papers"):
        table[column] = pd.to_numeric(table[column], errors="coerce").fillna(0)
    return table


def combined_observed(table: pd.DataFrame) -> pd.DataFrame:
    return table.loc[
        table["population"].eq("combined")
        & table["distribution"].eq("observed")
        & table["control_raw_value"].isin(STATUS_ORDER)
    ].copy()


def observability_rows(table: pd.DataFrame) -> list[list[object]]:
    selected = combined_observed(table)
    rows = []
    for dimension, label in DIMENSION_ORDER:
        cells = []
        for status in STATUS_ORDER:
            frame = selected.loc[
                selected["outcome_dimension"].eq(dimension)
                & selected["control_raw_value"].eq(status)
            ]
            if frame.empty:
                cells.append("Not available")
                continue
            denominator = int(frame["denominator"].iloc[0])
            control_n = int(frame["control_papers"].iloc[0])
            cells.append(f"{denominator:,}/{control_n:,} ({denominator/control_n:.1%})")
        rows.append([label, *cells])
    return rows


def leading_rows(table: pd.DataFrame) -> list[list[object]]:
    selected = combined_observed(table)
    rows = []
    for dimension, label in DIMENSION_ORDER:
        for status in STATUS_ORDER:
            frame = selected.loc[
                selected["outcome_dimension"].eq(dimension)
                & selected["control_raw_value"].eq(status)
            ].sort_values(["share", "papers"], ascending=False)
            if frame.empty:
                continue
            denominator = int(frame["denominator"].iloc[0])
            top = "; ".join(
                f"{row.category} {row.share:.1%} (n={int(row.papers):,})"
                for row in frame.head(3).itertuples()
            )
            rows.append([label, STATUS_LABELS[status], f"{denominator:,}", top])
    return rows


def detailed_dimension_rows(
    table: pd.DataFrame,
    dimension: str,
) -> list[list[object]]:
    selected = combined_observed(table)
    selected = selected[selected["outcome_dimension"].eq(dimension)].copy()
    order = (
        selected.groupby("category")["papers"].sum().sort_values(ascending=False).index.tolist()
    )
    rows = []
    for category in order:
        cells = []
        for status in STATUS_ORDER:
            match = selected.loc[
                selected["category"].eq(category)
                & selected["control_raw_value"].eq(status)
            ]
            if match.empty:
                cells.append("0 (0.0%)")
            else:
                item = match.iloc[0]
                cells.append(f"{int(item.papers):,} ({item.share:.1%})")
        rows.append([category, *cells])
    return rows


def strongest_population_contrasts() -> list[list[object]]:
    table = pd.read_csv(CORE_ADDITIONAL, keep_default_na=False)
    table["absolute_difference"] = table["core_minus_additional_pp"].abs()
    rows = []
    for status in STATUS_ORDER:
        for dimension, label in DIMENSION_ORDER:
            selected = table.loc[
                table["control_raw_value"].eq(status)
                & table["outcome_dimension"].eq(dimension)
            ].sort_values("absolute_difference", ascending=False)
            if selected.empty:
                continue
            item = selected.iloc[0]
            rows.append(
                [
                    STATUS_LABELS[status],
                    label,
                    item.category,
                    f"{item.core_share:.1%} (n={int(item.core_papers):,}/{int(item.core_denominator):,})",
                    f"{item.additional_share:.1%} (n={int(item.additional_papers):,}/{int(item.additional_denominator):,})",
                    f"{item.core_minus_additional_pp:+.1f} pp",
                ]
            )
    return rows


def pair_inventory_rows() -> list[list[object]]:
    table = pd.read_csv(PAIR_INVENTORY, keep_default_na=False)
    table = table.loc[
        table["population"].eq("combined") & table["distribution"].eq("observed")
    ].copy()
    table = table.sort_values(["row_label", "column_label"])
    return [
        [
            f"{row.row_label} by {row.column_label}",
            f"{int(row.analyzed_n):,}",
            int(row.nonzero_cells),
            f"{row.strongest_row_value} / {row.strongest_column_value}",
            f"{int(row.strongest_cell_papers):,}",
            f"{row.strongest_cell_share:.1%}",
        ]
        for row in table.itertuples()
    ]


def artifact_rows() -> list[list[object]]:
    manifest = json.loads(MANIFEST.read_text(encoding="utf-8"))
    names = [
        "study_status_conditioned_specification.csv",
        "nested_dimension_distributions.csv",
        "nested_dimension_pair_matrices.csv",
        "nested_dimension_pair_inventory.csv",
        "nested_specification_release.zip",
    ]
    purposes = {
        "study_status_conditioned_specification.csv": "All status-conditioned distributions used in Section A3",
        "nested_dimension_distributions.csv": "Every exact dimension-value control crossed with every remaining outcome dimension",
        "nested_dimension_pair_matrices.csv": "Cell-level counts and row, column, and analytical shares for all dimension pairs",
        "nested_dimension_pair_inventory.csv": "One-row audit summary for each population, view, and dimension pair",
        "nested_specification_release.zip": "Downloadable package containing the complete tidy release and readme",
    }
    rows = []
    for name in names:
        item = manifest["outputs"].get(name, {})
        checksum = str(item.get("sha256", ""))
        displayed_checksum = f"{checksum[:16]}..." if checksum else ""
        rows.append([name, f"{int(item.get('rows', 0)):,}", displayed_checksum, purposes[name]])
    return rows


def irr_coverage_rows() -> list[list[object]]:
    table = pd.read_csv(IRR_COVERAGE)
    return [
        [
            row.model_label,
            row.study_role,
            f"{int(row.successful_corpus_papers):,}",
            f"{row.coverage_share:.2%}",
            f"{int(row.missing_corpus_papers):,}",
            f"{int(row.balanced_common_papers):,}",
        ]
        for row in table.itertuples()
    ]


def irr_summary_rows() -> list[list[object]]:
    table = pd.read_csv(IRR_SUMMARY).sort_values(
        ["mean_krippendorff_alpha", "mean_exact_agreement"], ascending=False
    )
    return [
        [
            row.model_pair,
            f"{int(row.balanced_common_papers):,}",
            f"{row.mean_exact_agreement:.2%}",
            format_alpha(row.mean_krippendorff_alpha),
        ]
        for row in table.itertuples()
    ]


def irr_dimension_rows() -> list[list[object]]:
    table = pd.read_csv(IRR_DIMENSIONS)
    table["pair_order"] = table["model_pair"].astype("category").cat.codes
    table["dimension_order"] = table["dimension"].map(
        {
            "ai_method_or_phenomenon": 0,
            "ai_type_form": 1,
            "ai_role_function": 2,
            "ai_mechanism_analysis": 3,
            "level_of_analysis": 4,
            "scope_conditions": 5,
            "entrepreneurial_process_stage": 6,
            "definition_construct_clarity": 7,
        }
    )
    table = table.sort_values(["pair_order", "dimension_order"])
    return [
        [
            row.model_pair,
            row.dimension_label,
            row.classification,
            f"{int(row.comparable_papers):,}",
            f"{row.exact_agreement:.2%}",
            format_alpha(row.krippendorff_alpha),
            f"{row.observability_exact_agreement:.2%}",
            format_alpha(row.observability_krippendorff_alpha),
            f"{int(row.jointly_observed_papers):,}",
            f"{row.observed_category_exact_agreement:.2%}",
            format_alpha(row.observed_category_krippendorff_alpha),
        ]
        for row in table.itertuples()
    ]


def irr_consensus_rows() -> list[list[object]]:
    table = pd.read_csv(IRR_CONSENSUS)
    return [
        [
            row.dimension_label,
            row.classification,
            (
                f"{int(row.preferred_trio_agreement_papers):,} "
                f"({row.preferred_trio_agreement_share:.2%}); "
                f"unobserved {int(row.preferred_trio_unobserved_agreement_papers):,}; "
                f"observed {int(row.preferred_trio_observed_agreement_papers):,}"
            ),
            (
                f"{int(row.all_four_agreement_papers):,} "
                f"({row.all_four_agreement_share:.2%}); "
                f"unobserved {int(row.all_four_unobserved_agreement_papers):,}; "
                f"observed {int(row.all_four_observed_agreement_papers):,}"
            ),
        ]
        for row in table.itertuples()
    ]


def irr_artifact_rows() -> list[list[object]]:
    manifest = json.loads(IRR_MANIFEST.read_text(encoding="utf-8"))
    purposes = {
        IRR_COVERAGE.name: "Model-specific successful coverage and balanced-cohort construction",
        IRR_SUMMARY.name: "Six-core arithmetic means used only for orientation",
        IRR_DIMENSIONS.name: "All 48 pair-by-dimension exact-agreement and alpha estimates",
        IRR_CONSENSUS.name: "Preferred-trio and unanimous four-model exact consensus by dimension",
    }
    rows = []
    for name, item in manifest["outputs"].items():
        rows.append(
            [
                name,
                f"{int(item['rows']):,}",
                f"{item['sha256'][:16]}...",
                purposes[name],
            ]
        )
    return rows


def topic_selection_rows() -> list[list[object]]:
    """Return the tested grids and approved scope-specific topic decisions."""

    review = json.loads(TOPIC_SELECTION.read_text(encoding="utf-8"))
    final = json.loads(TOPIC_FINAL_MANIFEST.read_text(encoding="utf-8"))
    labels = {
        "full_corpus": "Full corpus",
        "query_1": "Broad business and management",
        "query_2": "FT50",
        "query_3": "Leading entrepreneurship journals",
        "query_4": "Additional entrepreneurship",
    }
    rows = []
    for scope_id, label in labels.items():
        grid = pd.read_csv(TOPIC_OPTIMISATION / scope_id / "grid_search_results.csv")
        candidates = ", ".join(str(int(value)) for value in grid["min_topic_size"])
        scope = review["scopes"][scope_id]
        model = final["models"][scope_id]
        papers = model.get("eligible_papers", model.get("papers"))
        rows.append(
            [
                label,
                f"{int(papers):,}",
                candidates,
                int(scope["selected_min_topic_size"]),
                int(scope["selected_topic_count"]),
                scope["decision"].replace("_", " "),
            ]
        )
    return rows


def topic_joint_grid_rows() -> list[list[object]]:
    """Return the common Full Corpus and Query 1 candidate comparison."""

    table = pd.read_csv(TOPIC_JOINT_GRID, encoding="utf-8-sig")
    return [
        [
            int(row.min_topic_size),
            int(row.n_topics_full_corpus),
            f"{row.composite_score_full_corpus:.4f}",
            int(row.n_topics_query_1),
            f"{row.composite_score_query_1:.4f}",
            f"{row.mean_composite_score:.4f}",
            "Selected" if bool(row.selected_by_joint_rule) else "",
        ]
        for row in table.itertuples()
    ]


def update_footers(document: Document) -> None:
    for section in document.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = "Supplementary Methods and Analysis Appendix | July 2026"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_landscape_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)


def add_portrait_section(document: Document) -> None:
    section = document.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)


def build() -> None:
    for path in (
        TEMPLATE,
        STATUS_SOURCE,
        CORE_ADDITIONAL,
        PAIR_INVENTORY,
        MANIFEST,
        IRR_COVERAGE,
        IRR_SUMMARY,
        IRR_DIMENSIONS,
        IRR_CONSENSUS,
        IRR_MANIFEST,
        TOPIC_SELECTION,
        TOPIC_FINAL_MANIFEST,
        TOPIC_JOINT_GRID,
        TOPIC_GRID_FIGURE,
    ):
        if not path.exists():
            raise FileNotFoundError(path)
    table = status_table()
    document = Document(TEMPLATE)
    clear_body(document)
    section = document.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Supplementary Methods and Analysis Appendix")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Supplementary to: What's theory got to do with it? Theory Elaboration on AI in Entrepreneurship Scholarship"
    ).italic = True
    document.add_paragraph(
        "This appendix preserves the detailed analysis behind the main paper without crowding the core theory "
        "narrative. It is generated from the same frozen tables and calculation functions as the interactive "
        "platform. It does not recode papers or introduce an additional sample."
    )

    document.add_heading("A1. Analytical Scope and Denominator Contract", level=1)
    document.add_paragraph(
        "The primary dataset contains 22,345 unique Scopus records coded by the current complete primary model. "
        "Nested analysis means analytical conditioning: one exact coded value defines the subset, after which "
        "all remaining distributions and relationship matrices are recalculated. It does not mean retrieving "
        "new papers or treating subgroups as independent samples."
    )
    add_table(
        document,
        "Table A1.1. Nested construct-specification analytical contract",
        ["Element", "Operational rule", "Interpretive boundary"],
        [
            ["Control", "One exact value of any of the eight registered dimensions", "Defines the conditioned paper subset"],
            ["Full view", "All papers remaining after population and control selections", "Retains missing and unspecified outcome categories"],
            ["Observed view", "Papers with a substantive value for the displayed outcome dimension", "Denominator differs by outcome dimension"],
            ["Matrix", "Any two different dimensions crossed inside the active controls", "Reports counts and within-row, within-column, and analytical shares"],
            ["Population", "Full corpus, Leading entrepreneurship journals, Additional entrepreneurship, or Combined entrepreneurship", "Populations may overlap with FT50 and business domains"],
            ["Inference", "Descriptive composition and theoretically guided comparison", "No temporal, causal, or independent-group claim"],
        ],
    )
    add_table(
        document,
        "Table A1.2. Registered dimensions available as controls and outcomes",
        ["Dimension", "Function in the analysis"],
        [
            ["Study status", "Separates AI as phenomenon, research method, both, or unclear"],
            ["AI role", "Identifies the theoretical work assigned to AI"],
            ["Technical type", "Identifies the named technical form"],
            ["Mechanism", "Identifies what AI observably changes or enables"],
            ["Level", "Locates the focal theoretical relation"],
            ["Process stage", "Identifies the visible entrepreneurial or organisational stage"],
            ["Scope", "Identifies the stated boundary or embedding condition"],
            ["Definition", "Identifies the definitional signal visible in the abstract record"],
        ],
    )

    document.add_heading("A2. Platform-to-Appendix Reproduction Procedure", level=1)
    document.add_paragraph(
        "The platform's Dataset scope, Coding model, Filter dimension, Filter value, Study status, and "
        "Distribution controls define the analytical state. Every visible bar or matrix cell is backed by the "
        "same paper-level frame used to generate this appendix. Selecting a cell opens its title, abstract, "
        "author keywords, coding evidence, metadata, and Scopus record. The evidence panel can retain all "
        "supporting papers, any exact two-, three-, or four-model match, or the registered Mini-Claude-Gemini "
        "sweet spot; every model's assignment remains visible. The frozen export records the control "
        "dimension and value, outcome dimension, denominator, category, paper count, and share so that every "
        "reported percentage can be reconstructed without relying on the interface."
    )
    add_table(
        document,
        "Table A2.1. Platform operation and frozen research artifact",
        ["Platform operation", "Frozen output", "Reproducible content"],
        [
            ["Filter one dimension value", "nested_dimension_distributions.csv", "25,813 conditional distribution rows"],
            ["Select phenomenon, method, or both", "study_status_conditioned_specification.csv", "1,447 status-conditioned rows"],
            ["Cross any two dimensions", "nested_dimension_pair_matrices.csv", "15,859 cell-level rows"],
            ["Audit available matrices", "nested_dimension_pair_inventory.csv", "224 population-view-pair summaries"],
            ["Compare model reliability", "full_corpus_pairwise_irr_dimensions.csv", "48 estimates on one 21,930-paper balanced intersection"],
            ["Inspect convergent evidence", "full_corpus_dimension_consensus.csv", "Preferred-trio and all-model exact agreement by dimension"],
            ["Download complete release", "nested_specification_release.zip", "All tables and denominator readme"],
        ],
    )

    document.add_heading("A3. Study-Status-Conditioned Construct Specification", level=1)
    document.add_paragraph(
        "The Combined entrepreneurship population contains 1,497 papers with a clear study status: 824 treat AI "
        "as the substantive phenomenon, 385 as a research method, and 288 as both. The remaining papers have an "
        "unclear status and are retained in the unconditioned full view but not assigned to one of these three "
        "status profiles."
    )
    add_picture(
        document,
        FIG_OBSERVABILITY,
        "Figure A3.1. Observability of each outcome dimension within the three clear study-status subsets.",
        6.5,
    )
    add_table(
        document,
        "Table A3.1. Observed denominator by outcome dimension and study status",
        ["Outcome dimension", "Phenomenon", "Method", "Both"],
        observability_rows(table),
    )
    add_note(
        document,
        "Each cell reports observed papers divided by all papers in that study-status subset. A lower value "
        "means that the outcome dimension is less visible in titles, abstracts, and author keywords; it does "
        "not establish absence from the full paper.",
    )
    add_picture(
        document,
        FIG_ALL,
        "Figure A3.2. Complete observed composition of all seven outcome dimensions after conditioning on study status.",
        6.8,
    )
    add_table(
        document,
        "Table A3.2. Leading categories within each observed dimension and study status",
        ["Outcome dimension", "Study status", "Observed n", "Leading three categories"],
        leading_rows(table),
        font_size=7.6,
    )
    document.add_paragraph(
        "Method papers combine high technical observability with low mechanism observability. Machine learning "
        "dominates their named types and prediction dominates the mechanisms that are stated, but only one third "
        "of method papers expose a substantive mechanism. Phenomenon papers are less likely to name a technical "
        "type but more likely to state a mechanism and distribute that mechanism across learning, prediction, "
        "judgment, uncertainty, stakeholder interaction, and access. Both papers retain tool and context roles "
        "while also carrying a research-method component."
    )

    for index, (dimension, label) in enumerate(DIMENSION_ORDER, start=3):
        selected = combined_observed(table)
        selected = selected[selected["outcome_dimension"].eq(dimension)]
        denominators = []
        for status in STATUS_ORDER:
            frame = selected[selected["control_raw_value"].eq(status)]
            denominators.append(int(frame["denominator"].iloc[0]) if len(frame) else 0)
        add_table(
            document,
            f"Table A3.{index}. Complete {label.lower()} composition by study status",
            ["Observed category", *[STATUS_LABELS[item] for item in STATUS_ORDER]],
            detailed_dimension_rows(table, dimension),
            font_size=7.5,
        )
        add_note(
            document,
            "Observed denominators: "
            + "; ".join(
                f"{STATUS_LABELS[status]} n={denominator:,}"
                for status, denominator in zip(STATUS_ORDER, denominators)
            )
            + ". Cells report papers and the column percentage within that dimension-specific denominator.",
        )

    document.add_heading("A4. Study Status by Entrepreneurship Population", level=1)
    document.add_paragraph(
        "Core and Additional entrepreneurship do not differ uniformly. The table selects, for each status and "
        "outcome dimension, the category with the largest absolute percentage-point difference. This is a "
        "descriptive boundary diagnostic, not a journal-quality ranking or significance test."
    )
    document.add_paragraph(
        "The largest differences occur among phenomenon and both papers. Additional entrepreneurship places "
        "greater weight on generative AI, learning, innovation, and country boundaries, whereas Core places "
        "greater weight on tools, prediction, and individual entrepreneurs. The role profile of method papers is "
        "much more similar across the two populations, indicating that the aggregate population contrast is "
        "partly produced by different mixtures of study status."
    )
    add_table(
        document,
        "Table A4.1. Strongest Core-Additional contrast within each study-status and outcome-dimension combination",
        ["Status", "Dimension", "Category", "Core", "Additional", "Core minus Additional"],
        strongest_population_contrasts(),
        font_size=7.2,
    )
    add_landscape_section(document)
    document.add_heading("A5. Complete Pairwise Dimension-Matrix Inventory", level=1)
    document.add_paragraph(
        "Eight dimensions yield 28 unique unordered pairs. Each pair was calculated for four populations and two "
        "denominator views, producing 224 matrices. Table A5.1 inventories the 28 Combined entrepreneurship "
        "observed-view matrices. The complete release retains every cell, including zero and rare cells, rather "
        "than selecting only visually prominent combinations."
    )
    add_table(
        document,
        "Table A5.1. Combined entrepreneurship observed-view matrix inventory",
        ["Dimension pair", "Analysed n", "Non-zero cells", "Largest cell", "Papers", "Share of analysed"],
        pair_inventory_rows(),
        font_size=7.0,
    )
    add_note(
        document,
        "The largest cell is an audit descriptor, not by itself a theoretically retained configuration. "
        "Theoretical interpretation additionally requires recurrence, adequate support, paper inspection, a "
        "meaningful contrast, and consistency with the evidence boundary.",
    )

    add_portrait_section(document)
    document.add_heading("A6. Four-Model Coverage, Reliability, and Convergent Evidence", level=1)
    document.add_paragraph(
        "Construct-specification distributions retain each model's actual successful-paper denominator. Model "
        "inter-rater reliability uses a different rule: every pair is restricted to one exact 21,930-paper "
        "intersection shared by Mini, Nano, Claude, and Gemini inside Claude's 21,940-paper successful cohort. "
        "This prevents pair-specific denominators from changing the apparent comparison. Exact agreement and "
        "nominal Krippendorff α are reported for all eight dimensions. Because a full-cohort comparison "
        "combines agreement about whether evidence is observable with agreement about which category applies, "
        "the dimension-level results are decomposed into full-category agreement, binary observability agreement, "
        "and conditional-category agreement where both models observed evidence. Arithmetic means across the six "
        "core dimensions are orientation summaries rather than omnibus reliability coefficients."
    )
    add_table(
        document,
        "Table A6.1. Model-specific coverage and balanced comparison cohort",
        ["Model", "Role", "Usable papers", "Coverage", "Non-responses", "Balanced IRR n"],
        irr_coverage_rows(),
        font_size=7.6,
    )
    add_table(
        document,
        "Table A6.2. Pairwise reliability orientation across the six core dimensions",
        ["Model pair", "Balanced papers", "Mean exact agreement", "Mean nominal α"],
        irr_summary_rows(),
        font_size=7.6,
    )
    add_note(
        document,
        "The means compare study status, technical type, AI role, mechanism, level, and scope. The complete "
        "dimension-level results below are the analytical record. Agreement measures coding convergence, not "
        "accuracy or ground truth.",
    )
    add_landscape_section(document)
    add_table(
        document,
        "Table A6.3. Complete pair-by-dimension model reliability",
        [
            "Model pair",
            "Dimension",
            "Use",
            "Full papers",
            "Full exact",
            "Full α",
            "Observability exact",
            "Observability α",
            "Both observed",
            "Conditional exact",
            "Conditional α",
        ],
        irr_dimension_rows(),
        font_size=5.6,
    )
    add_note(
        document,
        "Full-category agreement retains each dimension's unobserved value as a category. Observability "
        "agreement collapses every rating to observed versus unobserved. Conditional-category agreement "
        "uses only papers for which both models assigned an observed category; it therefore excludes "
        "disagreement about whether evidence was observable and must be interpreted alongside the "
        "observability result. Across Claude and Gemini, the decomposition distinguishes disagreement about "
        "presence from disagreement about category. Process-stage weakness is concentrated in the "
        "observability decision, whereas definition clarity remains weak at both signal detection and "
        "classification of the observed definitional form. Mechanism has the thinnest jointly observed base "
        "among the six core dimensions (6,676 papers; the next smallest is technical type at 12,819), but its "
        "conditional agreement remains materially stronger (74.07%; α = .63). This denominator applies "
        "to mechanism-dependent structuring results and must accompany their interpretation.",
    )
    add_portrait_section(document)
    document.add_paragraph(
        "The platform additionally supports paper-level convergence filters. General cross-model agreement "
        "requires any two or more models to assign the exact selected pattern. The preferred sweet spot requires "
        "Mini, Claude, and Gemini; Nano remains visible but does not veto that criterion. The stricter unanimous "
        "column requires all four models. These counts identify high-convergence evidence for inspection and do "
        "not replace Mini's primary codes with a consensus classification."
    )
    add_table(
        document,
        "Table A6.4. Exact preferred-trio and unanimous four-model agreement by dimension",
        [
            "Dimension",
            "Use",
            "Mini + Claude + Gemini: total; unobserved; observed",
            "All four models: total; unobserved; observed",
        ],
        irr_consensus_rows(),
        font_size=7.5,
    )
    add_note(
        document,
        "The unobserved and observed counts partition each exact-agreement total. For definition clarity, "
        "11,448 of the 11,497 unanimous four-model classifications are no-definition codes; only 49 are "
        "unanimous observed definition categories. Raw unanimity must therefore not be interpreted as "
        "reliability of definitional-form discrimination.",
    )
    add_table(
        document,
        "Table A6.5. Full-corpus model-reliability research artifacts",
        ["File", "Rows", "SHA-256", "Purpose"],
        irr_artifact_rows(),
        font_size=7.0,
    )

    document.add_heading("A7. Supplementary File Manifest", level=1)
    add_table(
        document,
        "Table A7.1. Nested-analysis research artifacts",
        ["File", "Rows represented", "SHA-256", "Purpose"],
        artifact_rows(),
        font_size=7.2,
    )

    document.add_heading("A8. Topic-Modelling Optimisation and Selection Audit", level=1)
    document.add_paragraph(
        "Topic modelling used a quantitative grid followed by semantic review. The composite score combined "
        "normalised silhouette (weight 0.25), top-term diversity (0.25), one minus the raw outlier rate (0.35), "
        "and inverse topic-size dispersion (0.15). The score selected a resolution for review; it did not measure "
        "classification accuracy. Candidates with fewer than five topics remained visible in the audit but were "
        "ineligible because the smaller entrepreneurship scopes otherwise collapsed into two broad clusters. "
        "Every selected candidate was reviewed using all leading terms and three centroid-nearest papers per "
        "topic."
    )
    add_picture(
        document,
        TOPIC_GRID_FIGURE,
        "Figure A8.1. Complete BERTopic grid-search review, showing automatic recommendations and approved selections.",
        6.6,
    )
    add_table(
        document,
        "Table A8.1. Tested grids and approved data-specific topic models",
        ["Scope", "Eligible papers", "Tested minimum topic sizes", "Selected", "Topics", "Decision basis"],
        topic_selection_rows(),
        font_size=7.2,
    )
    document.add_paragraph(
        "The broad business scope contains 96.02% of the papers eligible for the Full Corpus model. The final "
        "decision therefore required both scopes to use one common tested minimum topic size. Table A8.2 shows "
        "that 50 maximised the arithmetic mean of their composite scores and produced 53 Full Corpus topics and "
        "50 broad-scope topics. This replaced an unapproved automatic 24-topic Full Corpus solution and a "
        "withdrawn 37-versus-50 review proposal."
    )
    add_table(
        document,
        "Table A8.2. Joint Full Corpus and broad-scope resolution rule",
        ["Minimum size", "Full topics", "Full score", "Broad topics", "Broad score", "Mean score", "Decision"],
        topic_joint_grid_rows(),
        font_size=7.2,
    )
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            if run.text:
                run.text = british_spelling(run.text)
    update_footers(document)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)

    MARKDOWN.write_text(
        "\n".join(
            [
                "# Supplementary Methods and Analysis Appendix",
                "",
                "This mirror records the current supplementary deliverable and its authoritative artifacts.",
                "",
                "- Main nested table: `reports/analysis/tables/contrasting/study_status_conditioned_specification.csv`",
                "- All conditional distributions: `reports/analysis/tables/contrasting/nested_dimension_distributions.csv`",
                "- All pairwise cells: `reports/analysis/tables/contrasting/nested_dimension_pair_matrices.csv`",
                "- Pair inventory: `reports/analysis/tables/contrasting/nested_dimension_pair_inventory.csv`",
                "- Complete release: `reports/analysis/tables/contrasting/nested_specification_release.zip`",
                "- Model coverage: `reports/analysis/tables/model_validation/full_corpus_model_coverage.csv`",
                "- Pairwise IRR summary: `reports/analysis/tables/model_validation/full_corpus_pairwise_irr_core_summary.csv`",
                "- Complete dimension-level IRR: `reports/analysis/tables/model_validation/full_corpus_pairwise_irr_dimensions.csv`",
                "- Topic-selection review: `data/processed/topics/optimization/topic_selection_review.json`",
                "- Topic grid review figure: `data/processed/topics/optimization/grid_search_review_overview.png`",
                "- Full Corpus and broad-scope joint grid: `data/processed/topics/optimization/full_corpus_query_1_joint_grid.csv`",
                "- Preferred-trio and four-model consensus: `reports/analysis/tables/model_validation/full_corpus_dimension_consensus.csv`",
                "",
                "The DOCX contains the full study-status tables, Core-Additional conditional contrasts, all-dimension figures, the 28-pair Combined entrepreneurship matrix inventory, and the final four-model coverage, reliability, and consensus tables.",
            ]
        ),
        encoding="utf-8",
    )
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {MARKDOWN}")


if __name__ == "__main__":
    build()
