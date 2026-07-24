"""Audit every numbered manuscript table and figure against frozen sources.

The audit is read-only with respect to analytical data. It verifies population
identities, source-title coverage, model coverage, topic decisions, contrasting
tables, numbered exhibit presence, and image readability, then writes a small
machine-readable and human-readable audit record.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from PIL import Image

from aecsp.api.graph_service import GraphService


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
MASTER = ROOT / "data/processed/master_corpus.csv"
PRIMARY = ROOT / "data/processed/analysis/primary_analysis_dataset.csv"
DOMAIN_ASSIGNMENTS = ROOT / "data/processed/analysis/theory_elaboration/domains/business_domain_assignments.csv"
DOMAIN_MANIFEST = ROOT / "data/processed/analysis/theory_elaboration/domains/business_domain_manifest.json"
TOPIC_MANIFEST = ROOT / "data/processed/topics/final_run_manifest.json"
TOPIC_REVIEW = ROOT / "data/processed/topics/optimization/topic_selection_review.json"
CONTRAST = ROOT / "reports/analysis/tables/contrasting"
MODEL_TABLES = ROOT / "reports/analysis/tables/model_validation"
OUTPUT_JSON = ROOT / "reports/analysis/manuscript_exhibit_audit.json"
OUTPUT_MD = ROOT / "reports/analysis/MANUSCRIPT_EXHIBIT_AUDIT.md"

FIGURES = {
    "Figure 1": ROOT / "reports/analysis/figures/contrasting/publication_trend_combined_entrepreneurship.png",
    "Figure 2": ROOT / "reports/analysis/figures/specification/specification_combined_entrepreneurship.png",
    "Figure 3": ROOT / "reports/analysis/figures/contrasting/nested_status_central_dimensions.png",
    "Figure 4": ROOT / "reports/analysis/figures/contrasting/specification_type_by_role.png",
    "Figure 5": ROOT / "reports/analysis/figures/contrasting/horizontal_role_by_domain_with_ent.png",
    "Figure 6": ROOT / "reports/analysis/figures/contrasting/vertical_role_by_collapsed_level.png",
    "Figure 7": ROOT / "reports/analysis/figures/contrasting/structuring_role_by_mechanism.png",
    "Figure 8": ROOT / "reports/analysis/figures/contrasting/framework_diagram.png",
}

TABLE_HEADERS = {
    "Table 1": ["Tactic", "Question", "Primary population", "Main output"],
    "Table 2": ["Stage", "Step", "Records"],
    "Table 3": ["Population or domain", "Papers retained", "Source-title coverage in this corpus"],
    "Table 4": ["Model", "Provider/runtime", "Role"],
    "Table 5": ["Dimension", "Analytical question", "Examples", "Analytical status"],
    "Table 6": ["Outcome dimension", "Phenomenon", "Method", "Both"],
    "Table 7": ["Dimension", "Category", "Core", "Additional", "Core minus Additional"],
    "Table 8": [
        "Domain",
        "Dimension",
        "Category",
        "Domain observed n",
        "Full-corpus observed n",
        "Full-corpus share",
        "Within-domain share",
        "Difference from full corpus",
    ],
    "Table 9": ["AI role", "Papers with specified level", "Leading levels within role"],
    "Table 10": ["Relation and support", "Evidence paper", "Theoretical meaning"],
}


def truthy(series: pd.Series) -> pd.Series:
    return series.astype(str).str.strip().str.lower().isin({"1", "true", "yes", "y", "x"})


def require(condition: bool, message: str, checks: list[str]) -> None:
    if not condition:
        raise AssertionError(message)
    checks.append(message)


def table_by_header(document: Document, header: list[str]):
    for table in document.tables:
        actual = [cell.text.strip() for cell in table.rows[0].cells]
        if actual == header:
            return table
    raise AssertionError(f"Manuscript table is missing header: {header}")


def audit() -> dict[str, object]:
    checks: list[str] = []
    master = pd.read_csv(MASTER, dtype=str, keep_default_na=False)
    primary = pd.read_csv(PRIMARY, dtype=str, keep_default_na=False)
    require(len(master) == 22_345 and master["paper_id"].nunique() == 22_345, "Master corpus retains 22,345 unique papers", checks)
    require(set(master["paper_id"]) == set(primary["paper_id"]), "Primary analysis table retains every master-corpus paper ID", checks)

    raw_queries = {
        "Query 1": sum(len(pd.read_csv(path, low_memory=False)) for path in sorted((ROOT / "data/queries").glob("SQ1*.csv"))),
        "Query 2": len(pd.read_csv(ROOT / "data/queries/SQ2.csv", low_memory=False)),
        "Query 3": len(pd.read_csv(ROOT / "data/queries/SQ3.csv", low_memory=False)),
        "Query 4": len(pd.read_csv(ROOT / "data/queries/SQ4.csv", low_memory=False)),
    }
    require(raw_queries == {"Query 1": 29_294, "Query 2": 818, "Query 3": 1_097, "Query 4": 1_509}, "Table 2 query counts match the retained Scopus exports", checks)
    require(len(pd.read_csv(ROOT / "data/interim/stage0_5_merged.csv", low_memory=False)) == 30_673, "Table 2 unique-publication count matches the deduplicated stage", checks)
    require(len(pd.read_csv(ROOT / "data/interim/stage1_rejected_source_title.csv", low_memory=False)) == 20, "Table 2 source-title exclusions match the rejection ledger", checks)
    require(len(pd.read_csv(ROOT / "data/interim/stage1_5_rejected_relevance.csv", low_memory=False)) == 8_308, "Table 2 relevance exclusions match the rejection ledger", checks)

    q2, q3, q4 = (truthy(primary[column]) for column in ("in_query_2", "in_query_3", "in_query_4"))
    core, additional, combined, ft50 = primary[q3], primary[q4], primary[q3 | q4], primary[q2]
    require((len(core), len(additional), len(combined), len(ft50)) == (646, 986, 1_632, 438), "Table 3 analytical-population counts match exact query flags", checks)
    require(not (q3 & q4).any() and len(combined) == len(core) + len(additional), "Combined entrepreneurship is the lossless, disjoint Core-Additional union", checks)
    source_counts = tuple(frame["Source title"].str.strip().nunique() for frame in (core, additional, combined, ft50))
    require(source_counts == (15, 13, 28, 37), "Table 3 source-title-label coverage is 15, 13, 28, and 37", checks)

    domain_manifest = json.loads(DOMAIN_MANIFEST.read_text(encoding="utf-8"))
    assignments = pd.read_csv(DOMAIN_ASSIGNMENTS, dtype=str, keep_default_na=False)
    for domain_id, item in domain_manifest["domains"].items():
        selected = assignments[assignments["domain_id"].eq(domain_id)]
        require(selected["paper_id"].nunique() == int(item["papers"]), f"Table 3 {item['label']} paper count matches its frozen assignments", checks)
        require(selected["source_title"].nunique() == int(item["represented_source_titles"]), f"Table 3 {item['label']} source-title count matches its registry", checks)
    validation = domain_manifest["validation"]
    require(
        validation["unique_assigned_papers"] == 19_553
        and validation["residual_papers"] == 2_792
        and validation["all_corpus_papers_have_official_asjc"],
        "Domain audit distinguishes the 19,553-paper ten-domain union from the 2,792-paper ASJC-classified residual",
        checks,
    )

    coverage = pd.read_csv(MODEL_TABLES / "full_corpus_model_coverage.csv")
    require(dict(zip(coverage["model_label"], coverage["successful_corpus_papers"])) == {"GPT-5.4 Mini": 22_345, "GPT-4.1 Nano": 22_335, "Claude Sonnet 5": 21_940, "Gemini 3.1 Pro Preview": 22_345}, "Table 5 model coverage matches frozen model artifacts", checks)
    irr = pd.read_csv(MODEL_TABLES / "full_corpus_pairwise_irr_core_summary.csv")
    require(len(irr) == 6 and irr["balanced_common_papers"].eq(21_930).all(), "Table 6 contains all six pairs on one 21,930-paper intersection", checks)

    topic_manifest = json.loads(TOPIC_MANIFEST.read_text(encoding="utf-8"))
    topic_review = json.loads(TOPIC_REVIEW.read_text(encoding="utf-8"))
    expected_topics = {"full_corpus": (50, 53), "query_1": (50, 50), "query_2": (8, 13), "query_3": (18, 6), "query_4": (20, 8)}
    for scope_id, (minimum, topics) in expected_topics.items():
        model = topic_manifest["models"][scope_id]
        require((int(model["min_topic_size"]), int(model["topics"])) == (minimum, topics), f"Table 8 {scope_id} parameters match the final topic manifest", checks)
        review = topic_review["scopes"][scope_id]
        require((int(review["selected_min_topic_size"]), int(review["selected_topic_count"])) == (minimum, topics), f"Table 8 {scope_id} parameters match the approved review", checks)

    status = pd.read_csv(CONTRAST / "study_status_conditioned_specification.csv", keep_default_na=False)
    require(set(status["population"]) == {"full_corpus", "core", "other", "combined"}, "Table 9 source preserves all four registered analytical populations", checks)
    horizontal = pd.read_csv(CONTRAST / "horizontal_domain_contrast_full_corpus.csv", keep_default_na=False)
    require({"Leading entrepreneurship journals", "Additional entrepreneurship", "Combined entrepreneurship", "FT50"}.issubset(set(horizontal["domain_label"])), "Figure 5 and Table 11 sources retain Leading, Additional, Combined, and FT50 explicitly", checks)
    require(len(pd.read_csv(CONTRAST / "nested_dimension_pair_inventory.csv")) == 224, "Pairwise exhibit inventory retains all 28 pairs across four populations and two views", checks)

    document = Document(MANUSCRIPT)
    for table_name, header in TABLE_HEADERS.items():
        table_by_header(document, header)
        checks.append(f"{table_name} is present with its registered header")
    captions = "\n".join(paragraph.text for paragraph in document.paragraphs)
    manuscript_text = captions + "\n" + "\n".join(
        cell.text
        for table in document.tables
        for row in table.rows
        for cell in row.cells
    )
    require(
        "Leading entrepreneurship journals increased from 47 papers through 2000 to 646 through 2026 (1,274.47%)" in manuscript_text
        and "Combined population increased from 65 to 1,632 papers (2,410.77%)" in manuscript_text,
        "Trend prose matches the platform's cumulative publication-stock calculation",
        checks,
    )
    internal_phrases = {
        "frozen corpus",
        "checksum-locked",
        "registered retrieval rules",
        "full-intersection",
        "model-irr",
        "preferred sweet spot",
        "preferred-trio",
        "checksummed",
        "final checkpoint",
        "pilot and production caches",
        "artifact identifiers",
        "paper ids",
    }
    lower_text = manuscript_text.lower()
    found_internal = sorted(phrase for phrase in internal_phrases if phrase in lower_text)
    require(
        not found_internal,
        "Reader-facing manuscript excludes internal workflow language",
        checks,
    )
    for figure_name, path in FIGURES.items():
        require(figure_name in captions, f"{figure_name} caption is present in the manuscript", checks)
        require(path.exists() and path.stat().st_size > 10_000, f"{figure_name} artifact exists and is non-empty", checks)
        with Image.open(path) as image:
            require(image.width >= 800 and image.height >= 500, f"{figure_name} PNG is readable at publication-scale dimensions", checks)

    table3 = table_by_header(document, TABLE_HEADERS["Table 3"])
    table3_rows = {row.cells[0].text.strip(): row.cells[1].text.strip() for row in table3.rows[1:]}
    require(table3_rows["Leading entrepreneurship journals"] == "646" and table3_rows["Additional entrepreneurship"] == "986" and table3_rows["Combined entrepreneurship"] == "1,632", "Rendered Table 3 displays the complete entrepreneurship populations", checks)
    for item in domain_manifest["domains"].values():
        label = (
            "Ethics and corporate social responsibility"
            if item["label"] == "Ethics and CSR"
            else item["label"]
        )
        require(
            table3_rows[label] == f"{int(item['papers']):,}",
            f"Rendered Table 3 {label} count matches the ASJC aggregation manifest",
            checks,
        )

    require(
        "2,792 papers (12.49%) all retain official ASJC codes but fall outside the selected domain rows" in manuscript_text
        and "They remain in the full-corpus comparison baseline" in manuscript_text,
        "Manuscript discloses the genuine ASJC-classified residual and full-baseline rule",
        checks,
    )

    table8 = table_by_header(document, TABLE_HEADERS["Table 8"])
    table8_rows = {
        (row.cells[0].text.strip(), row.cells[1].text.strip(), row.cells[2].text.strip()): row
        for row in table8.rows[1:]
    }
    service = GraphService()
    dimension_ids = {
        "AI role": "ai_role",
        "Mechanism": "mechanism",
        "Technical type": "technical_type",
    }
    baseline_denominators = {
        label: service.theory_horizontal_contrast(
            "gpt-5.4-mini-2026-03-17", dimension_id, "observed"
        )["baseline"]["denominator"]
        for label, dimension_id in dimension_ids.items()
    }
    for (domain, dimension, category), row in table8_rows.items():
        source = horizontal[
            horizontal["domain_label"].eq(domain)
            & horizontal["dimension_label"].eq(dimension)
            & horizontal["category"].eq(category)
            & horizontal["distribution"].eq("observed")
        ]
        require(len(source) == 1, f"Table 8 {domain} {dimension} {category} has one source row", checks)
        item = source.iloc[0]
        expected_baseline = float(item["share"]) - float(item["percentage_point_difference"]) / 100
        require(
            row.cells[3].text.strip() == f"{int(item['denominator']):,}"
            and row.cells[4].text.strip() == f"{baseline_denominators[dimension]:,}"
            and row.cells[5].text.strip() == f"{expected_baseline:.1%}"
            and row.cells[6].text.strip() == f"{float(item['share']):.1%}"
            and row.cells[7].text.strip()
            == f"{float(item['percentage_point_difference']):+.1f} pp",
            f"Table 8 {domain} {dimension} baseline is explicit and source-matched",
            checks,
        )

    require(
        "The remaining 135 papers have an unclear study status" in manuscript_text,
        "Results report the 135 Combined-entrepreneurship papers with unclear study status",
        checks,
    )
    require(
        "Nineteen of the 1,414 entrepreneurship papers with an observed AI-role code" in manuscript_text,
        "Agency interpretation discloses its 19-paper observed-role base",
        checks,
    )
    require(
        "no blind human validation was available for the current eight-dimensional instrument" in manuscript_text,
        "Human-validation limitation matches Supplementary Appendix A6",
        checks,
    )

    return {
        "status": "passed",
        "manuscript": str(MANUSCRIPT.relative_to(ROOT)),
        "tables_verified": list(TABLE_HEADERS),
        "figures_verified": list(FIGURES),
        "checks_passed": len(checks),
        "checks": checks,
    }


def main() -> None:
    result = audit()
    OUTPUT_JSON.write_text(json.dumps(result, indent=2) + "\n", encoding="utf-8")
    lines = [
        "# Manuscript Exhibit Audit",
        "",
        f"Status: **{result['status'].upper()}**",
        "",
        f"Verified {len(result['tables_verified'])} numbered tables, {len(result['figures_verified'])} numbered figures, and {result['checks_passed']} source-integrity checks.",
        "",
        "## Checks",
        "",
        *[f"- {item}" for item in result["checks"]],
        "",
    ]
    OUTPUT_MD.write_text("\n".join(lines), encoding="utf-8")
    print(f"PASS: {result['checks_passed']} checks; {len(TABLE_HEADERS)} tables; {len(FIGURES)} figures")
    print(f"Wrote {OUTPUT_JSON}")
    print(f"Wrote {OUTPUT_MD}")


if __name__ == "__main__":
    main()
