"""Build the full current-evidence methodology, results, and discussion draft.

The July 17 manuscript supplies the introduction, theory section, and document
styles.  This builder replaces the empty Methods/Results/Discussion markers
with the complete registered theory-elaboration blueprint.  It reads only
frozen corpus, validation, topic, domain, and contrasting artifacts; it does
not recode papers or change any analytical dataset.

The output is deliberately labelled ``current evidence`` because the full
Claude/Gemini extension and blind human annotation remain incomplete.  Those
pending layers are disclosed inside the validation section rather than used to
block results already supported by the frozen corpus and primary coding.
"""

from __future__ import annotations

import json
import sqlite3
import tempfile
import zipfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/ETP draft - July 17.docx"
OUTPUT = ROOT / "docs/ETP draft - full methodology results discussion - current evidence 2026-07-21.docx"
MARKDOWN = ROOT / "reports/analysis/ETP_FULL_BLUEPRINT_CURRENT_EVIDENCE.md"

PRIMARY = ROOT / "data/processed/analysis/primary_analysis_dataset.csv"
DOMAIN_MANIFEST = ROOT / "data/processed/analysis/theory_elaboration/domains/business_domain_manifest.json"
TOPIC_MANIFEST = ROOT / "data/processed/topics/final_run_manifest.json"
CONTRAST = ROOT / "reports/analysis/tables/contrasting"
VALIDATION = ROOT / "data/processed/analysis/model_validation"
VALIDATION_TABLES = ROOT / "reports/analysis/tables/model_validation"
HUMAN_DB = ROOT / "data/interim/human_validation/human_annotations.sqlite3"

FIG_SPEC = ROOT / "reports/analysis/figures/specification/specification_combined_entrepreneurship.png"
FIG_TYPE_ROLE = ROOT / "reports/analysis/figures/contrasting/specification_type_by_role.png"
FIG_HORIZONTAL = ROOT / "reports/analysis/figures/contrasting/horizontal_role_by_domain_with_ent.png"
FIG_VERTICAL = ROOT / "reports/analysis/figures/contrasting/vertical_role_by_collapsed_level.png"
FIG_STRUCTURE = ROOT / "reports/analysis/figures/contrasting/structuring_role_by_mechanism.png"
FIG_FRAMEWORK = ROOT / "reports/analysis/figures/contrasting/framework_diagram.png"

ROLE_MISSING = {"", "AI as unspecified label"}
TYPE_MISSING = {"", "unspecified AI"}
MECH_MISSING = {"", "mechanism missing"}
LEVEL_MISSING = {"", "unspecified level"}
STAGE_MISSING = {"", "process unspecified"}
SCOPE_MISSING = {"", "generalised without scope"}
DEFINITION_MISSING = {"", "no definition"}


def truthy(series: pd.Series) -> pd.Series:
    return series.astype(str).str.strip().str.lower().isin({"1", "true", "yes", "y", "x"})


def remove_from_marker(document: Document, marker: str) -> None:
    paragraph = next((p for p in document.paragraphs if p.text.strip() == marker), None)
    if paragraph is None:
        raise RuntimeError(f"Could not find manuscript marker: {marker}")
    body = document._element.body
    remove = False
    for child in list(body):
        if child is paragraph._p:
            remove = True
        if remove and child.tag != qn("w:sectPr"):
            body.remove(child)


def shade(cell, fill: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    element = OxmlElement("w:shd")
    element.set(qn("w:fill"), fill)
    properties.append(element)


def set_cell(cell, value: object, *, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(value))
    run.bold = bold
    run.font.size = Pt(8.3)


def format_table(table, *, prevent_splitting: bool = False) -> None:
    properties = table._tbl.tblPr
    borders = properties.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:color"), "B7C9D6")
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    header_properties.append(OxmlElement("w:tblHeader"))
    if prevent_splitting:
        for row in table.rows[1:]:
            row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def inject_footnotes(path: Path, footnotes: list[tuple[int, str]]) -> None:
    """Insert true Word footnotes into an already saved DOCX package."""

    if not footnotes:
        return
    word_namespace = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with tempfile.NamedTemporaryFile(
        dir=path.parent,
        prefix=f".{path.stem}-",
        suffix=".docx",
        delete=False,
    ) as temporary_file:
        temporary_path = Path(temporary_file.name)
    try:
        with zipfile.ZipFile(path, "r") as source, zipfile.ZipFile(
            temporary_path, "w", zipfile.ZIP_DEFLATED
        ) as destination:
            for item in source.infolist():
                data = source.read(item.filename)
                if item.filename == "word/footnotes.xml":
                    root = etree.fromstring(data)
                    for footnote_id, footnote_text in footnotes:
                        footnote = etree.SubElement(
                            root,
                            f"{{{word_namespace}}}footnote",
                            {f"{{{word_namespace}}}id": str(footnote_id)},
                        )
                        paragraph = etree.SubElement(
                            footnote, f"{{{word_namespace}}}p"
                        )
                        properties = etree.SubElement(
                            paragraph, f"{{{word_namespace}}}pPr"
                        )
                        etree.SubElement(
                            properties,
                            f"{{{word_namespace}}}pStyle",
                            {f"{{{word_namespace}}}val": "FootnoteText"},
                        )
                        marker_run = etree.SubElement(
                            paragraph, f"{{{word_namespace}}}r"
                        )
                        marker_properties = etree.SubElement(
                            marker_run, f"{{{word_namespace}}}rPr"
                        )
                        etree.SubElement(
                            marker_properties,
                            f"{{{word_namespace}}}rStyle",
                            {f"{{{word_namespace}}}val": "FootnoteReference"},
                        )
                        etree.SubElement(
                            marker_run, f"{{{word_namespace}}}footnoteRef"
                        )
                        text_run = etree.SubElement(
                            paragraph, f"{{{word_namespace}}}r"
                        )
                        text_element = etree.SubElement(
                            text_run, f"{{{word_namespace}}}t"
                        )
                        text_element.set(
                            "{http://www.w3.org/XML/1998/namespace}space", "preserve"
                        )
                        text_element.text = f" {footnote_text}"
                    data = etree.tostring(
                        root,
                        xml_declaration=True,
                        encoding="UTF-8",
                        standalone=True,
                    )
                destination.writestr(item, data)
        temporary_path.replace(path)
    finally:
        temporary_path.unlink(missing_ok=True)


class Writer:
    """Write matching Word and Markdown methods/results/discussion sections."""

    def __init__(self, document: Document) -> None:
        self.document = document
        self.footnotes: list[tuple[int, str]] = []
        self.markdown_footnotes: list[str] = []
        self.md: list[str] = [
            "# Full methodology, results, and discussion: current evidence",
            "",
            "> This mirror contains the completed sections added to the July 17 manuscript. "
            "The source manuscript remains unchanged.",
            "",
        ]

    def heading(self, text: str, level: int) -> None:
        self.document.add_heading(text, level=level)
        self.md.extend([f"{'#' * (level + 1)} {text}", ""])

    def paragraph(self, text: str) -> None:
        self.document.add_paragraph(text)
        self.md.extend([text, ""])

    def paragraph_with_footnote(self, text: str, footnote_text: str) -> None:
        paragraph = self.document.add_paragraph(text)
        footnote_id = len(self.footnotes) + 1
        reference_run = OxmlElement("w:r")
        reference_properties = OxmlElement("w:rPr")
        reference_style = OxmlElement("w:rStyle")
        reference_style.set(qn("w:val"), "FootnoteReference")
        reference_properties.append(reference_style)
        reference_run.append(reference_properties)
        reference = OxmlElement("w:footnoteReference")
        reference.set(qn("w:id"), str(footnote_id))
        reference_run.append(reference)
        paragraph._p.append(reference_run)
        self.footnotes.append((footnote_id, footnote_text))
        marker = f"fn{footnote_id}"
        self.md.extend([f"{text}[^{marker}]", ""])
        self.markdown_footnotes.append(f"[^{marker}]: {footnote_text}")

    def bullets(self, values: list[str]) -> None:
        for value in values:
            self.document.add_paragraph(value, style="List Bullet")
            self.md.append(f"- {value}")
        self.md.append("")

    def table(self, caption: str, headers: list[str], rows: list[list[object]]) -> None:
        p = self.document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = "Caption" if "Caption" in [s.name for s in self.document.styles] else "Normal"
        table = self.document.add_table(rows=1, cols=len(headers))
        table_style_names = {style.name for style in self.document.styles if style.type == 3}
        if "Table Grid" in table_style_names:
            table.style = "Table Grid"
        for index, header in enumerate(headers):
            set_cell(table.rows[0].cells[index], header, bold=True)
            shade(table.rows[0].cells[index], "D9EAF7")
        for row in rows:
            cells = table.add_row().cells
            for index, value in enumerate(row):
                set_cell(cells[index], value)
        format_table(table)
        self.document.add_paragraph()
        self.md.extend([f"**{caption}**", ""])
        self.md.append("| " + " | ".join(headers) + " |")
        self.md.append("| " + " | ".join("---" for _ in headers) + " |")
        for row in rows:
            clean = [str(value).replace("|", "\\|").replace("\n", " ") for value in row]
            self.md.append("| " + " | ".join(clean) + " |")
        self.md.append("")

    def picture(self, path: Path, caption: str, *, width: float = 6.8) -> None:
        if not path.exists():
            raise FileNotFoundError(path)
        p = self.document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        c = self.document.add_paragraph(caption)
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.style = "Caption" if "Caption" in [s.name for s in self.document.styles] else "Normal"
        self.md.extend([f"![{caption}]({path.relative_to(ROOT).as_posix()})", "", f"*{caption}*", ""])

    def evidence_table(self, caption: str, rows: list[dict[str, str]]) -> None:
        p = self.document.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.style = "Caption" if "Caption" in [s.name for s in self.document.styles] else "Normal"
        headers = ["Configuration and support", "Evidence paper", "Theoretical meaning"]
        table = self.document.add_table(rows=1, cols=len(headers))
        table_style_names = {style.name for style in self.document.styles if style.type == 3}
        if "Table Grid" in table_style_names:
            table.style = "Table Grid"
        for index, header in enumerate(headers):
            set_cell(table.rows[0].cells[index], header, bold=True)
            shade(table.rows[0].cells[index], "D9EAF7")
        for item in rows:
            cells = table.add_row().cells
            set_cell(
                cells[0],
                f"{item['configuration']}\nn={item['papers']}; Core / Additional={item['split']}",
            )
            cells[1].text = ""
            paragraph = cells[1].paragraphs[0]
            paragraph.add_run(item["title"] + "\n").bold = True
            paragraph.add_run(item["evidence"] + "\n")
            add_hyperlink(paragraph, "Scopus record", item["url"])
            set_cell(cells[2], item["meaning"])
        format_table(table, prevent_splitting=True)
        self.document.add_paragraph()
        self.md.extend([f"**{caption}**", ""])
        self.md.append("| Configuration | Papers | Core / Additional | Evidence paper | Theoretical meaning |")
        self.md.append("| --- | ---: | ---: | --- | --- |")
        for item in rows:
            paper = f"[{item['title']}]({item['url']}). {item['evidence']}"
            vals = [item["configuration"], item["papers"], item["split"], paper, item["meaning"]]
            vals = [str(v).replace("|", "\\|").replace("\n", " ") for v in vals]
            self.md.append("| " + " | ".join(vals) + " |")
        self.md.append("")


def distribution(frame: pd.DataFrame, column: str, excluded: set[str]) -> tuple[int, pd.Series]:
    values = frame[column].astype(str).str.strip()
    values = values[~values.isin(excluded)]
    return len(values), values.value_counts(normalize=True).mul(100)


def model_validation_rows() -> list[list[object]]:
    macro = pd.read_csv(VALIDATION_TABLES / "pairwise_macro_agreement.csv")
    pairwise = pd.read_csv(VALIDATION / "agreement_pairwise.csv")
    rows = []
    for row in macro.sort_values("percent_agreement", ascending=False).itertuples():
        match = pairwise[
            (pairwise.comparison_set == "probability_sample")
            & (pairwise.left_model == row.left_model)
            & (pairwise.right_model == row.right_model)
        ]
        if match.empty:
            match = pairwise[
                (pairwise.comparison_set == "probability_sample")
                & (pairwise.left_model == row.right_model)
                & (pairwise.right_model == row.left_model)
            ]
        comparable = int(match.iloc[0].comparable)
        rows.append(
            [
                f"{row.left_model} : {row.right_model}",
                comparable,
                f"{row.percent_agreement:.2f}",
                f"{row.krippendorff_alpha:.2f}",
            ]
        )
    return rows


def annotation_count() -> int:
    if not HUMAN_DB.exists():
        return 0
    with sqlite3.connect(HUMAN_DB) as connection:
        return int(connection.execute("SELECT COUNT(*) FROM annotations").fetchone()[0])


def representative_rows(frame: pd.DataFrame) -> list[dict[str, str]]:
    selected = [
        (
            "AI as tool × improves prediction",
            "eid:2-s2.0-105034800856",
            "AI improves performance on some knowledge-work tasks but lowers it on tasks outside the model's competence frontier.",
            "Prediction support is conditional; users still have to judge where the tool is reliable.",
        ),
        (
            "AI as firm capability × supports learning",
            "eid:2-s2.0-105018343754",
            "AI is described as an organizing capability emerging from relations among human and algorithmic actors.",
            "The capability resides in organized relations and learning, not in the technical artefact alone.",
        ),
        (
            "AI as research method × improves prediction",
            "eid:2-s2.0-105039062930",
            "Interpretable machine learning is used to predict entrepreneurial behaviour across countries and years.",
            "Here AI produces evidence for theory but is not the phenomenon being explained.",
        ),
        (
            "AI as tool × alters judgment",
            "eid:2-s2.0-105034750556",
            "AI-generated representations create artificial certainty in expert decision making.",
            "More detailed output can relocate, rather than remove, the burden of judgment and responsibility.",
        ),
        (
            "AI as context × transforms stakeholder interaction",
            "eid:2-s2.0-85217167240",
            "The review traces how intelligent machines alter affordances and value for organizational stakeholders.",
            "AI is a condition of stakeholder relations rather than simply a decision tool.",
        ),
        (
            "AI as tool × reduces uncertainty",
            "eid:2-s2.0-105017260726",
            "Generative AI expands entrepreneurial ideation while hallucinations and alien outputs create new epistemic risks.",
            "Option generation moves the bottleneck toward evaluation, plausibility, and selective commitment.",
        ),
    ]
    q3 = truthy(frame["in_query_3"])
    q4 = truthy(frame["in_query_4"])
    ent = frame[q3 | q4].copy()
    role = ent["ai_role_function"].str.strip()
    mechanism = ent["ai_mechanism_analysis"].str.strip()
    observed = ent[~role.isin(ROLE_MISSING) & ~mechanism.isin(MECH_MISSING)].copy()
    observed["configuration"] = (
        observed["ai_role_function"].str.strip() + " × " + observed["ai_mechanism_analysis"].str.strip()
    )
    rows: list[dict[str, str]] = []
    for configuration, paper_id, evidence, meaning in selected:
        group = observed[observed.configuration == configuration]
        paper = frame[frame.paper_id == paper_id]
        if paper.empty:
            continue
        item = paper.iloc[0]
        rows.append(
            {
                "configuration": configuration,
                "papers": f"{len(group):,}",
                "split": f"{int(truthy(group.in_query_3).sum())} / {int(truthy(group.in_query_4).sum())}",
                "title": item["Title"],
                "url": item["Link"] or (f"https://doi.org/{item['DOI']}" if item["DOI"] else ""),
                "evidence": evidence,
                "meaning": meaning,
            }
        )
    return rows


def build() -> None:
    frame = pd.read_csv(PRIMARY, dtype=str, keep_default_na=False)
    domain_manifest = json.loads(DOMAIN_MANIFEST.read_text(encoding="utf-8"))
    topic_manifest = json.loads(TOPIC_MANIFEST.read_text(encoding="utf-8"))
    core = frame[truthy(frame.in_query_3)].copy()
    additional = frame[truthy(frame.in_query_4)].copy()
    combined = frame[truthy(frame.in_query_3) | truthy(frame.in_query_4)].copy()
    ft50 = frame[truthy(frame.in_query_2)].copy()

    document = Document(SOURCE)
    for paragraph in document.paragraphs:
        if "increased by XXX%" in paragraph.text:
            paragraph.text = (
                "AI is ubiquitous across business disciplines (Ooi et al., 2025), but its uptake within "
                "entrepreneurship research is noteworthy. Using a broad corpus of 22,345 AI-related papers in "
                "business and management journals, we track cumulative publication growth across analytical "
                "populations. Between 2000 and 2026, the publication stock in Combined entrepreneurship grew "
                "from 65 to 1,632 papers (2,410.77%), compared with 1,172.94% among the remaining broad-corpus "
                "papers. Combined entrepreneurship grew by 113.89% between 2010 and 2020, 97.73% between 2020 "
                "and 2023, and 167.98% between 2023 and 2026. The corresponding increases were 278.49%, "
                "112.78%, and 105.87% in marketing, and 132.53%, 113.47%, and 157.28% in finance. All 2026 "
                "values reflect records available at the 8 July 2026 retrieval. These patterns show rapid "
                "growth across business research, with entrepreneurship forming an increasingly important part "
                "of that expansion."
            )
        elif paragraph.text.startswith("To answer this question, we developed an interactive web-based platform"):
            paragraph.text = (
                "To answer this question, we developed an interactive web-based platform that serves as both "
                "the analytical infrastructure for this study and a research resource. The platform contains "
                "the broad 22,345-paper business and management corpus and distinguishes the FT50, Core "
                "entrepreneurship, Additional entrepreneurship, and Combined entrepreneurship populations. For "
                "each paper, the title, abstract, and author keywords are examined across study status, AI role, "
                "technical type, observable mechanism, level of analysis, process stage, scope, and definition "
                "form. Users can reproduce filters, inspect matrices, and open the evidence papers behind each "
                "aggregate."
            )
        elif paragraph.text.startswith("Our paper contributes to the entrepreneurship literature in a few ways. First"):
            paragraph.text = (
                "Our paper contributes to entrepreneurship research in several ways. First, it provides an "
                "interactive platform for examining AI-related scholarship within a broad business and management "
                "corpus, while keeping the Core, Additional, and Combined entrepreneurship populations explicit. "
                "Publication years run through 2026, with separately labelled advance 2027 records already "
                "indexed by Scopus at the 8 July 2026 retrieval. The platform connects growth, construct "
                "specification, contrasting, topic navigation, and evidence-paper inspection, making the review "
                "transparent, reusable, and extendable rather than static."
            )
        elif paragraph.text.startswith("We use structuring in a bounded way."):
            paragraph.text = (
                "We use structuring in a bounded way. We do not claim a causal process model. Instead, we use "
                "structuring to organize the relations among the dimensions of the construct-clarification "
                "framework. This is appropriate because our aim is to clarify AI as a construct in "
                "entrepreneurship research before stronger causal relationships can be developed or tested."
            )
        elif "(Suddaby, 2017)" in paragraph.text:
            # The Zotero-backed July working file records the cited construct-
            # clarity and surplus-meaning source as Suddaby (2010).
            paragraph.text = paragraph.text.replace("(Suddaby, 2017)", "(Suddaby, 2010)")
    remove_from_marker(document, "Methods")
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    writer = Writer(document)

    writer.heading("Methods", 1)
    writer.heading("3.1 Research design and theory-elaboration logic", 2)
    writer.paragraph(
        "We used theory elaboration because the literature is neither theoretically empty nor sufficiently "
        "settled for direct theory testing. Following Fisher and Aguinis (2017), the design combines four "
        "analytical tactics. Construct specification establishes what AI represents in a paper. Horizontal "
        "contrasting asks whether that meaning changes across substantive business domains. Vertical "
        "contrasting asks whether it changes across levels of analysis. Structuring identifies recurring "
        "relations among role, mechanism, level, and scope. The tactics are sequential as an interpretive "
        "logic, but the data are cross-sectional; arrows in the framework do not represent observed temporal "
        "or causal sequences."
    )
    writer.table(
        "Table 1. Theory-elaboration design and analytical questions",
        ["Tactic", "Question", "Primary population", "Main output"],
        [
            ["Construct specification", "What does each study mean by AI?", "Core, Additional, Combined entrepreneurship", "Eight-dimensional observed portrait and AI type × role"],
            ["Horizontal contrasting", "Does the same dimension vary across business domains?", "Full corpus domains; FT50 restriction", "Within-domain matrix and percentage-point contrasts"],
            ["Vertical contrasting", "Does AI's theoretical meaning change across levels?", "Combined entrepreneurship", "Role/type/mechanism/status/stage/scope × level matrices"],
            ["Structuring", "Which theoretically interpretable combinations recur?", "Combined entrepreneurship", "Pairwise and selected three-way configurations with evidence papers"],
        ],
    )

    writer.heading("3.2 Corpus construction and integrity controls", 2)
    writer.paragraph(
        "The analytical corpus contains 22,345 unique Scopus records retrieved on 8 July 2026. Search results "
        "were combined, deduplicated, and retained when the registered retrieval rules linked an AI term with "
        "the relevant business or entrepreneurship vocabulary and source-title boundaries. Publication years "
        "run through 2026, with 15 advance records already indexed by Scopus as 2027 at retrieval. These records "
        "are retained and labelled by indexed publication year rather than described as papers published after "
        "the retrieval date. The primary analysis dataset is checksum-locked and contains one row per paper."
    )
    writer.paragraph(
        "A known limitation is that acronym matching can admit retrieval false positives. Five book-review "
        "records entered because an author name contained the standalone token BERT, which the search treated "
        "as the AI model. The study does not silently delete those rows or revise the frozen denominator at this "
        "stage. Their presence, and the possibility of similar acronym-only cases, is reported as a corpus "
        "precision limitation. This affects the interpretation of broad-corpus prevalence and topic estimates "
        "but does not convert an unspecified model code into evidence of AI relevance."
    )
    writer.table(
        "Table 2. Identification, screening, and inclusion of records",
        ["Stage", "Step", "Records"],
        [
            ["Identification", "Query 1: broad business and management sources", "29,294"],
            ["Identification", "Query 2: FT50 journals", "818"],
            ["Identification", "Query 3: Core entrepreneurship", "1,097"],
            ["Identification", "Query 4: Additional entrepreneurship", "1,509"],
            ["Identification", "Total records identified", "32,718"],
            ["Screening", "Duplicate records collapsed", "2,045"],
            ["Screening", "Unique publications", "30,673"],
            ["Screening", "Outside validated source-title universe", "20"],
            ["Screening", "Assessed for AI and business or entrepreneurship relevance", "30,653"],
            ["Screening", "Did not meet the registered relevance rule", "8,308"],
            ["Included", "Primary analytical corpus", "22,345"],
            ["Sensitivity", "Strict AI-and-entrepreneurship indicator", "2,509"],
        ],
    )
    writer.paragraph(
        "The identification and screening flow is PRISMA 2020–aligned rather than a claim that a computational "
        "bibliometric corpus is a conventional intervention review. Four searches yielded 32,718 records. "
        "Deduplication used Scopus EID first, DOI second, and normalized title-year matching as the fallback, "
        "while retaining every query-membership flag. Source-title validation and the registered relevance rule "
        "then produced the 22,345-paper analytical corpus. The stricter 2,509-paper AI-and-entrepreneurship "
        "indicator is a sensitivity lens and is not substituted for the Core, Additional, or Combined journal "
        "populations."
    )

    writer.heading("3.3 Analytical populations and business domains", 2)
    writer.paragraph(
        "The 22,345-paper table is the broad business and management corpus, not an entrepreneurship-only "
        "corpus. Core entrepreneurship contains 646 papers in the registered leading-journal view; Additional "
        "entrepreneurship contains 986 papers in the wider entrepreneurship-journal view; their union contains "
        "1,632 papers with no cross-view overlap. FT50 contains 438 papers and is used as a robustness and "
        "boundary restriction, not as the primary horizontal contrast."
    )
    domain_rows = [
        [entry["label"], f"{entry['papers']:,}", entry["represented_source_titles"]]
        for entry in domain_manifest["domains"].values()
    ]
    writer.table(
        "Table 3. Analytical populations and represented business domains",
        ["Population or domain", "Papers", "Represented source titles"],
        [
            ["Full corpus", f"{len(frame):,}", "All retained source titles"],
            ["Core entrepreneurship", f"{len(core):,}", 15],
            ["Additional entrepreneurship", f"{len(additional):,}", 13],
            ["Combined entrepreneurship", f"{len(combined):,}", 28],
            ["FT50 restriction", f"{len(ft50):,}", "Registered FT50 view"],
        ] + domain_rows,
    )
    writer.paragraph(
        f"Business-domain assignments were applied only to journals represented in the existing corpus; no "
        f"new papers were retrieved to fill categories. The registry assigns "
        f"{domain_manifest['validation']['unique_assigned_papers']:,} unique papers to ten substantive domains. "
        f"Seventy-nine papers belong to more than one domain, so domain counts are not additive. Source "
        "classification follows the preserved Scopus-aligned journal registry and reviewed source-title aliases."
    )
    writer.paragraph(
        "The official ASJC layer was constructed from Elsevier's June 2026 Scopus Source Title List before the "
        "domain aggregation. Across 640 distinct source-title and ISSN pairs, 22,199 papers matched by exact "
        "normalized title and ISSN and 146 papers from six source pairs used documented reviewed overrides. All "
        "22,345 papers received at least one ASJC code, producing 72,734 paper-code links across 111 distinct "
        "codes. Papers inherited every source-level code rather than being forced into one primary category. The "
        "June 2026 list supplies a current source classification; it does not reconstruct the code Scopus may "
        "have assigned in each historical publication year."
    )

    writer.heading("3.4 Construct-specification instrument", 2)
    writer.paragraph(
        "Each paper was coded from its title, abstract, and author keywords only. Journal and year were supplied "
        "as descriptive metadata, while index keywords, full text, citation counts, query membership, topic "
        "assignments, and other model outputs were withheld from the coder. The evidence boundary means "
        "that a missing code denotes non-observability in those fields, not absence from the full paper. The "
        "instrument records one primary value per dimension; it cannot establish claim-level multiplicity or "
        "show that a paper simultaneously assigns several roles to AI. Observed composition is the substantive "
        "layer. Unspecified and missing categories form a diagnostic layer and remain in the full-distribution "
        "view."
    )
    writer.table(
        "Table 4. Construct-specification dimensions",
        ["Dimension", "Analytical question", "Examples"],
        [
            ["Study status", "Is AI the phenomenon, a research method, both, or unclear?", "phenomenon; method; both"],
            ["AI role", "What theoretical work is assigned to AI?", "tool; research method; context; capability; infrastructure; actor/agent"],
            ["Technical type", "Which technical form is named?", "machine learning; generative AI; LLM; NLP; automation"],
            ["Mechanism", "What does AI observably change or enable?", "prediction; learning; uncertainty; judgment; interaction"],
            ["Level", "Where is the focal relation located?", "individual; team; firm; platform; ecosystem; institution"],
            ["Process stage", "Which entrepreneurial or organizational stage is visible?", "opportunity evaluation; resource acquisition; innovation; scaling"],
            ["Scope", "What boundary or embedding condition is stated?", "sector; country; established firm; startup; platform"],
            ["Definition form", "What definitional signal is visible?", "partial; example only; explicit and aligned"],
        ],
    )
    writer.paragraph(
        "Mechanism is reported as one reader-facing dimension. Internally, the analysis applies a preregistered "
        "empty-logic rule so a generic mechanism label is not treated as substantive when the accompanying "
        "causal logic is absent. Definition form and process stage are retained as exploratory diagnostics; "
        "definition visibility in an abstract is not used as a verdict on full-paper quality."
    )
    writer.paragraph(
        "The coding discipline required evidence before classification. For every dimension the response stored "
        "a short supporting quotation or close paraphrase, an epistemic status of stated, inferred, or absent, "
        "and a dimension-specific confidence value. The structured response also retained mechanism logic, named "
        "theories, full-text-review flags, and an adversarial self-review. Instrument calibration occurred on "
        "separate pilot outputs from the target corpus; pilot and production caches were never pooled. The final "
        "system prompt fixed the category definitions, evidence boundary, response schema, and rule that a "
        "substantive mechanism required a non-empty statement of what AI changed or enabled."
    )

    writer.heading("3.5 Model execution and validation", 2)
    writer.paragraph(
        "GPT-5.4 Mini supplies the current complete 22,345-paper coding used for population analyses. GPT-4.1 "
        "Nano supplies an additional near-complete baseline. Claude Sonnet 5 and Gemini 3.1 Pro Preview were "
        "run on the fixed 2,235-paper stratified probability sample. The sample was drawn independently of model "
        "outputs using publication era, query provenance, abstract length, journal coverage, and metadata "
        "completeness. Mini and Gemini cover all 2,235 sampled papers; Nano and Claude each have one "
        "non-response, leaving an exact four-model intersection of 2,233 papers."
    )
    writer.paragraph(
        "Before proprietary coding, the adequacy of the probability-sample fraction was tested on the "
        "22,335-paper Mini-Nano intersection. With seed 20260711, 1,000 stratified samples were drawn at "
        "10%, 25%, and 40%. For all eight coding fields, replicate exact agreement and nominal Krippendorff "
        "alpha estimates were compared with their full-intersection values. The declared limits were absolute "
        "bias no greater than 0.01 and empirical 95% width no greater than 0.05 for exact agreement, and bias "
        "no greater than 0.02 and width no greater than 0.10 for alpha. All fractions passed. At 10%, the "
        "largest absolute bias was 0.0056 and the largest empirical 95% width was 0.0574, supporting the "
        "smallest tested fraction rather than treating 10% as a universal reliability rule."
    )
    writer.paragraph(
        "The frozen draw used seed 20260712. It contains 500 papers through 2015, 283 from 2016-2020, and "
        "1,452 from 2021 onward; the short, medium, and long abstract bands contain 742, 748, and 745 papers. "
        "It includes 88 FT50, 93 Core entrepreneurship, and 108 Additional entrepreneurship memberships, "
        "which may overlap across query flags; 106, 441, and 1,688 papers come from small, medium, and large "
        "journal strata; and 167 papers lack author keywords. The maximum absolute sample-population share "
        "difference was 1.98 percentage points. Selection probabilities and inverse-probability weights are "
        "retained, and non-responses remain missing rather than being replaced to improve agreement."
    )
    writer.table(
        "Table 5. Probability-sample model reliability across six core dimensions",
        ["Model pair", "Common papers", "Mean exact agreement", "Mean nominal alpha"],
        model_validation_rows(),
    )
    writer.paragraph(
        "The macro statistic averages study status, technical type, AI role, mechanism, level, and scope. "
        "Technical type is the most stable individual dimension. Process stage and definition form are shown "
        "separately because their chance-corrected reliability is weak. Agreement demonstrates coding "
        "consistency, not truth. No blind human annotation has yet been completed "
        f"(current saved annotation records: {annotation_count()}); claims of human-validated accuracy are "
        "therefore withheld. Claude and Gemini are being extended to the full corpus, but their pending coverage "
        "does not change the coding definitions used in the current analyses."
    )

    writer.heading("3.6 Topic modelling as a navigation layer", 2)
    models = topic_manifest["models"]
    writer.paragraph(
        "Topic modelling followed specification coding and was kept analytically separate from it. Titles, "
        "abstracts, and keywords were converted into hybrid phrase documents and embedded with "
        "sentence-transformers/all-MiniLM-L6-v2. BERTopic used UMAP dimensionality reduction and HDBSCAN "
        "clustering. Candidate minimum topic sizes were compared using silhouette, topic diversity, raw outlier "
        "rate, and topic-size balance. The declared composite weighted these components 0.25, 0.25, 0.35, and "
        "0.15 respectively. The full corpus and broad-query model used the common best reviewed minimum size of "
        "50; the smaller scopes were optimized independently."
    )
    writer.table(
        "Table 6. Approved data-specific topic models",
        ["Scope", "Papers", "Minimum topic size", "Topics"],
        [
            ["Full corpus", f"{models['full_corpus']['eligible_papers']:,}", models['full_corpus']['min_topic_size'], models['full_corpus']['topics']],
            ["Broad business and management", f"{models['query_1']['papers']:,}", models['query_1']['min_topic_size'], models['query_1']['topics']],
            ["FT50", f"{models['query_2']['papers']:,}", models['query_2']['min_topic_size'], models['query_2']['topics']],
            ["Core entrepreneurship", f"{models['query_3']['papers']:,}", models['query_3']['min_topic_size'], models['query_3']['topics']],
            ["Additional entrepreneurship", f"{models['query_4']['papers']:,}", models['query_4']['min_topic_size'], models['query_4']['topics']],
        ],
    )
    writer.paragraph(
        "The full-corpus model retained 4,322 conservative outliers and one title-only record below the usable-"
        "text threshold; all records remain in the master table with blank topic fields where unassigned. Low "
        "silhouette scores indicate overlapping thematic boundaries. Topics are therefore used to navigate "
        "papers and inspect conversations, not as objective theoretical categories or the source of the four "
        "theory-elaboration conclusions."
    )
    writer.paragraph(
        "The final configuration used seed 42, UMAP with 15 neighbours, five components, and cosine distance, "
        "and HDBSCAN with excess-of-mass cluster selection. Topic representations used one-to-three-word n-grams "
        "with a minimum document frequency of two. Outlier rates reported during grid search are not final "
        "unassigned rates: for the broad-query model, 10,585 papers were initially marked as outliers, 6,974 "
        "were conservatively reassigned only when their highest non-outlier topic probability reached 0.05, "
        "and 3,611 remained unassigned. Topic labels remain researcher-reviewable display labels; changing a "
        "label does not change a fitted paper assignment or stable topic identity."
    )

    writer.heading("3.7 Operationalisation of the four tactics", 2)
    writer.paragraph(
        "Construct specification was run separately for Core, Additional, and Combined entrepreneurship using "
        "both full and observed distributions and an AI-type-by-role matrix. Horizontal contrasts apply the "
        "same selected dimension to every represented business domain and report within-domain percentages, "
        "declared denominators, percentage-point differences from the full-corpus baseline, and a separate FT50 "
        "restriction. Vertical contrasts keep the registered level of analysis on one axis while allowing study "
        "status, role, technical type, mechanism, process stage, or scope on the other. Structuring uses pairwise "
        "role-mechanism, role-level, mechanism-level, and role-scope relations plus selected three-way "
        "combinations. Exact five-way cells are supplementary because only two satisfy the current support rule."
    )
    writer.paragraph(
        "For the main role-mechanism evidence table, a recurring pair required at least 20 Combined "
        "entrepreneurship papers. Meeting this threshold was necessary but not sufficient: retained pairs also "
        "had to illuminate a contrast, have evidence that could be traced to the displayed fields, and contribute "
        "to the construct-clarification or bottleneck-relocation argument. Core and Additional support counts are "
        "reported for every retained pair."
    )

    writer.heading("3.8 Researcher-led interpretation and platform implementation", 2)
    writer.paragraph(
        "Matrices were interpreted rather than narrated cell by cell. A pattern was retained only when it used "
        "substantive categories, disclosed an adequate denominator, appeared in more than one comparison or was "
        "meaningfully bounded by Core, Additional, or FT50 results, linked to supporting and contrasting papers, "
        "and did not require unseen full text. The historical workbook contained 154 rows representing 153 "
        "unique papers because one paper appeared in two topic locations. Conservative exact matching by "
        "normalized title, Scopus EID, and DOI found 136 papers in the current corpus and left 17 unmatched; "
        "no fuzzy match was accepted. The matched records also agreed with the original Scopus exports on title, "
        "abstract, author keywords, source title, and year. These 136 papers serve as purposive qualitative "
        "anchors and counterexamples, not a prevalence sample. Twenty-three occurred naturally in the fixed "
        "probability sample and remain a possible blind-human anchor; the other 113 are a separately labelled "
        "targeted-read evidence set and are not added to the probability design."
    )
    writer.paragraph(
        "The interactive platform is part of the methodology. Its Construct Specification page reproduces "
        "model, dataset, study-status, full-versus-observed, and nested dimension filters. Its Construct "
        "Contrasting page reproduces the horizontal, vertical, and structuring matrices. Aggregate cells open "
        "evidence panels containing the title, abstract, author keywords, coding evidence, metadata, and Scopus "
        "record. Research artifacts and filtered reports are generated from the same tables used by the visible "
        "charts. Human annotations and topic-label decisions are stored separately from model output, and "
        "reviewed topic names modify only the display label, never the fitted assignments. Stable paper, scope, "
        "topic, model, and artifact identifiers preserve provenance across downloads and rebuilt figures. Thus "
        "the platform functions as an inspectable, data-specific implementation of the analysis rather than a "
        "detached visualization website."
    )

    writer.heading("Results", 1)
    writer.heading("4.1 Growth and analytical population", 2)
    writer.paragraph(
        "The frozen corpus contains 22,345 papers. Cumulative indexed records increased from 1,691 by 2000 to "
        "22,330 through 2026, a 1,220.52% increase in the publication stock; the remaining 15 records are advance "
        "2027 records already present at retrieval. Growth establishes the scale of the conversation, but it is "
        "not itself evidence of theoretical accumulation. The theory-elaboration results therefore focus on the "
        "1,632-paper Combined entrepreneurship population and compare it with the broader domain structure."
    )

    writer.heading("4.2 Construct specification within entrepreneurship", 2)
    status_n, status = distribution(combined, "ai_method_or_phenomenon", {"", "unclear"})
    role_n, role = distribution(combined, "ai_role_function", ROLE_MISSING)
    type_n, ai_type = distribution(combined, "ai_type_form", TYPE_MISSING)
    mech_n, mechanism = distribution(combined, "ai_mechanism_analysis", MECH_MISSING)
    level_n, level = distribution(combined, "level_of_analysis", LEVEL_MISSING)
    stage_n, stage = distribution(combined, "entrepreneurial_process_stage", STAGE_MISSING)
    scope_n, scope = distribution(combined, "scope_conditions", SCOPE_MISSING)
    writer.paragraph(
        f"Among the {status_n:,} entrepreneurship papers with a clear study status, {status['phenomenon']:.1f}% "
        f"treat AI as a substantive phenomenon, {status['method']:.1f}% as a research method, and "
        f"{status['both']:.1f}% as both. Among {role_n:,} substantive role codes, AI is most often a tool "
        f"({role['AI as tool']:.1f}%), followed by a research method ({role['AI as research method']:.1f}%), "
        f"context ({role['AI as context']:.1f}%), and firm capability ({role['AI as firm capability']:.1f}%). "
        f"Machine learning accounts for {ai_type['machine learning']:.1f}% of {type_n:,} named technical types, "
        f"and generative AI for {ai_type['generative AI']:.1f}%. The leading observable mechanisms are prediction "
        f"({mechanism['improves prediction']:.1f}%) and learning ({mechanism['supports learning']:.1f}%) among "
        f"{mech_n:,} papers. Firm-level claims account for {level['firm']:.1f}% of {level_n:,} specified levels. "
        f"Innovation is the leading specified stage ({stage['innovation']:.1f}% of {stage_n:,}), while sector and "
        f"country boundaries account for {scope['sector-specific']:.1f}% and {scope['country-specific']:.1f}% of "
        f"{scope_n:,} observed scope codes."
    )
    writer.picture(FIG_SPEC, "Figure 1. Observed construct composition in Combined entrepreneurship.")

    # Core versus Additional differences, observed denominators.
    contrast_rows = []
    contrast_specs = [
        ("Study status", "ai_method_or_phenomenon", {"", "unclear"}, ["method", "both"]),
        ("AI role", "ai_role_function", ROLE_MISSING, ["AI as research method", "AI as firm capability"]),
        ("Technical type", "ai_type_form", TYPE_MISSING, ["machine learning", "generative AI"]),
        ("Mechanism", "ai_mechanism_analysis", MECH_MISSING, ["improves prediction", "supports learning"]),
        ("Level", "level_of_analysis", LEVEL_MISSING, ["individual entrepreneur", "firm"]),
        ("Process stage", "entrepreneurial_process_stage", STAGE_MISSING, ["innovation", "resource acquisition"]),
        ("Scope", "scope_conditions", SCOPE_MISSING, ["country-specific", "sector-specific"]),
    ]
    for label, column, excluded, categories in contrast_specs:
        core_n, core_dist = distribution(core, column, excluded)
        add_n, add_dist = distribution(additional, column, excluded)
        for category in categories:
            contrast_rows.append(
                [
                    label,
                    category,
                    f"{core_dist.get(category, 0):.1f}% (n={core_n:,})",
                    f"{add_dist.get(category, 0):.1f}% (n={add_n:,})",
                    f"{core_dist.get(category, 0) - add_dist.get(category, 0):+.1f} pp",
                ]
            )
    writer.table(
        "Table 7. Selected Core versus Additional entrepreneurship contrasts (observed view)",
        ["Dimension", "Category", "Core", "Additional", "Core minus Additional"],
        contrast_rows,
    )
    writer.paragraph(
        "The within-entrepreneurship comparison is not flat. Core entrepreneurship is more method-oriented: "
        "research-method roles, machine learning, prediction mechanisms, individual-level analysis, resource "
        "acquisition, and opportunity evaluation are all more prominent. Additional entrepreneurship is more "
        "likely to code AI as both phenomenon and method, as a firm capability or context, as generative AI, "
        "through learning mechanisms, at the innovation stage, and within country-specific boundaries. Journal "
        "population therefore changes the observed identity of the construct and should be treated as a boundary "
        "condition rather than pooled away."
    )

    writer.picture(FIG_TYPE_ROLE, "Figure 2. AI technical type by theoretical role in Combined entrepreneurship.")
    writer.paragraph(
        "The type-by-role matrix provides the clearest evidence that a technical label does not carry a single "
        "theoretical meaning. Of 403 papers in which machine learning and a substantive role are both visible, "
        "50.1% use it as a research method and 41.2% as a tool. Of 123 comparable generative-AI papers, 56.1% "
        "use it as a tool, 20.3% as context, and 12.2% as a firm capability. General AI is distributed across "
        "capability, context, tool, infrastructure, and actor roles. Technical sameness is therefore insufficient "
        "for theoretical comparability."
    )
    type_mech = combined[
        ~combined["ai_type_form"].fillna("").isin(TYPE_MISSING)
        & ~combined["ai_mechanism_analysis"].fillna("").isin(MECH_MISSING)
    ]
    ml_mech_n, ml_mech = distribution(
        type_mech[type_mech["ai_type_form"] == "machine learning"],
        "ai_mechanism_analysis",
        MECH_MISSING,
    )
    gen_mech_n, gen_mech = distribution(
        type_mech[type_mech["ai_type_form"] == "generative AI"],
        "ai_mechanism_analysis",
        MECH_MISSING,
    )
    writer.paragraph(
        f"Technical type also differentiates observable mechanisms. Among {ml_mech_n:,} entrepreneurship "
        f"papers naming machine learning and specifying a mechanism, "
        f"{ml_mech.get('improves prediction', 0):.1f}% emphasize prediction and "
        f"{ml_mech.get('supports learning', 0):.1f}% learning. Among the {gen_mech_n:,} comparable "
        f"generative-AI papers, learning accounts for {gen_mech.get('supports learning', 0):.1f}%, "
        f"experimentation for {gen_mech.get('reshapes experimentation', 0):.1f}%, and stakeholder "
        f"interaction for {gen_mech.get('transforms stakeholder interaction', 0):.1f}%. This mechanism "
        "comparison is calculated separately from Figure 2; it is not inferred from the type-by-role chart."
    )

    writer.heading("4.3 Horizontal contrasting across business domains", 2)
    writer.picture(FIG_HORIZONTAL, "Figure 3. Horizontal contrast in AI-role composition across domains and entrepreneurship populations.")
    horizontal = pd.read_csv(CONTRAST / "horizontal_domain_contrast_full_corpus.csv")
    selections = [
        ("Operations", "ai_role", "AI as tool"),
        ("Marketing", "mechanism", "transforms stakeholder interaction"),
        ("Innovation", "technical_type", "generative AI"),
        ("Organization studies", "mechanism", "alters judgment"),
        ("Finance", "technical_type", "machine learning"),
        ("Environmental and sustainability", "mechanism", "improves prediction"),
        ("Core entrepreneurship", "ai_role", "AI as research method"),
        ("Additional entrepreneurship", "ai_role", "AI as firm capability"),
    ]
    horizontal_rows = []
    for domain, dimension, category in selections:
        match = horizontal[
            (horizontal.domain_label == domain)
            & (horizontal.dimension_id == dimension)
            & (horizontal.category == category)
        ]
        if match.empty:
            continue
        row = match.iloc[0]
        horizontal_rows.append(
            [domain, row.dimension_label, category, f"{int(row.denominator):,}", f"{row.share:.1%}", f"{row.percentage_point_difference:+.1f} pp"]
        )
    writer.table(
        "Table 8. Selected theoretically meaningful horizontal contrasts",
        ["Domain", "Dimension", "Category", "Observed denominator", "Within-domain share", "Difference from full corpus"],
        horizontal_rows,
    )
    writer.paragraph(
        "The horizontal comparison shows specialization rather than one management-wide AI construct. Operations "
        "is strongly tool-oriented; environmental and sustainability research concentrates on prediction; "
        "marketing emphasizes transformed stakeholder interaction; innovation gives greater weight to generative "
        "AI; organization studies gives greater weight to judgment; and finance remains machine-learning heavy. "
        "Entrepreneurship is comparatively less tool-dominated than the full corpus and gives more weight to "
        "research method, context, capability, learning, judgment, and generative AI. These differences indicate "
        "that domain context changes the theoretical work assigned to the same broad AI label."
    )
    writer.paragraph(
        "The FT50 restriction does not erase heterogeneity, but it reduces several domain denominators sharply. "
        "For example, AI as tool is 63.0% in FT50 information-systems papers with an observed role but 25.0% in "
        "FT50 strategy; the corresponding observed denominators are 46 and 32. These results are directional "
        "robustness evidence, not equally precise replications of the full-corpus matrix."
    )

    writer.heading("4.4 Vertical contrasting across levels", 2)
    writer.picture(FIG_VERTICAL, "Figure 4. Vertical contrast in AI role across aggregated registered levels.")
    vertical = pd.read_csv(CONTRAST / "vertical_dimension_by_level.csv")
    vertical_rows = []
    for category in ["AI as tool", "AI as research method", "AI as firm capability", "AI as actor/agent", "AI as context"]:
        subset = vertical[(vertical.row_dimension == "ai_role") & (vertical.row_value == category)].sort_values("papers", ascending=False)
        if subset.empty:
            continue
        total = int(subset.papers.sum())
        leading = "; ".join(
            f"{r.column_value} {r.share_within_row:.1%} (n={int(r.papers)})" for r in subset.head(3).itertuples()
        )
        vertical_rows.append([category, total, leading])
    writer.table(
        "Table 9. Where entrepreneurship AI roles are located",
        ["AI role", "Papers with specified level", "Leading levels within role"],
        vertical_rows,
    )
    writer.paragraph(
        "Role changes with level. Firm-capability claims are overwhelmingly firm-level (90.8%), whereas tool "
        "claims span firms (44.2%) and individual entrepreneurs (32.7%). Research-method uses extend from firms "
        "to individuals and industries. Actor/agent is rare but relatively concentrated at the individual and "
        "team levels. Mechanisms show a related pattern: learning is split mainly between firms and individuals, "
        "prediction also reaches industry-level analyses, and judgment is most concentrated at the individual "
        "level. A finding about AI-supported individual judgment therefore cannot be treated as equivalent to a "
        "finding about AI as a firm capability merely because both use the label AI."
    )

    writer.heading("4.5 Structuring recurring configurations", 2)
    writer.picture(FIG_STRUCTURE, "Figure 5. Structuring matrix linking AI role and observable mechanism.")
    writer.evidence_table(
        "Table 10. Recurring configurations, evidence papers, and theoretical meanings",
        representative_rows(frame),
    )
    writer.paragraph(
        "The pairwise configurations expose different causal stories. Tool × prediction and research method × "
        "prediction may share a technical method but place it on opposite sides of the explanandum. Capability × "
        "learning locates the effect in organizational relations and routines. Tool × judgment and tool × "
        "uncertainty show that better outputs can create new evaluation problems. Context × stakeholder "
        "interaction locates AI in a changed relational environment. The exact five-dimensional combinations are "
        "too sparse for the main argument; only two meet the current support rule, so they are not presented as "
        "stable archetypes."
    )
    writer.paragraph(
        "The field therefore has recurring role-bound fragments rather than a widely shared complete "
        "configuration. Prediction is meaningful when attached to a tool or research-method role; learning is "
        "meaningful when attached to a capability or tool role; and stakeholder interaction becomes visible "
        "mainly when AI is treated as context. These co-occurrences identify families of theoretical explanation, "
        "not temporal sequences."
    )

    writer.heading("4.5.1 Researcher-led interpretation of the strongest configurations", 3)
    writer.paragraph(
        "Close reading makes the recurring configurations theoretically legible. The 136 previously read papers "
        "are used here as qualitative anchors and counterexamples, never as a prevalence sample. The retained "
        "interpretation has one central entrepreneurship insight, one organizational condition, and one domain "
        "boundary. An agency pattern remains an open frontier. Twenty of the 23 unique papers cited in this "
        "interpretive section belong to the Core or Additional entrepreneurship populations. The three papers "
        "outside those populations are identified only as contrasting domain cases."
    )
    writer.paragraph(
        "Bottleneck relocation is the central insight. The dominant configurations pair AI with informational "
        "mechanisms: tools improve prediction, support learning, and reduce uncertainty, while AI-supported "
        "judgment is concentrated closer to individual action. Yet the Core entrepreneurship anchors show that "
        "more information does not settle the entrepreneurial decision. Chalmers et al. (2021) connect cheaper "
        "solution generation to a greater need for evaluation and selection. Ramoglou et al. (2026) describe "
        "opportunity search as adjudication among machine-generated possibilities, and Rady et al. (2026) expose "
        "the difficulty of separating plausible opportunities from hallucinated outputs. De Véricourt and Gurkan "
        "(2026) show that superior prediction need not resolve verification and reliance, while Boussioux et al. "
        "(2024) show that human-AI generation still leaves selection work to be done. Blohm et al. (2022), "
        "contrasted with the information-systems study by Fu et al. (2021), further demonstrates that predictive "
        "performance and bias-sensitive evaluation are distinct problems. The insight is therefore not that AI "
        "eliminates uncertainty, but that expanded search, generation, and prediction relocate the binding "
        "constraint toward plausibility judgment, selective commitment, calibrated reliance, and responsibility."
    )
    writer.paragraph(
        "Organizational embedding is the condition under which firms can handle the relocated bottleneck. The "
        "capability-supports-learning configuration contains 57 entrepreneurship papers, 51 from Additional and "
        "six from Core entrepreneurship. De Fano et al. (2025), Shore et al. (2024), and Abbas et al. (2026) "
        "locate value in routines, resilience, knowledge bases, and learning rather than in isolated access to a "
        "tool. Grashof and Kopka (2023) connect AI outcomes to absorptive capacity. Schwaeke et al. (2025), "
        "Ledesma Chaves et al. (2026), and Metzger et al. (2025) show the condition from below: infrastructure, "
        "skills, readiness, trust, and governance gaps constrain what SMEs can do with AI. The resulting insight "
        "is plain: firms benefit when AI is integrated across organizational routines, data, skills, learning, "
        "and governance, rather than used as a single independent option. This is the organizational condition "
        "for managing bottleneck relocation, not a separate storyline."
    )
    writer.paragraph(
        "Domain context bounds the mechanism through which the central insight operates. Entrepreneurship anchors "
        "describe AI as an innovation intermediary, an ecosystem force, and a set of technical forms with "
        "different organizational consequences (Just, 2024; Hunt & Kurdoglu, 2025; Chalmers et al., 2026). The "
        "horizontal matrices show why these accounts cannot be generalized from a shared AI label alone. The "
        "outside-domain contrast cases locate AI in loan-text signal extraction, lending-bias detection, and "
        "national information-processing institutions (Netzer et al., 2019; Fu et al., 2021; Yoon et al., 2025). "
        "These are different mechanism claims, not merely different application settings. Domain variation thus "
        "specifies where bottleneck relocation takes a predictive, learning, interactional, or judgment-centered "
        "form; it does not create an additional competing narrative."
    )
    writer.paragraph_with_footnote(
        "Agency remains an open frontier. Core and Additional entrepreneurship papers range from AI augmenting "
        "entrepreneurial cognition to AI as a teammate, relational nonhuman actor, or co-agent (Shepherd & "
        "Majchrzak, 2022; Murtinu & De Massis, 2025; Al-Bashrawi et al., 2026; Spurrier et al., 2025). Because the "
        "present evidence is abstract-observable and the instrument records one primary role per paper, these "
        "cases identify an unresolved theoretical boundary rather than support a settled claim about autonomous "
        "AI agency.",
        "To align an earlier human evaluation with the current corpus, Lada et al. (2023), which is absent from "
        "the frozen corpus, was excluded and reliability was recalculated for the remaining 14 papers. The "
        "exercise assessed allocation to the three interpretive insight families plus fragmentation, not the "
        "eight-dimension instrument. Using the Task 2 decisions recorded in the KS and MK tabs of the final IRR "
        "workbook, the researcher and human coder agreed on all 14 papers (100.0%; Cohen's kappa = 1.000). "
        "Eleven papers "
        "map to the current Core or Additional entrepreneurship populations and three to FT50 only; six of the "
        "23 current interpretive anchors were included. We use this result only as prior human triangulation of "
        "narrative allocation and as evidence that the boundaries between the insights required interpretation."
    )

    writer.heading("4.6 Integrated empirical finding", 2)
    writer.paragraph(
        "The literature does not contain one consistently portable AI construct. It contains recurring, "
        "theoretically different configurations that often share a technical label. Construct specification "
        "shows the multiplicity; horizontal contrasting shows that domains emphasize different meanings; "
        "vertical contrasting shows that roles and mechanisms move differently across levels; and structuring "
        "shows which combinations recur. The empirical answer is therefore configurational: an AI finding is "
        "theoretically interpretable only when the paper's study status, technical form, role, mechanism, level, "
        "and scope are sufficiently explicit."
    )

    writer.heading("Discussion", 1)
    writer.heading("5.1 From topical accumulation to theoretical accumulation", 2)
    writer.paragraph(
        "The rapid growth of AI publications demonstrates topical accumulation, not necessarily theoretical "
        "accumulation. Studies can use identical technical language while explaining different phenomena, using "
        "AI as an instrument, assigning agency differently, or operating at incompatible levels. Cumulation "
        "requires a comparability rule, not merely a common keyword. The present evidence suggests that findings "
        "should be compared only when their construct configurations are sufficiently aligned or when the "
        "difference between configurations is itself the theoretical object."
    )

    writer.heading("5.2 AI as a configurational construct", 2)
    writer.paragraph(
        "The construct contribution is not a new universal definition of AI. It is a framework for establishing "
        "the theoretical identity of an AI claim. Technical type says what system is named; role says what work it "
        "performs in the argument; mechanism says what it changes; level says where the relation operates; scope "
        "states its boundary; and study status separates substantive AI theory from AI-enabled research methods. "
        "None is sufficient alone. Together they specify the unit that can enter cumulative theory."
    )

    writer.heading("5.3 Entrepreneurship and bottleneck relocation", 2)
    writer.paragraph(
        "The configurations sharpen the earlier bottleneck-relocation insight. AI can expand search, generate "
        "options, improve prediction, and reduce some information costs. Yet the papers linking tools to judgment, "
        "uncertainty, and experimentation show that the entrepreneurial problem does not disappear. It moves "
        "toward evaluating possibilities, assessing plausibility, calibrating reliance, selecting what deserves "
        "commitment, and retaining responsibility for action. The same tool can improve performance inside its "
        "competence frontier and mislead outside it. More possible futures therefore create a stronger need to "
        "decide which futures are credible and worth pursuing."
    )
    writer.paragraph(
        "Organizational integration conditions whether firms benefit from this relocation. The capability × "
        "learning configuration indicates that value arises when AI is integrated across routines, data, skills, "
        "learning, governance, and human-algorithm relations rather than used as an isolated optional tool. This "
        "is not a separate story from bottleneck relocation. Integration supplies the organizational means for "
        "handling the new evaluation and commitment burden."
    )

    writer.heading("5.4 What is distinctive about entrepreneurship?", 2)
    writer.paragraph(
        "Entrepreneurship does not simply reproduce the full business corpus. It is less dominated by AI-as-tool "
        "and prediction, and it gives greater weight to research-method use, context, firm capability, learning, "
        "judgment, generative AI, and country boundaries. Its internal populations also differ: Core journals are "
        "more method-, machine-learning-, prediction-, and individual-oriented, whereas Additional journals are "
        "more generative-, capability-, learning-, innovation-, and country-oriented. Entrepreneurship's "
        "distinctiveness therefore lies less in a unique AI category than in the mix of configurations and the "
        "evaluation problems they foreground."
    )

    writer.heading("5.5 Implications for cumulative theory development", 2)
    writer.paragraph(
        "Future research should report a minimum construct configuration: whether AI is the phenomenon, method, "
        "or both; the named technical form; its theoretical role; the mechanism through which it matters; the "
        "level of analysis; and the relevant scope conditions. Entrepreneurship studies should additionally "
        "state where evaluation, commitment, and responsibility remain after AI changes prediction, search, or "
        "generation. This reporting rule enables meaningful replication, identifies genuine boundary conditions, "
        "and prevents technically similar studies from being pooled when their theoretical objects differ."
    )

    writer.heading("5.6 Methodological contribution of the platform", 2)
    writer.paragraph(
        "The platform operationalizes the theory-elaboration workflow. Researchers can move from a population-"
        "level pattern to its cells, evidence papers, coding excerpts, and source records; change the model and "
        "analytical population; compare full and observed distributions; alter the dimensions on a vertical "
        "matrix; apply the FT50 restriction; and download the exact filtered artifact. Topic humanization and "
        "blind annotation are human-in-the-loop functions rather than hidden preprocessing. Because charts, "
        "evidence panels, reports, and downloads use shared source tables, the interface is an inspectability and "
        "method-reproduction layer as well as a communication tool."
    )

    writer.heading("5.7 Limitations and current evidence boundary", 2)
    writer.paragraph(
        "The evidence is restricted to titles, abstracts, and author keywords. One primary code per dimension "
        "cannot recover multiple roles or claim-level relations within a paper. Agency configuration is not an "
        "independently observed full-text construct; the observable mechanism is used only to describe what AI "
        "changes or enables. Domains overlap and depend on journal registries. Model agreement does not establish "
        "truth, human validation is incomplete, and the full Claude/Gemini extension remains pending. Definition "
        "visibility is genre-sensitive. Topic boundaries overlap and serve navigation rather than inference. The "
        "frozen corpus also contains known acronym-driven false positives, including five book-review records, "
        "which may slightly affect broad-corpus estimates. The 136 previously read papers are purposive evidence "
        "and cannot supply prevalence estimates."
    )

    writer.picture(FIG_FRAMEWORK, "Figure 6. Construct-clarification framework and entrepreneurship implication.", width=5.7)
    writer.heading("References", 1)
    writer.paragraph(
        "Fisher, G., & Aguinis, H. (2017). Using theory elaboration to make theoretical advancements. "
        "Organizational Research Methods, 20(3), 438–464. https://doi.org/10.1177/1094428116689707"
    )
    writer.paragraph(
        "Armstrong, M. E., Tornblad, M. K., & Jones, K. S. (2020). The accuracy of interrater reliability "
        "estimates found using a subset of the total data sample: A bootstrap analysis. Proceedings of the "
        "Human Factors and Ergonomics Society Annual Meeting, 64(1), 1377–1382."
    )
    writer.paragraph(
        "Hayes, A. F., & Krippendorff, K. (2007). Answering the call for a standard reliability measure for "
        "coding data. Communication Methods and Measures, 1(1), 77–89."
    )
    writer.paragraph(
        "Hughes, J. (2021). krippendorffsalpha: An R package for measuring agreement using Krippendorff's alpha "
        "coefficient. The R Journal, 13(1), 413–425."
    )
    writer.paragraph(
        "Page, M. J., et al. (2021). The PRISMA 2020 statement: An updated guideline for reporting systematic "
        "reviews. BMJ, 372, n71. https://doi.org/10.1136/bmj.n71"
    )
    writer.paragraph(
        "Chalmers, D., MacKenzie, N. G., & Carter, S. (2021). Artificial intelligence and entrepreneurship: "
        "Implications for venture creation in the fourth industrial revolution. Entrepreneurship Theory and "
        "Practice. https://doi.org/10.1177/1042258720934581"
    )
    writer.paragraph(
        "Ramoglou, S., Chandra, Y., & Jin, Q. (2026). Opportunity search in the era of GenAI: Navigating "
        "uncertainty in an expanding universe of imaginable but unknowable futures. Journal of Management "
        "Studies. https://doi.org/10.1111/joms.70011"
    )
    writer.paragraph(
        "Rady, J., Townsend, D., & Hunt, R. (2026). From algorithmic hallucinations to alien minds: Addressing "
        "the ideator's dilemma through entrepreneurial work. Journal of Business Venturing. "
        "https://doi.org/10.1016/j.jbusvent.2025.106550"
    )
    writer.paragraph(
        "de Véricourt, F., & Gurkan, H. (2026). Is your machine better than you? You may never know. Management "
        "Science. https://doi.org/10.1287/mnsc.2023.4791"
    )
    writer.paragraph(
        "Boussioux, L., Lane, J. N., Zhang, M., Jacimovic, V., & Lakhani, K. R. (2024). The crowdless future? "
        "Generative AI and creative problem-solving. Organization Science. "
        "https://doi.org/10.1287/orsc.2023.18430"
    )
    writer.paragraph(
        "Blohm, I., Antretter, T., Sirén, C., Grichnik, D., & Wincent, J. (2022). It's a peoples game, isn't it?! "
        "A comparison between the investment returns of business angels and machine learning algorithms. "
        "Entrepreneurship Theory and Practice. https://doi.org/10.1177/1042258720945206"
    )
    writer.paragraph(
        "Fu, R., Huang, Y., & Singh, P. V. (2021). Crowds, lending, machine, and bias. Information Systems "
        "Research. https://doi.org/10.1287/isre.2020.0990"
    )
    writer.paragraph(
        "De Fano, D., Schena, R., & Russo, A. (2025). Harnessing AI ambidexterity for competitive advantage: "
        "The role of dynamic capabilities in digital innovation ecosystems. European Journal of Innovation "
        "Management. https://doi.org/10.1108/EJIM-11-2024-1404"
    )
    writer.paragraph(
        "Shore, A., Tiwari, M., Tandon, P., & Foropon, C. (2024). Building entrepreneurial resilience during "
        "crisis using generative AI: An empirical study on SMEs. Technovation. "
        "https://doi.org/10.1016/j.technovation.2024.103063"
    )
    writer.paragraph(
        "Abbas, J., Dabić, M., & Stojčić, N. (2026). Digital divide in industry 5.0: Role of generative AI "
        "knowledge bases and intellectual capital in organizational resilience performance under territorial "
        "proximity. Technovation. https://doi.org/10.1016/j.technovation.2025.103357"
    )
    writer.paragraph(
        "Grashof, N., & Kopka, A. (2023). Artificial intelligence and radical innovation: An opportunity for all "
        "companies? Small Business Economics. https://doi.org/10.1007/s11187-022-00698-3"
    )
    writer.paragraph(
        "Schwaeke, J., Peters, A., Kanbach, D. K., Kraus, S., & Jones, P. (2025). The new normal: The status quo "
        "of AI adoption in SMEs. Journal of Small Business Management. "
        "https://doi.org/10.1080/00472778.2024.2379999"
    )
    writer.paragraph(
        "Ledesma Chaves, P., Gil-Cordero, E., Navarro-García, A., & Higuera Reina, J. A. (2026). Risk factors in "
        "the adoption of artificial intelligence by SMEs: A comprehensive study. European Journal of Innovation "
        "Management. https://doi.org/10.1108/EJIM-06-2025-0719"
    )
    writer.paragraph(
        "Metzger, M., O'Reilly, S., & Mac an Bhaird, C. (2025). Generative artificial intelligence augmenting "
        "SME financial management. Technovation. https://doi.org/10.1016/j.technovation.2025.103313"
    )
    writer.paragraph(
        "Just, J. (2024). Natural language processing for innovation search: Reviewing an emerging non-human "
        "innovation intermediary. Technovation. https://doi.org/10.1016/j.technovation.2023.102883"
    )
    writer.paragraph(
        "Hunt, R. A., & Kurdoglu, R. S. (2025). Font of innovation or algorithmic deforestation? The ecosystem "
        "impacts of artificial intelligence in entrepreneurship. Journal of Business Venturing Insights. "
        "https://doi.org/10.1016/j.jbvi.2025.e00575"
    )
    writer.paragraph(
        "Chalmers, D., Hunt, R., Pachidi, S., Potočnik, K., & Townsend, D. (2026). The acceleration of artificial "
        "intelligence: Rethinking organization and work in an era of rapid technological change. Journal of "
        "Management Studies. https://doi.org/10.1111/joms.70063"
    )
    writer.paragraph(
        "Netzer, O., Lemaire, A., & Herzenstein, M. (2019). When words sweat: Identifying signals for loan "
        "default in the text of loan applications. Journal of Marketing Research. "
        "https://doi.org/10.1177/0022243719852959"
    )
    writer.paragraph(
        "Yoon, H. D., Belkhouja, M., & Dau, L. A. (2025). Privacy protection laws, national culture, and "
        "artificial intelligence innovation around the world. Journal of International Business Studies. "
        "https://doi.org/10.1057/s41267-025-00790-2"
    )
    writer.paragraph(
        "Shepherd, D. A., & Majchrzak, A. (2022). Machines augmenting entrepreneurs: Opportunities and threats "
        "at the nexus of artificial intelligence and entrepreneurship. Journal of Business Venturing. "
        "https://doi.org/10.1016/j.jbusvent.2022.106227"
    )
    writer.paragraph(
        "Murtinu, S., & De Massis, A. (2025). Artificial intelligence machines as relational nonhuman actors in "
        "entrepreneurial teams. Journal of Small Business Management. "
        "https://doi.org/10.1080/00472778.2025.2461031"
    )
    writer.paragraph(
        "Al-Bashrawi, M. A., et al. (2026). Agentic AI systems and the future of entrepreneurship: A perspective "
        "on co-agency, innovation, and ecosystem transformation. International Entrepreneurship and Management "
        "Journal. https://doi.org/10.1007/s11365-026-01164-2"
    )
    writer.paragraph(
        "Spurrier, H., Kamineni, R., Dissanayake, M., & Lindsay, N. (2025). New venture team trust: Perceptions "
        "of new human, AI and AI-augmented teammates during funding activities. Journal of Small Business "
        "Management. https://doi.org/10.1080/00472778.2025.2579191"
    )
    writer.paragraph(
        "Bertoni, F., Bonini, S., Capizzi, V., Colombo, M. G., & Manigart, S. (2022). Digitization in the market "
        "for entrepreneurial finance: Innovative business models and new financing channels. Entrepreneurship "
        "Theory and Practice, 46(5), 1120–1135. https://doi.org/10.1177/10422587211038480"
    )
    writer.paragraph(
        "Grimes, M., von Krogh, G., Feuerriegel, S., Rink, F., & Gruber, M. (2023). From scarcity to abundance: "
        "Scholars and scholarship in an age of generative artificial intelligence. Academy of Management "
        "Journal. https://doi.org/10.5465/amj.2023.4006"
    )
    writer.paragraph(
        "Kaminski, J. C., & Hopp, C. (2020). Predicting outcomes in crowdfunding campaigns with textual, visual, "
        "and linguistic signals. Small Business Economics, 55(3), 627–649. "
        "https://doi.org/10.1007/s11187-019-00218-w"
    )
    writer.paragraph(
        "Obschonka, M., & Fisch, C. (2022). Artificial intelligence and entrepreneurship research. In Oxford "
        "Research Encyclopedia of Business and Management. "
        "https://doi.org/10.1093/acrefore/9780190224851.013.298"
    )
    writer.paragraph(
        "Obschonka, M., Grégoire, D. A., Nikolaev, B., Ooms, F., Lévesque, M., Pollack, J. M., & Behrend, T. S. "
        "(2025). Artificial intelligence and entrepreneurship: A call for research to prospect and establish "
        "the scholarly AI frontiers. Entrepreneurship Theory and Practice, 49(3), 620–641. "
        "https://doi.org/10.1177/10422587241304676"
    )
    writer.paragraph(
        "Ooi, K.-B., Koohang, A., Aw, E. C.-X., Cham, T.-H., Cobanoglu, C., Dennis, C., Dwivedi, Y. K., Hew, "
        "J.-J., Linton Kelly, H., Hughes, L., Lin, C.-Y., Mishra, A., Phau, I., Raman, R., Sigala, M., Tang, "
        "Y.-C., Wong, L.-W., & Tan, G. W.-H. (2025). Unveiling the potential of generative artificial "
        "intelligence: A multidimensional journey into the future. Industrial Management & Data Systems. "
        "https://doi.org/10.1108/IMDS-10-2023-0703"
    )
    writer.paragraph(
        "Prüfer, J., & Prüfer, P. (2020). Data science for entrepreneurship research: Studying demand dynamics "
        "for entrepreneurial skills in the Netherlands. Small Business Economics, 55(3), 651–672. "
        "https://doi.org/10.1007/s11187-019-00208-y"
    )
    writer.paragraph(
        "Short, C. E., & Short, J. C. (2023). The artificially intelligent entrepreneur: ChatGPT, prompt "
        "engineering, and entrepreneurial rhetoric creation. Journal of Business Venturing Insights, 19, "
        "e00388. https://doi.org/10.1016/j.jbvi.2023.e00388"
    )
    writer.paragraph(
        "Suddaby, R. (2010). Editor's comments: Construct clarity in theories of management and organization. "
        "Academy of Management Review, 35(3), 346–357."
    )
    writer.paragraph(
        "Vossen, A., & Ihl, C. (2020). More than words! How narrative anchoring and enrichment help to balance "
        "differentiation and conformity of entrepreneurial products. Journal of Business Venturing, 35(6), "
        "106050. https://doi.org/10.1016/j.jbusvent.2020.106050"
    )
    writer.paragraph(
        "Elsevier. (2026). Scopus content: Source title list. Retrieved July 2026 from "
        "https://downloads.ctfassets.net/o78em1y1w4i4/7xtaTxNiNcWRTeZkV86eNy/710bfd3c7f7c7c9c88eeb3638ba4be43/ext_list_Jun_2026.xlsx"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    MARKDOWN.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)
    inject_footnotes(OUTPUT, writer.footnotes)
    if writer.markdown_footnotes:
        writer.md.extend(["", *writer.markdown_footnotes, ""])
    MARKDOWN.write_text("\n".join(writer.md), encoding="utf-8")
    print(f"Wrote {OUTPUT}")
    print(f"Wrote {MARKDOWN}")


if __name__ == "__main__":
    build()
