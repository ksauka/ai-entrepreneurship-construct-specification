"""Add approved entrepreneurship topic inventories to the KS supplement.

The two tables are generated from the final native Query 3 and Query 4 topic
assignments and the completed human topic-label review. The author-edited DOCX
is changed only between the end of Appendix A8.2 and Appendix A9.
"""

from __future__ import annotations

import hashlib
import os
import shutil
import tempfile
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
BACKUP = (
    ROOT
    / "docs/ETP supplementary material july2026 ks.before-entrepreneurship-topic-tables.docx"
)
REVIEW = ROOT / "data/processed/analysis/stage4/topic_label_review.csv"
MASTER = ROOT / "data/processed/master_corpus.csv"
OUTPUT_DIR = ROOT / "reports/analysis/tables/stage4"
LEADING_OUTPUT = OUTPUT_DIR / "leading_entrepreneurship_topic_inventory.csv"
ADDITIONAL_OUTPUT = (
    OUTPUT_DIR / "additional_entrepreneurship_topic_inventory.csv"
)

SCOPE_CONFIG = {
    "query_3": {
        "label": "Leading entrepreneurship journals",
        "expected_topics": 6,
        "expected_assigned": 646,
        "output": LEADING_OUTPUT,
        "caption": (
            "Table A8.3. Leading entrepreneurship journals: BERTopic and "
            "humanized topic inventory"
        ),
    },
    "query_4": {
        "label": "Additional entrepreneurship journals",
        "expected_topics": 8,
        "expected_assigned": 985,
        "output": ADDITIONAL_OUTPUT,
        "caption": (
            "Table A8.4. Additional entrepreneurship journals: BERTopic and "
            "humanized topic inventory"
        ),
    },
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def insert_paragraph_before(
    reference: Paragraph,
    text: str,
    *,
    bold: bool = False,
    centered: bool = False,
    keep_with_next: bool = False,
) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addprevious(element)
    paragraph = Paragraph(element, reference._parent)
    paragraph.style = "Normal"
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.JUSTIFY
    )
    paragraph.paragraph_format.keep_with_next = keep_with_next
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(9 if bold else 10)
    run.bold = bold
    return paragraph


def remove_between(start, end) -> None:
    current = start
    parent = current.getparent()
    while current is not None and current is not end:
        following = current.getnext()
        parent.remove(current)
        current = following
    if current is None:
        raise RuntimeError("Topic-table block did not terminate at Appendix A9")


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell(cell, text: object, *, bold: bool = False, size: float = 7.0) -> None:
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    if bold:
        shade_cell(cell)


def prevent_row_split(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:cantSplit")) is None:
        properties.append(OxmlElement("w:cantSplit"))


def repeat_table_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    if properties.find(qn("w:tblHeader")) is None:
        header = OxmlElement("w:tblHeader")
        header.set(qn("w:val"), "true")
        properties.append(header)


def insert_table_before(
    document: Document,
    reference: Paragraph,
    rows: list[list[object]],
) -> None:
    headers = [
        "BERTopic raw topic",
        "Humanized topic",
        "Journal count",
        "Top 10 keywords",
    ]
    table = document.add_table(rows=1, cols=4)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True, size=7.0)
    repeat_table_header(table.rows[0])
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell(row.cells[index], value, size=7.0)
        prevent_row_split(row)
    widths = [1.85, 1.85, 0.72, 3.30]
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        for row in table.rows:
            row.cells[index].width = Inches(width)
    reference._p.addprevious(table._tbl)


def normalize_terms(value: str) -> list[str]:
    terms = [term.strip() for term in str(value).split(";") if term.strip()]
    if len(terms) != 10:
        raise RuntimeError(
            f"Expected exactly ten BERTopic terms; found {len(terms)} in {value!r}"
        )
    return terms


def build_topic_tables() -> dict[str, pd.DataFrame]:
    review = pd.read_csv(REVIEW, dtype=str, keep_default_na=False)
    master = pd.read_csv(
        MASTER,
        usecols=["paper_id", "Source title"],
        dtype=str,
        keep_default_na=False,
    )
    outputs: dict[str, pd.DataFrame] = {}
    for scope, config in SCOPE_CONFIG.items():
        assignments = pd.read_csv(
            ROOT / f"data/processed/topics/native/{scope}/assignments.csv",
            dtype=str,
            keep_default_na=False,
        )
        if assignments["paper_id"].duplicated().any():
            raise RuntimeError(f"{scope} topic assignments contain duplicate papers")
        assignment_detail = assignments.merge(
            master,
            on="paper_id",
            how="left",
            validate="one_to_one",
        )
        if assignment_detail["Source title"].eq("").any():
            raise RuntimeError(f"{scope} contains assignments without source journals")
        journal_counts = (
            assignment_detail.groupby("native_topic_id")["Source title"]
            .nunique()
            .to_dict()
        )

        selected = review[review["scope"].eq(scope)].copy()
        selected["topic_sort"] = pd.to_numeric(
            selected["topic_id"], errors="raise"
        )
        selected = selected.sort_values("topic_sort")
        if len(selected) != config["expected_topics"]:
            raise RuntimeError(
                f"{scope} expected {config['expected_topics']} reviewed topics; "
                f"found {len(selected)}"
            )
        if not selected["review_status"].eq("approved").all():
            raise RuntimeError(f"{scope} contains unapproved humanized labels")
        if int(selected["final_assigned_papers"].astype(int).sum()) != config[
            "expected_assigned"
        ]:
            raise RuntimeError(f"{scope} assigned-paper total changed")

        records = []
        for row in selected.itertuples(index=False):
            terms = normalize_terms(row.top_terms)
            topic_id = str(row.topic_id)
            records.append(
                {
                    "bertopic_topic": f"T{topic_id}: {row.automatic_label}",
                    "humanized_topic": row.approved_label,
                    "journal_count": int(journal_counts.get(topic_id, 0)),
                    "top_10_keywords": "; ".join(terms),
                    "assigned_papers": int(row.final_assigned_papers),
                    "scope": scope,
                    "scope_label": config["label"],
                }
            )
        output = pd.DataFrame(records)
        if output["journal_count"].le(0).any():
            raise RuntimeError(f"{scope} contains a topic with no represented journal")
        config["output"].parent.mkdir(parents=True, exist_ok=True)
        output.to_csv(config["output"], index=False, encoding="utf-8-sig")
        outputs[scope] = output
    return outputs


def update_research_file_manifest(document: Document) -> None:
    matches = [
        table
        for table in document.tables
        if len(table.rows) > 0
        and len(table.rows[0].cells) == 4
        and table.rows[0].cells[0].text.strip() == "File"
        and table.rows[0].cells[1].text.strip() == "Rows represented"
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one research-file manifest table; found {len(matches)}"
        )
    table = matches[0]
    additions = {
        LEADING_OUTPUT.name: (
            "6",
            sha256(LEADING_OUTPUT)[:16] + "…",
            "Raw and approved topic labels, journal counts, and ten defining terms for Leading entrepreneurship",
        ),
        ADDITIONAL_OUTPUT.name: (
            "8",
            sha256(ADDITIONAL_OUTPUT)[:16] + "…",
            "Raw and approved topic labels, journal counts, and ten defining terms for Additional entrepreneurship",
        ),
    }
    existing = {
        row.cells[0].text.strip(): row
        for row in table.rows[1:]
        if row.cells[0].text.strip()
    }
    for filename, values in additions.items():
        row = existing.get(filename) or table.add_row()
        for index, value in enumerate((filename, *values)):
            set_cell(row.cells[index], value, size=6.0)
        prevent_row_split(row)
    repeat_table_header(table.rows[0])


def update_document(outputs: dict[str, pd.DataFrame]) -> None:
    document = Document(SUPPLEMENT)
    a9 = find_paragraph(document, "A9. Systematic Close Reading")
    existing = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(
            "Tables A8.3 and A8.4 report the final entrepreneurship"
        )
    ]
    if existing:
        if len(existing) != 1:
            raise RuntimeError("Multiple entrepreneurship topic-table blocks found")
        remove_between(existing[0]._p, a9._p)
    else:
        stale_caption = [
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith("Table A8.3.")
        ]
        if stale_caption:
            remove_between(stale_caption[0]._p, a9._p)

    insert_paragraph_before(
        a9,
        "Tables A8.3 and A8.4 report the final entrepreneurship topic "
        "inventories. The BERTopic column preserves the unedited automatic "
        "topic label, the humanized column reports the approved researcher "
        "label, journal count is the number of distinct source titles among "
        "papers assigned to that topic, and keywords retain the ten defining "
        "BERTopic terms in their fitted order.",
    )
    for scope in ("query_3", "query_4"):
        config = SCOPE_CONFIG[scope]
        caption = insert_paragraph_before(
            a9,
            config["caption"],
            bold=True,
            centered=True,
            keep_with_next=True,
        )
        if scope == "query_4":
            caption.paragraph_format.page_break_before = True
        output = outputs[scope]
        rows = [
            [
                row.bertopic_topic,
                row.humanized_topic,
                int(row.journal_count),
                row.top_10_keywords,
            ]
            for row in output.itertuples(index=False)
        ]
        insert_table_before(document, a9, rows)
        note = (
            f"Note. The {config['label']} model contains "
            f"{len(output)} approved topics. Journal counts are topic-specific "
            "and are not additive because the same journal may contribute papers "
            "to multiple topics. Underscores identify multiword terms retained "
            "from the fitted BERTopic vocabulary."
        )
        if scope == "query_4":
            note += (
                " One of the 986 Additional entrepreneurship papers remained "
                "without a final topic assignment; the eight topic rows therefore "
                "cover 985 papers."
            )
        insert_paragraph_before(a9, note)

    update_research_file_manifest(document)
    if not BACKUP.exists():
        shutil.copy2(SUPPLEMENT, BACKUP)
    with tempfile.NamedTemporaryFile(
        prefix=SUPPLEMENT.stem + ".",
        suffix=".docx",
        dir=SUPPLEMENT.parent,
        delete=False,
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        os.replace(temporary, SUPPLEMENT)
    finally:
        temporary.unlink(missing_ok=True)


def main() -> None:
    outputs = build_topic_tables()
    update_document(outputs)
    print(f"Updated {SUPPLEMENT}")
    print(f"Wrote {LEADING_OUTPUT}")
    print(f"Wrote {ADDITIONAL_OUTPUT}")


if __name__ == "__main__":
    main()
