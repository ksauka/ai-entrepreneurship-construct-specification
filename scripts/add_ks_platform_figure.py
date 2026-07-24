#!/usr/bin/env python3
"""Insert the annotated methodological-platform Figure 1 into the KS manuscript."""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.shared import Inches, Pt


PROJECT_ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = PROJECT_ROOT / "docs/ETP draft - July2026ks.docx"
BACKUP = PROJECT_ROOT / "docs/ETP draft - July2026ks.before-platform-figure.docx"
FIGURE = PROJECT_ROOT / "reports/analysis/figures/platform/platform_implementation_annotated.png"
CAPTION = (
    "Figure 1. Interactive methodological platform. (A) Dynamic Construct Specification "
    "controls and state-specific exports. (B) Paper-level evidence, named model agreement, "
    "and the Mini-Claude-Gemini convergence filter. (C) Horizontal, vertical, and "
    "structuring matrices. (D) Paper-grounded topic-label humanization and approval. "
    "(E) Blind human annotation using the fixed instrument and separate resumable records. "
    "(F) Bounded Knowledge Graph navigation across publications, metadata, topics, and "
    "construct codes. Selecting a bar, matrix cell, topic, or graph node retains the route "
    "from an aggregate result to its supporting papers; convergence filters do not replace "
    "the primary codes."
)


def main() -> None:
    if not FIGURE.exists():
        raise FileNotFoundError(FIGURE)

    document = Document(MANUSCRIPT)
    if any(p.text.startswith("Figure 1. Interactive methodological platform") for p in document.paragraphs):
        if not BACKUP.exists():
            raise RuntimeError("Figure 1 is already present and no pre-figure backup is available")
        shutil.copy2(BACKUP, MANUSCRIPT)
        document = Document(MANUSCRIPT)

    graph_paragraph = next(
        (p for p in document.paragraphs if p.text.startswith("The Knowledge Graph workspace provides")),
        None,
    )
    if graph_paragraph is None:
        raise RuntimeError("Could not locate the Knowledge Graph methodology paragraph")
    graph_paragraph.text = graph_paragraph.text.replace(
        "It retrieves a bounded, dataset-specific seed from Neo4j rather than loading the complete graph into the browser.",
        "When Neo4j is connected, it retrieves a bounded, dataset-specific seed rather than loading the complete graph into the browser; when Neo4j is unavailable, the same scoped controls operate on a bounded dataframe seed.",
    )

    results = next((p for p in document.paragraphs if p.text.strip() == "Results"), None)
    if results is None:
        raise RuntimeError("Could not locate the Results heading")

    if not BACKUP.exists():
        shutil.copy2(MANUSCRIPT, BACKUP)

    page_break = results.insert_paragraph_before()
    page_break.add_run().add_break(WD_BREAK.PAGE)

    image_paragraph = results.insert_paragraph_before()
    image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    image_paragraph.paragraph_format.keep_with_next = True
    image_paragraph.add_run().add_picture(str(FIGURE), width=Inches(6.3))

    caption = results.insert_paragraph_before()
    caption_run = caption.add_run(CAPTION)
    caption_run.font.name = "Times New Roman"
    caption_run.font.size = Pt(10)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Inches(0.08)

    document.save(MANUSCRIPT)

    check = Document(MANUSCRIPT)
    captions = [p.text for p in check.paragraphs if p.text.startswith("Figure 1.")]
    if captions != [CAPTION]:
        raise RuntimeError(f"Unexpected Figure 1 captions after save: {captions}")
    if len(check.inline_shapes) < len(document.inline_shapes):
        raise RuntimeError("The inserted image was not retained after saving")
    print(f"Inserted Figure 1 into {MANUSCRIPT}")


if __name__ == "__main__":
    main()
