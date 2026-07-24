"""Add the Mini-to-Gemini robustness result to the researcher-edited DOCX files.

The manuscript and supplementary material are the current KS documents rather
than generated replacements.  This updater therefore inserts only identified
paragraphs, corrects the scope denominator in Table 7 to the platform's
observed-category rule, and rebuilds an Appendix A10 at the end of the
supplementary material.
"""

from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.table import Table
from docx.text.paragraph import Paragraph

from update_ks_irr_decomposition_documents import (
    find_paragraph,
    replace_paragraph,
    replace_table_contents,
    set_cell,
)


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
ROBUSTNESS = (
    ROOT / "reports/analysis/tables/model_validation/coder_robustness"
)


def insert_paragraph_after(
    reference: Paragraph,
    text: str,
    style: str | None = None,
) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addnext(element)
    paragraph = Paragraph(element, reference._parent)
    if style:
        paragraph.style = style
    paragraph.add_run(text)
    return paragraph


def upsert_paragraph_after(
    document: Document,
    reference_prefix: str,
    marker_prefix: str,
    text: str,
) -> Paragraph:
    existing = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(marker_prefix)
    ]
    if len(existing) > 1:
        raise RuntimeError(f"Duplicate robustness paragraph: {marker_prefix}")
    if existing:
        replace_paragraph(existing[0], text)
        return existing[0]
    reference = find_paragraph(document, reference_prefix)
    return insert_paragraph_after(reference, text)


def table_after(document: Document, caption_prefix: str) -> Table:
    caption = find_paragraph(document, caption_prefix)
    sibling = caption._p.getnext()
    while sibling is not None:
        if sibling.tag.endswith("}tbl"):
            return Table(sibling, caption._parent)
        if sibling.tag.endswith("}p"):
            paragraph = Paragraph(sibling, caption._parent)
            if paragraph.text.strip():
                break
        sibling = sibling.getnext()
    raise RuntimeError(f"No table follows caption {caption_prefix!r}")


def _format_top(row: object, prefix: str) -> str:
    return (
        f"{getattr(row, f'{prefix}_category')}: "
        f"{getattr(row, f'{prefix}_share'):.2%} "
        f"(n={int(getattr(row, f'{prefix}_denominator')):,})"
    )


def _leading_three(cells: pd.DataFrame, model_role: str, role: str) -> str:
    selected = cells[
        cells["model_role"].eq(model_role) & cells["role"].eq(role)
    ].sort_values("rank").head(3)
    return "; ".join(
        f"{row.level} {row.share_within_role:.2%} (n={int(row.papers):,})"
        for row in selected.itertuples()
    )


def update_manuscript() -> None:
    document = Document(MANUSCRIPT)

    upsert_paragraph_after(
        document,
        "Generative Pre-trained Transformer (GPT)-5.4 Mini",
        "GPT-5.4 Mini was designated",
        "GPT-5.4 Mini was designated as the primary coder after the fixed challenge-set check and before "
        "Claude Sonnet 5 and Gemini 3.1 Pro Preview were extended to the complete corpus and the full cross-"
        "model agreement analysis was conducted. The choice was based on compliance with the evidence and "
        "mechanism-logic rules, not on later model convergence. Because convergence does not establish accuracy, "
        "the primary designation was not changed after the agreement results were known. Instead, the five main "
        "construct-specification analyses were recomputed from the complete Gemini coding record under unchanged "
        "populations, observed-category rules, selected contrasts, and support threshold (Supplementary Appendix A10).",
    )

    upsert_paragraph_after(
        document,
        "For Claude-Gemini, conditional α",
        "Study status requires particular caution",
        "Study status requires particular caution because it conditions the nested analysis in Section 4.2.1. "
        "Its full-category nominal α is .32 for Mini-Claude, .47 for Mini-Gemini, and .63 for Claude-Gemini. "
        "It is therefore the weakest core dimension in the Mini-Claude comparison and the second weakest in "
        "Mini-Gemini, where mechanism is marginally lower. The nested results are reported as primary-coder "
        "estimates and tested directly against Gemini rather than treated as coder-invariant findings.",
    )

    upsert_paragraph_after(
        document,
        "Among the 1,497 entrepreneurship papers",
        "The Gemini re-estimation retains",
        "The Gemini re-estimation retains the same leading category in all seven aggregate dimensions, but the "
        "study-status split is not stable in magnitude. Mini classifies the 1,497 clear-status papers as 55.04% "
        "phenomenon, 25.72% method, and 19.24% both; Gemini classifies 1,455 clear-status papers as 62.34% "
        "phenomenon, 35.60% method, and 2.06% both. The aggregate identity of the leading role, technical type, "
        "mechanism, level, process stage, and scope is reproduced, while conclusions that depend on the size of "
        "the both category require coder-specific reporting (Supplementary Table A10.1).",
    )

    upsert_paragraph_after(
        document,
        "Papers coded as both combine",
        "The nested results are not completely coder-invariant",
        "The nested results are not completely coder-invariant. Gemini reproduces 17 of the 21 leading cells "
        "and 12 of the 15 cells involving core outcome dimensions. Within phenomenon papers, the leading named "
        "technical type changes from machine learning to generative AI and the leading mechanism from learning "
        "to prediction. Within both papers, the leading scope changes from sector-specific to country-specific, "
        "but Gemini assigns only 30 papers to both, with 19 observed scope codes, compared with 288 and 221 under "
        "Mini. The exploratory definition result for method papers also changes. These cells are therefore "
        "reported as coder-dependent rather than as replicated differences (Supplementary Table A10.2).",
    )

    upsert_paragraph_after(
        document,
        "The nested comparison shows that these differences",
        "The Core-Additional boundary result is the most stable",
        "The Core-Additional boundary result is the most stable part of the re-estimation. All 14 selected "
        "contrast directions are reproduced by Gemini, although several effect sizes change. The conclusion that "
        "Leading entrepreneurship journals are more method-, machine-learning-, prediction-, individual-, and resource-"
        "acquisition-oriented, while Additional entrepreneurship gives more weight to capability, generative AI, "
        "learning, innovation, and country boundaries, does not depend on using Mini as the coder "
        "(Supplementary Table A10.3).",
    )

    upsert_paragraph_after(
        document,
        "Role changes with level",
        "The leading vertical result is also reproduced",
        "The leading vertical result is also reproduced. Gemini retains the same leading level for all five "
        "reported roles: firm for tool, research-method, capability, and context roles, and individual "
        "entrepreneur for actor-like roles. Secondary level rankings and shares vary, so Table 9 is robust at the "
        "leading-location level rather than identical cell by cell (Supplementary Table A10.4).",
    )

    upsert_paragraph_after(
        document,
        "The field therefore contains recurring role-bound relations",
        "The recurring-relation results are more coder-sensitive",
        "The recurring-relation results are more coder-sensitive. At the registered 20-paper support threshold, "
        "Gemini retains three of the six selected relations: tool × prediction (96 papers), research method × "
        "prediction (75), and tool × judgement (47). Capability × learning (12), context × stakeholder "
        "interaction (6), and tool × uncertainty reduction (15) remain observable but fall below the threshold. "
        "Only two of the six primary-coder evidence papers receive the same role-mechanism relation from Gemini. "
        "The first three relations therefore have cross-coder recurrence support; the remaining relations are "
        "retained only as primary-coder patterns and close-reading interpretations, not as coder-invariant "
        "frequency findings (Supplementary Table A10.5).",
    )

    table7 = table_after(
        document,
        "Table 7. Selected Core versus Additional entrepreneurship contrasts",
    )
    table7_heading = find_paragraph(
        document, "4.2.2 Core and Additional entrepreneurship boundaries"
    )
    table7_caption = find_paragraph(
        document,
        "Table 7. Selected Core versus Additional entrepreneurship contrasts",
    )
    table7_heading.paragraph_format.keep_with_next = True
    table7_caption.paragraph_format.keep_with_next = True
    for row in table7.rows[1:]:
        dimension = row.cells[0].text.strip()
        category = row.cells[1].text.strip()
        if dimension == "Scope" and category == "country-specific":
            values = ["16.2% (n=444)", "31.8% (n=801)", "-15.6 pp"]
        elif dimension == "Scope" and category == "sector-specific":
            values = ["57.2% (n=444)", "50.2% (n=801)", "+7.0 pp"]
        else:
            continue
        for index, value in enumerate(values, start=2):
            set_cell(row.cells[index], value, font_size=7.5)

    table10_caption = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith(
                (
                    "Table 10. Recurring relations, evidence papers, and theoretical meanings",
                    "Table 10. Recurring relations under the primary coding record",
                )
            )
        ),
        None,
    )
    if table10_caption is None:
        raise RuntimeError("Could not locate the Table 10 caption")
    replace_paragraph(
        table10_caption,
        "Table 10. Recurring relations under the primary coding record, evidence papers, and theoretical meanings",
    )

    output = MANUSCRIPT.with_suffix(".tmp.docx")
    document.save(output)
    Document(output)
    output.replace(MANUSCRIPT)


def _append_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.size = Pt(9)


def _append_note(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(8.5)


def _append_table(
    document: Document,
    caption: str,
    headers: list[str],
    rows: list[list[str]],
    font_size: float = 7.0,
    page_break_before: bool = False,
    widths: list[float] | None = None,
) -> None:
    if page_break_before:
        document.add_page_break()
    _append_caption(document, caption)
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    replace_table_contents(table, headers, rows, font_size, widths)


def _remove_existing_a10(document: Document) -> None:
    start = next(
        (
            paragraph
            for paragraph in document.paragraphs
            if paragraph.text.strip().startswith("A10. Primary-Coder Robustness")
        ),
        None,
    )
    if start is None:
        return
    body = document._element.body
    previous = start._p.getprevious()
    while previous is not None and previous.tag.endswith("}p"):
        has_text = any(str(text).strip() for text in previous.itertext())
        has_page_break = any(child.tag.endswith("}br") for child in previous.iter())
        if has_text or not has_page_break:
            break
        before_previous = previous.getprevious()
        body.remove(previous)
        previous = before_previous
    element = start._p
    while element is not None:
        next_element = element.getnext()
        body.remove(element)
        element = next_element


def update_supplement() -> None:
    document = Document(SUPPLEMENT)
    _remove_existing_a10(document)

    aggregate = pd.read_csv(ROBUSTNESS / "aggregate_leading_categories.csv")
    nested = pd.read_csv(
        ROBUSTNESS / "nested_leading_categories_by_study_status.csv"
    )
    contrasts = pd.read_csv(
        ROBUSTNESS / "core_additional_selected_contrasts.csv"
    )
    role_cells = pd.read_csv(ROBUSTNESS / "role_by_level_cells.csv")
    role_comparison = pd.read_csv(
        ROBUSTNESS / "role_by_level_leading_comparison.csv"
    )
    relations = pd.read_csv(ROBUSTNESS / "selected_recurring_relations.csv")
    summary = json.loads((ROBUSTNESS / "summary.json").read_text(encoding="utf-8"))

    document.add_page_break()
    heading = document.add_paragraph()
    run = heading.add_run("A10. Primary-Coder Robustness: Gemini Re-estimation")
    run.bold = True
    run.font.size = Pt(15)
    document.add_paragraph(
        "GPT-5.4 Mini was designated as the primary coder before the complete Gemini run and full cross-model "
        "agreement analysis. The designation followed the fixed challenge-set compliance check and was not "
        "selected from the later agreement results. To test whether that choice changes the substantive findings, "
        "the five analyses named in the Results were recomputed from Gemini's complete 22,345-paper record. The "
        "Core, Additional, and Combined entrepreneurship populations, observed-category exclusions, selected "
        "contrasts, and 20-paper relation threshold were held constant. This is a robustness re-estimation, not a "
        "consensus coding rule and not evidence that either coder is correct."
    )

    aggregate_rows = [
        [
            row.dimension_label,
            _format_top(row, "primary"),
            _format_top(row, "alternative"),
            "Yes" if row.same_leading_category else "No",
        ]
        for row in aggregate.itertuples(index=False)
    ]
    _append_table(
        document,
        "Table A10.1. Aggregate leading-category robustness in Combined entrepreneurship",
        ["Dimension", "GPT-5.4 Mini", "Gemini 3.1 Pro Preview", "Same leader"],
        aggregate_rows,
        widths=[1.2, 1.8, 1.8, 0.8],
    )
    _append_note(
        document,
        "Note. All seven aggregate leading categories are reproduced, but study-status shares differ materially: "
        "Mini assigns 19.24% of clear-status papers to both, compared with 2.06% under Gemini.",
    )

    nested_rows = [
        [
            row.study_status,
            f"{row.dimension_label} ({row.analytical_status})",
            _format_top(row, "primary"),
            _format_top(row, "alternative"),
            "Yes" if row.same_leading_category else "No",
        ]
        for row in nested.itertuples(index=False)
    ]
    _append_table(
        document,
        "Table A10.2. Nested leading-category robustness after conditioning on study status",
        ["Study status", "Outcome", "GPT-5.4 Mini", "Gemini 3.1 Pro Preview", "Same leader"],
        nested_rows,
        5.8,
        widths=[0.75, 1.25, 1.7, 1.7, 0.65],
    )
    _append_note(
        document,
        "Note. Seventeen of 21 leading cells are reproduced, including 12 of 15 core outcome cells. Gemini's both "
        "subset contains 30 papers, compared with 288 under Mini; inferences within that subgroup are therefore "
        "coder-dependent as well as based on a small alternative-coder denominator.",
    )

    contrast_rows = [
        [
            f"{row.dimension_label}: {row.category}",
            f"Core {row.primary_core_share:.2%}; Additional {row.primary_additional_share:.2%}; {row.primary_difference_pp:+.2f} pp",
            f"Core {row.alternative_core_share:.2%}; Additional {row.alternative_additional_share:.2%}; {row.alternative_difference_pp:+.2f} pp",
            "Yes" if row.direction_preserved else "No",
        ]
        for row in contrasts.itertuples(index=False)
    ]
    _append_table(
        document,
        "Table A10.3. Core-Additional contrast robustness",
        ["Dimension and category", "GPT-5.4 Mini", "Gemini 3.1 Pro Preview", "Direction retained"],
        contrast_rows,
        5.8,
        page_break_before=True,
        widths=[1.85, 1.75, 1.75, 0.75],
    )
    _append_note(
        document,
        "Note. All 14 selected contrast directions are reproduced. Shares use each coder's observed denominator "
        "for the displayed dimension; missing and unspecified categories are not included.",
    )

    role_rows = [
        [
            row.role,
            _leading_three(role_cells, "primary", row.role),
            _leading_three(role_cells, "alternative", row.role),
            "Yes" if row.same_leading_level else "No",
        ]
        for row in role_comparison.itertuples(index=False)
    ]
    _append_table(
        document,
        "Table A10.4. Role-by-level robustness",
        ["AI role", "GPT-5.4 Mini: leading levels", "Gemini: leading levels", "Same leading level"],
        role_rows,
        5.8,
        widths=[1.25, 2.2, 2.2, 0.65],
    )
    _append_note(
        document,
        "Note. The leading level is reproduced for all five reported roles, although the secondary rankings and "
        "within-role shares are not identical.",
    )

    relation_rows = [
        [
            row.relation,
            f"{int(row.primary_papers):,} ({int(row.primary_core_papers):,}/{int(row.primary_additional_papers):,})",
            f"{int(row.alternative_papers):,} ({int(row.alternative_core_papers):,}/{int(row.alternative_additional_papers):,})",
            "Yes" if row.retained_by_both else "No",
            "Yes" if row.alternative_evidence_paper_matches_relation else "No",
        ]
        for row in relations.itertuples(index=False)
    ]
    _append_table(
        document,
        "Table A10.5. Selected recurring-relation robustness",
        [
            "Role-mechanism relation",
            "Mini papers (Core/Additional)",
            "Gemini papers (Core/Additional)",
            "At least 20 under both",
            "Selected evidence paper matches under Gemini",
        ],
        relation_rows,
        5.6,
        widths=[1.75, 1.25, 1.25, 0.8, 1.0],
    )
    _append_note(
        document,
        "Note. Three of six selected relations meet the 20-paper threshold under both coders, and two of the six "
        "primary-coder evidence papers receive the same role-mechanism relation from Gemini. The complete frozen "
        "distributions and matrix cells are archived in reports/analysis/tables/model_validation/coder_robustness/. "
        f"The overall result is therefore {summary['summary']['conclusion']}, not coder-invariant replication.",
    )

    output = SUPPLEMENT.with_suffix(".tmp.docx")
    document.save(output)
    Document(output)
    output.replace(SUPPLEMENT)


def main() -> None:
    update_manuscript()
    update_supplement()
    print(f"Updated {MANUSCRIPT}")
    print(f"Updated {SUPPLEMENT}")


if __name__ == "__main__":
    main()
