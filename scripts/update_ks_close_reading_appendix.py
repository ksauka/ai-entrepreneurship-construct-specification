"""Align the current manuscript and Appendix A9 with the auditable reading set.

The script deliberately avoids reconstructing the retired workbook-selection
history. It reports the current-corpus reading ledger, its registered
entrepreneurship-population membership, current data-specific topic
assignments, probability-sample overlap, and the papers' positions in the
available Query 3 and Query 4 VOSviewer document maps. VOS total link strength
is reported as a post hoc network-position check, not rewritten as the
historical selection criterion.
"""

from __future__ import annotations

import html
import math
import os
from pathlib import Path
import re
import tempfile
import unicodedata

import pandas as pd
from docx import Document
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
MATCHED = ROOT / "data/interim/theory_elaboration/theory_elaboration_matched_papers.csv"
OVERLAP = ROOT / "data/interim/theory_elaboration/theory_elaboration_probability_overlap_23.csv"
LEADING_TOPICS = ROOT / "data/processed/topics/native/query_3/assignments.csv"
ADDITIONAL_TOPICS = ROOT / "data/processed/topics/native/query_4/assignments.csv"
LEADING_VOS_MAP = ROOT / "data/vosdata/query 3.txt"
ADDITIONAL_VOS_MAP = ROOT / "data/vosdata/query 4.txt"
AUDIT_OUTPUT = (
    ROOT
    / "reports/analysis/tables/contrasting/close_reading_current_population_audit.csv"
)
SUPPLEMENT = ROOT / "docs/ETP supplementary material july2026 ks.docx"
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"


def flag(series: pd.Series) -> pd.Series:
    return pd.to_numeric(series, errors="coerce").fillna(0).astype(int).eq(1)


def normalize_doi(value: object) -> str:
    text = str(value or "").strip().lower()
    text = re.sub(r"^https?://(dx\.)?doi\.org/", "", text)
    return re.sub(r"\s+", "", text).rstrip(".,;)")


def normalize_title(value: object) -> str:
    text = unicodedata.normalize("NFKD", str(value or ""))
    text = text.encode("ascii", "ignore").decode().lower()
    return re.sub(r"[^a-z0-9]+", " ", text).strip()


def load_vos_map(path: Path) -> pd.DataFrame:
    frame = pd.read_csv(path, sep="\t", dtype=str, keep_default_na=False)
    required = {
        "id",
        "description",
        "url",
        "cluster",
        "weight<Links>",
        "weight<Total link strength>",
    }
    if missing := required.difference(frame.columns):
        raise RuntimeError(f"{path} is missing VOS columns: {sorted(missing)}")

    def description_title(value: object) -> str:
        match = re.search(
            r"<td>Title:</td><td>(.*?)</td>",
            html.unescape(str(value or "")),
            flags=re.IGNORECASE | re.DOTALL,
        )
        return match.group(1) if match else ""

    frame["doi_norm"] = frame["url"].map(normalize_doi)
    frame["title_norm"] = frame["description"].map(description_title).map(
        normalize_title
    )
    frame["vos_links"] = pd.to_numeric(frame["weight<Links>"], errors="raise").astype(
        int
    )
    frame["vos_total_link_strength"] = pd.to_numeric(
        frame["weight<Total link strength>"], errors="raise"
    ).astype(int)
    frame["vos_tls_rank"] = (
        frame["vos_total_link_strength"]
        .rank(method="min", ascending=False)
        .astype(int)
    )
    return frame


def add_vos_positions(
    reading: pd.DataFrame,
    *,
    membership_column: str,
    population_label: str,
    vos_map: pd.DataFrame,
) -> None:
    by_doi = {
        row["doi_norm"]: row
        for _, row in vos_map.iterrows()
        if row["doi_norm"]
    }
    by_title = {
        row["title_norm"]: row
        for _, row in vos_map.iterrows()
        if row["title_norm"]
    }
    population_size = len(vos_map)
    for index, paper in reading[reading[membership_column]].iterrows():
        match = by_doi.get(normalize_doi(paper["DOI"]))
        if match is None:
            match = by_title.get(normalize_title(paper["Title"]))
        if match is None:
            raise RuntimeError(
                f"No {population_label} VOS node for {paper['paper_id']}: "
                f"{paper['Title']}"
            )
        reading.at[index, "vos_population"] = population_label
        reading.at[index, "vos_node_id"] = match["id"]
        reading.at[index, "vos_cluster"] = match["cluster"]
        reading.at[index, "vos_links"] = int(match["vos_links"])
        reading.at[index, "vos_total_link_strength"] = int(
            match["vos_total_link_strength"]
        )
        reading.at[index, "vos_tls_rank"] = int(match["vos_tls_rank"])
        reading.at[index, "vos_population_nodes"] = population_size


def build_audit() -> tuple[pd.DataFrame, dict[str, int | float]]:
    reading = pd.read_csv(MATCHED, dtype=str, keep_default_na=False)
    overlap_ids = set(
        pd.read_csv(OVERLAP, dtype=str, keep_default_na=False)["paper_id"]
    )
    leading = pd.read_csv(LEADING_TOPICS, dtype=str, keep_default_na=False)
    additional = pd.read_csv(ADDITIONAL_TOPICS, dtype=str, keep_default_na=False)
    leading_vos = load_vos_map(LEADING_VOS_MAP)
    additional_vos = load_vos_map(ADDITIONAL_VOS_MAP)

    if len(reading) != 136 or reading["paper_id"].nunique() != 136:
        raise RuntimeError("The close-reading ledger is not the expected 136 unique papers")

    reading["in_leading_entrepreneurship_journals"] = flag(reading["in_query_3"])
    reading["in_additional_entrepreneurship_journals"] = flag(reading["in_query_4"])
    reading["in_combined_entrepreneurship"] = (
        reading["in_leading_entrepreneurship_journals"]
        | reading["in_additional_entrepreneurship_journals"]
    )
    reading["in_probability_annotation_target"] = reading["paper_id"].isin(
        overlap_ids
    )
    reading["interpretive_use"] = "cross-domain contrast only"
    reading.loc[
        reading["in_combined_entrepreneurship"], "interpretive_use"
    ] = "entrepreneurship interpretation"

    leading_map = leading.set_index("paper_id")
    additional_map = additional.set_index("paper_id")
    reading["current_topic_population"] = ""
    reading["current_topic_id"] = ""
    reading["current_topic_label"] = ""
    for column in (
        "vos_population",
        "vos_node_id",
        "vos_cluster",
        "vos_links",
        "vos_total_link_strength",
        "vos_tls_rank",
        "vos_population_nodes",
    ):
        reading[column] = ""
    for index, row in reading.iterrows():
        paper_id = row["paper_id"]
        if row["in_leading_entrepreneurship_journals"]:
            topic = leading_map.loc[paper_id]
            reading.at[index, "current_topic_population"] = (
                "Leading entrepreneurship journals"
            )
        elif row["in_additional_entrepreneurship_journals"]:
            topic = additional_map.loc[paper_id]
            reading.at[index, "current_topic_population"] = (
                "Additional entrepreneurship journals"
            )
        else:
            continue
        reading.at[index, "current_topic_id"] = topic["native_topic_id"]
        reading.at[index, "current_topic_label"] = topic["native_topic_label"]

    add_vos_positions(
        reading,
        membership_column="in_leading_entrepreneurship_journals",
        population_label="Leading entrepreneurship journals",
        vos_map=leading_vos,
    )
    add_vos_positions(
        reading,
        membership_column="in_additional_entrepreneurship_journals",
        population_label="Additional entrepreneurship journals",
        vos_map=additional_vos,
    )

    leading_read = reading[reading["in_leading_entrepreneurship_journals"]]
    additional_read = reading[reading["in_additional_entrepreneurship_journals"]]
    leading_topic_total = leading["native_topic_id"].loc[lambda x: x.ne("")].nunique()
    additional_topic_total = (
        additional["native_topic_id"].loc[lambda x: x.ne("")].nunique()
    )
    counts = {
        "reading_papers": len(reading),
        "leading_papers": int(reading["in_leading_entrepreneurship_journals"].sum()),
        "additional_papers": int(
            reading["in_additional_entrepreneurship_journals"].sum()
        ),
        "combined_papers": int(reading["in_combined_entrepreneurship"].sum()),
        "outside_combined_papers": int(
            (~reading["in_combined_entrepreneurship"]).sum()
        ),
        "probability_annotation_overlap": int(
            reading["in_probability_annotation_target"].sum()
        ),
        "combined_probability_annotation_overlap": int(
            (
                reading["in_combined_entrepreneurship"]
                & reading["in_probability_annotation_target"]
            ).sum()
        ),
        "leading_probability_annotation_overlap": int(
            (
                reading["in_leading_entrepreneurship_journals"]
                & reading["in_probability_annotation_target"]
            ).sum()
        ),
        "additional_probability_annotation_overlap": int(
            (
                reading["in_additional_entrepreneurship_journals"]
                & reading["in_probability_annotation_target"]
            ).sum()
        ),
        "leading_topics_represented": leading_read["current_topic_id"]
        .loc[lambda x: x.ne("")]
        .nunique(),
        "leading_topics_available": leading_topic_total,
        "additional_topics_represented": additional_read["current_topic_id"]
        .loc[lambda x: x.ne("")]
        .nunique(),
        "additional_topics_available": additional_topic_total,
        "leading_vos_nodes": len(leading_vos),
        "additional_vos_nodes": len(additional_vos),
        "leading_vos_matches": int(leading_read["vos_node_id"].ne("").sum()),
        "additional_vos_matches": int(additional_read["vos_node_id"].ne("").sum()),
        "leading_vos_tls_median": float(
            pd.to_numeric(leading_read["vos_total_link_strength"]).median()
        ),
        "additional_vos_tls_median": float(
            pd.to_numeric(additional_read["vos_total_link_strength"]).median()
        ),
        "leading_vos_tls_rank_min": int(
            pd.to_numeric(leading_read["vos_tls_rank"]).min()
        ),
        "leading_vos_tls_rank_max": int(
            pd.to_numeric(leading_read["vos_tls_rank"]).max()
        ),
        "additional_vos_tls_rank_min": int(
            pd.to_numeric(additional_read["vos_tls_rank"]).min()
        ),
        "additional_vos_tls_rank_max": int(
            pd.to_numeric(additional_read["vos_tls_rank"]).max()
        ),
        "leading_vos_top_quartile": int(
            (
                pd.to_numeric(leading_read["vos_tls_rank"])
                <= math.ceil(len(leading_vos) * 0.25)
            ).sum()
        ),
        "additional_vos_top_quartile": int(
            (
                pd.to_numeric(additional_read["vos_tls_rank"])
                <= math.ceil(len(additional_vos) * 0.25)
            ).sum()
        ),
    }
    expected = {
        "reading_papers": 136,
        "leading_papers": 51,
        "additional_papers": 73,
        "combined_papers": 124,
        "outside_combined_papers": 12,
        "probability_annotation_overlap": 23,
        "combined_probability_annotation_overlap": 20,
        "leading_probability_annotation_overlap": 9,
        "additional_probability_annotation_overlap": 11,
        "leading_topics_represented": 5,
        "leading_topics_available": 6,
        "additional_topics_represented": 8,
        "additional_topics_available": 8,
        "leading_vos_nodes": 622,
        "additional_vos_nodes": 961,
        "leading_vos_matches": 51,
        "additional_vos_matches": 73,
        "leading_vos_tls_median": 145.0,
        "additional_vos_tls_median": 181.0,
        "leading_vos_tls_rank_min": 3,
        "leading_vos_tls_rank_max": 545,
        "additional_vos_tls_rank_min": 3,
        "additional_vos_tls_rank_max": 900,
        "leading_vos_top_quartile": 25,
        "additional_vos_top_quartile": 37,
    }
    if counts != expected:
        raise RuntimeError(f"Unexpected close-reading audit counts: {counts}")

    columns = [
        "paper_id",
        "Title",
        "Authors",
        "Year",
        "Source title",
        "DOI",
        "Link",
        "in_leading_entrepreneurship_journals",
        "in_additional_entrepreneurship_journals",
        "in_combined_entrepreneurship",
        "current_topic_population",
        "current_topic_id",
        "current_topic_label",
        "vos_population",
        "vos_node_id",
        "vos_cluster",
        "vos_links",
        "vos_total_link_strength",
        "vos_tls_rank",
        "vos_population_nodes",
        "in_probability_annotation_target",
        "interpretive_use",
    ]
    AUDIT_OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    reading[columns].to_csv(AUDIT_OUTPUT, index=False)
    return reading, counts


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def paragraph_starting(document: Document, prefix: str):
    matches = [p for p in document.paragraphs if p.text.strip().startswith(prefix)]
    if len(matches) != 1:
        raise RuntimeError(f"Expected one paragraph beginning {prefix!r}; found {len(matches)}")
    return matches[0]


def paragraph_starting_any(document: Document, prefixes: tuple[str, ...]):
    matches = [
        p
        for p in document.paragraphs
        if any(p.text.strip().startswith(prefix) for prefix in prefixes)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning one of {prefixes!r}; found {len(matches)}"
        )
    return matches[0]


def write_cell(cell, text: str) -> None:
    set_paragraph_text(cell.paragraphs[0], text)
    for paragraph in cell.paragraphs[1:]:
        set_paragraph_text(paragraph, "")


def save_atomic(document: Document, path: Path) -> None:
    with tempfile.NamedTemporaryFile(
        prefix=path.stem + ".", suffix=".docx", dir=path.parent, delete=False
    ) as handle:
        temporary = Path(handle.name)
    try:
        document.save(temporary)
        os.replace(temporary, path)
    finally:
        temporary.unlink(missing_ok=True)


def update_supplement(counts: dict[str, int | float]) -> None:
    document = Document(SUPPLEMENT)
    set_paragraph_text(
        paragraph_starting(document, "The full-corpus model retained"),
        "The full-corpus model retained 4,322 conservative outliers and one title-only record below the usable-text threshold; these records remain in the master table with blank topic fields. In the broad-scope transformation, 10,585 papers were initially marked as outliers, 6,974 were reassigned when their highest non-outlier probability reached 0.05, and 3,611 remained unassigned. Low silhouette values indicate overlapping thematic boundaries. Topic labels are researcher-reviewed display labels; changing a label does not alter the fitted paper assignment or topic identifier. Current data-specific topic assignments organise the entrepreneurship reading set; they are navigation devices rather than theoretical constructs or sampling weights.",
    )
    set_paragraph_text(
        paragraph_starting_any(
            document,
            ("The interpretive reading set contained", "The auditable reading ledger contains"),
        ),
        "The auditable reading ledger contains 136 papers in the current corpus. Registered population flags place 51 papers in Leading entrepreneurship journals and 73 in Additional entrepreneurship journals, producing a 124-paper Combined entrepreneurship interpretation base. The remaining 12 papers are retained only as cross-domain contrasts. Current data-specific topic assignments organise the 124 papers across five of the six fitted Leading-journal topics and all eight assigned Additional-journal topics. Topic placement structures comparison but does not establish statistical representativeness.",
    )

    target_tables = [
        table
        for table in document.tables
        if table.cell(0, 0).text.strip() == "Step"
        and any(
            row.cells[0].text.strip() in {"Reading-set construction", "Current reading base"}
            for row in table.rows
        )
    ]
    if len(target_tables) != 1:
        raise RuntimeError(f"Expected one Appendix A9 table; found {len(target_tables)}")
    table = target_tables[0]
    rows = [
        (
            "Current reading base",
            "136 current-corpus papers: 51 Leading entrepreneurship journals, 73 Additional entrepreneurship journals, and 12 outside the Combined entrepreneurship population",
            "Entrepreneurship interpretation uses the 124 Combined papers; the 12 outside papers are retained only as cross-domain contrasts",
        ),
        (
            "Topic organisation",
            "The 124 entrepreneurship papers were linked to their current data-specific topic assignments: five of six Leading-journal topics and all eight assigned Additional-journal topics are represented",
            "Topics organise reading and comparison; they are not constructs, sampling strata, or evidence of representativeness",
        ),
        (
            "Bibliometric network-position check",
            "All 51 Leading and 73 Additional reading papers appear in the Query 3 and Query 4 VOSviewer document maps. Their total-link-strength ranks span 3-545 of 622 and 3-900 of 961, respectively; 25 Leading and 37 Additional papers fall in their map's top quartile",
            "This post hoc check demonstrates coverage across the available document networks, including highly connected papers. Total link strength was not the historical selection rule, and the reading set is not presented as the highest-total-link-strength subset",
        ),
        (
            "Evidence recorded",
            "AI role, technical type, mechanism, level, central claim, supporting passage or abstract evidence, and relationship to the selected pattern",
            "The reading set interprets theoretical meaning and counterexamples; it does not replace corpus-level coding or estimate prevalence",
        ),
        (
            "Interpretation retention",
            "Recognisable recurrence across multiple current topic assignments and evidence in both entrepreneurship populations when an interpretation is framed as field-wide; analytically important counterexamples were retained",
            "Prevents one topic, one journal population, or an isolated paper from determining an entrepreneurship-wide conclusion",
        ),
        (
            "Use in the manuscript",
            "Supporting cases, contrasting cases, counterexamples, and bounded entrepreneurship interpretation",
            "The reading set explains recurring configurations; all prevalence claims come from the full coded population and state their denominator",
        ),
        (
            "Probability-sample and annotation overlap",
            "23 of the 136 papers occur in the audited probability-sample annotation target. Of these, 20 are in Combined entrepreneurship: 9 Leading and 11 Additional; 3 are outside Combined entrepreneurship",
            "Natural overlap permits direct comparison on the same papers; targeted reading papers were not added to the probability sample",
        ),
        (
            "Human triangulation of interpretation",
            "A second researcher independently reviewed 14 papers: 11 from Combined entrepreneurship and 3 cross-domain contrasts. Insight-family allocation agreed for 14/14 papers (Cohen’s kappa = 1.00); 4 of the 14 also occur in the 23-paper annotation target",
            "This checks allocation to the interpretive insight families only; it does not validate the eight model-coded dimensions",
        ),
        (
            "Agency frontier",
            "Agency allocation was examined as a theoretical question when role, mechanism, and level raised issues of initiative, authority, decision rights, or responsibility",
            "Agency was not coded from titles, abstracts, and author keywords; stronger claims require full-text and claim-level evidence",
        ),
    ]
    if len(table.rows) == len(rows):
        table.add_row()
    if len(table.rows) != len(rows) + 1:
        raise RuntimeError("Appendix A9 table structure changed unexpectedly")
    for target_row, values in zip(table.rows[1:], rows):
        for cell, value in zip(target_row.cells, values):
            write_cell(cell, value)
    for target_row in table.rows:
        for cell in target_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)

    set_paragraph_text(
        paragraph_starting(document, "A candidate interpretation was retained"),
        "A candidate interpretation was retained only when it met the recurrence and evidence criteria in Table A9.1. The procedure produced one central entrepreneurship insight, bottleneck relocation, with organisational embedding as the condition under which firms can handle the relocated evaluation burden and domain context as the boundary shaping the mechanism through which relocation occurs. Agency remains a frontier insight because actor-like, judgment-related, and interactional claims raise questions about human-machine authority that cannot be resolved by the current abstract-level instrument. The post hoc VOSviewer audit located all 124 entrepreneurship reading papers in the available Query 3 and Query 4 document maps. Their total-link-strength ranks extend from highly connected to peripheral positions, so the result supports network coverage rather than a claim that the reading set comprises the highest-ranked papers. The VOS network construction and counting method must be disclosed before total link strength is given a stronger substantive interpretation.",
    )
    save_atomic(document, SUPPLEMENT)


def update_manuscript() -> None:
    document = Document(MANUSCRIPT)
    research_design = paragraph_starting(document, "This study uses theory elaboration")
    set_paragraph_text(
        research_design,
        research_design.text.replace(
            "Core entrepreneurship, Additional entrepreneurship",
            "Leading entrepreneurship journals, Additional entrepreneurship",
        ),
    )
    set_paragraph_text(
        paragraph_starting(document, "Topic modeling followed construct-specification coding"),
        "Topic modeling followed construct-specification coding and was kept analytically separate so that discovered topics could not influence the instrument. BERTopic models were fitted to titles, abstracts, and author keywords and were selected through quantitative diagnostics followed by semantic review. The resulting topics are used to navigate research conversations and organise close reading, not as objective theoretical categories or as the source of the theory-elaboration conclusions. Current data-specific topic assignments structure comparison within the entrepreneurship reading set. Complete model settings, grid searches, outlier rules, and selected resolutions are reported in Supplementary Figure A8.1 and Tables A8.1-A8.2.",
    )
    set_paragraph_text(
        paragraph_starting_any(
            document,
            ("The systematic close-reading set contained", "The auditable reading ledger contains"),
        ),
        "The auditable reading ledger contains 136 papers in the current corpus. Of these, 51 belong to Leading entrepreneurship journals and 73 to Additional entrepreneurship journals, producing a 124-paper Combined entrepreneurship interpretation base. The remaining 12 papers are used only as cross-domain contrasts. Current data-specific topic assignments organise the 124 entrepreneurship papers across five of six fitted Leading-journal topics and all eight assigned Additional-journal topics; topic placement structures comparison but is not treated as a construct or a statistical sampling weight. All 124 entrepreneurship papers are present in the available Query 3 and Query 4 VOSviewer document maps. Their total-link-strength ranks span broad portions of both maps, including 25 Leading and 37 Additional papers in the respective top quartiles; this is reported as a post hoc network-position check, not as the reading-set selection rule.",
    )
    set_paragraph_text(
        paragraph_starting(document, "A candidate interpretation was retained when"),
        "A candidate interpretation was retained when it recurred across multiple current topic assignments, was consequential for entrepreneurship theory, remained analytically distinct, and was supported in both entrepreneurship populations when framed as field-wide. Counterexamples and the 12 outside-domain papers were retained to bound interpretation rather than estimate prevalence. This procedure produced one central entrepreneurship insight, bottleneck relocation, an organisational embedding condition, a domain-mechanism boundary, and an agency frontier. The complete procedure, population audit, and human insight-allocation check are reported in Supplementary Appendix A9.",
    )
    save_atomic(document, MANUSCRIPT)


def main() -> None:
    _, counts = build_audit()
    update_supplement(counts)
    update_manuscript()
    print(f"Wrote {AUDIT_OUTPUT.relative_to(ROOT)}")
    print(f"Updated {SUPPLEMENT.relative_to(ROOT)}")
    print(f"Updated {MANUSCRIPT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
