"""Restructure the KS manuscript results around construct specification.

This updater operates only on the researcher-edited manuscript. It removes the
duplicated publication-growth subsection from Results, expands the Combined
entrepreneurship construct portrait using the frozen observed distributions,
and rewrites horizontal contrasting in the same dimension-led interpretive
style. Existing figures are retained because the growth figure belongs with
the Introduction in the assembled paper.
"""

from __future__ import annotations

import re
import shutil
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT = ROOT / "docs/ETP draft - July2026ks.docx"
BACKUP = ROOT / "docs/ETP draft - July2026ks.before-results-restructure.docx"
SPECIFICATION = (
    ROOT
    / "reports/analysis/tables/contrasting/construct_specification_by_population.csv"
)
COMBINED_PAPERS = 1_632


DIMENSION_ORDER = [
    "study_status",
    "ai_role",
    "technical_type",
    "mechanism",
    "level",
    "process_stage",
    "scope",
    "definition",
]


def find_paragraph(document: Document, prefix: str) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefix)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning {prefix!r}; found {len(matches)}"
        )
    return matches[0]


def find_paragraph_any(document: Document, prefixes: tuple[str, ...]) -> Paragraph:
    matches = [
        paragraph
        for paragraph in document.paragraphs
        if paragraph.text.strip().startswith(prefixes)
    ]
    if len(matches) != 1:
        raise RuntimeError(
            f"Expected one paragraph beginning with {prefixes!r}; found {len(matches)}"
        )
    return matches[0]


def replace_paragraph(paragraph: Paragraph, text: str) -> None:
    paragraph.clear()
    paragraph.add_run(text)


def insert_paragraph_after(
    reference: Paragraph,
    text: str,
    *,
    style: str | None = None,
    italic: bool = False,
) -> Paragraph:
    element = OxmlElement("w:p")
    reference._p.addnext(element)
    paragraph = Paragraph(element, reference._parent)
    if style:
        paragraph.style = style
    run = paragraph.add_run(text)
    run.italic = italic
    return paragraph


def remove_from_start_until(start: Paragraph, stop: Paragraph) -> None:
    """Remove every body element from start through the element before stop."""

    element = start._p
    while element is not stop._p:
        next_element = element.getnext()
        if next_element is None:
            raise RuntimeError("Reached document end before the section stop")
        element.getparent().remove(element)
        element = next_element


def shade_cell(cell, fill: str = "D9EAF7") -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.font.size = Pt(7.5)
    if bold:
        shade_cell(cell)


def make_observed_table(document: Document, frame: pd.DataFrame) -> Table:
    observed = frame[
        frame["population"].eq("combined")
        & frame["distribution"].eq("observed")
    ].copy()
    rows: list[list[str]] = []
    for dimension_id in DIMENSION_ORDER:
        selected = observed[observed["dimension_id"].eq(dimension_id)]
        if selected.empty:
            raise RuntimeError(f"Missing Combined observed distribution: {dimension_id}")
        denominator = int(selected["denominator"].iloc[0])
        leading = selected.sort_values(
            ["papers", "category"], ascending=[False, True]
        ).head(3)
        categories = "; ".join(
            f"{row.category} {float(row.share):.1%}"
            for row in leading.itertuples()
        )
        rows.append(
            [
                str(selected["dimension_label"].iloc[0]),
                f"{denominator:,} ({denominator / COMBINED_PAPERS:.1%})",
                categories,
            ]
        )

    table = document.add_table(rows=1, cols=3)
    table.style = "Normal Table"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [1.25, 1.30, 4.05]
    headers = [
        "Dimension",
        "Papers specifying the dimension",
        "Leading observed categories",
    ]
    for index, header in enumerate(headers):
        set_cell(table.rows[0].cells[index], header, bold=True)
    for values in rows:
        row = table.add_row()
        for index, value in enumerate(values):
            set_cell(row.cells[index], value)
        properties = row._tr.get_or_add_trPr()
        properties.append(OxmlElement("w:cantSplit"))
    for index, width in enumerate(widths):
        table.columns[index].width = Inches(width)
        for row in table.rows:
            row.cells[index].width = Inches(width)
    return table


def renumber_existing_results(document: Document) -> None:
    table_numbers = {"6": "7", "7": "8", "8": "9", "9": "10", "10": "11"}
    pattern = re.compile(r"\bTable (10|9|8|7|6)\b")
    for paragraph in document.paragraphs:
        text = paragraph.text
        updated = pattern.sub(lambda match: f"Table {table_numbers[match.group(1)]}", text)
        updated = updated.replace("Section 4.2.1", "Section 4.1.1")
        if updated != text:
            replace_paragraph(paragraph, updated)

    heading_updates = {
        "4.2 Construct specification within entrepreneurship": (
            "4.1 Construct specification within entrepreneurship"
        ),
        "4.2.1 Nested specification by study status": (
            "4.1.1 Nested specification by study status"
        ),
        "4.2.2 Core and Additional entrepreneurship boundaries": (
            "4.1.2 Leading and Additional entrepreneurship boundaries"
        ),
        "4.2.3 Technical type, role, and mechanism": (
            "4.1.3 Technical type, role, and mechanism"
        ),
        "4.3 Horizontal contrasting across business domains": (
            "4.2 Horizontal contrasting across business domains"
        ),
        "4.4 Vertical contrasting across levels": (
            "4.3 Vertical contrasting across levels"
        ),
        "4.5 Structuring recurring configurations": (
            "4.4 Structuring recurring configurations"
        ),
        "4.6 Entrepreneurship interpretation": "4.5 Entrepreneurship interpretation",
    }
    for old, new in heading_updates.items():
        replace_paragraph(find_paragraph(document, old), new)


def expand_construct_specification(document: Document, frame: pd.DataFrame) -> None:
    opening = find_paragraph(document, "Among the 1,497 entrepreneurship papers")
    replace_paragraph(
        opening,
        "The construct-specification results begin with the 1,632 papers in the Combined "
        "entrepreneurship population. Figure 2 and Table 6 report the observed composition "
        "of the eight dimensions. Observed refers to papers receiving a substantive category "
        "for the displayed dimension; each dimension therefore has its own denominator. "
        "Unobserved categories remain available in the complete distributions reported in "
        "the supplementary material and methodological platform.",
    )

    caption = insert_paragraph_after(
        opening,
        "Table 6. Observed construct specification in Combined entrepreneurship",
        style=opening.style.name,
    )
    caption.paragraph_format.keep_with_next = True
    table = make_observed_table(document, frame)
    caption._p.addnext(table._tbl)
    note = document.add_paragraph()
    note.style = opening.style
    note_run = note.add_run(
        "Note. Percentages in the final column use the observed denominator for "
        "that dimension, not all 1,632 papers. Process stage and definition clarity "
        "are exploratory dimensions."
    )
    note_run.italic = True
    table._tbl.addnext(note._p)

    figure_caption = find_paragraph(document, "Figure 2.")
    detail_texts = [
        (
            "Study status. Among the 1,497 papers with a clear study status, 55.0% examine AI "
            "as a substantive phenomenon, 25.7% use it as a research method, and 19.2% combine "
            "both uses. Entrepreneurship scholarship therefore predominantly studies AI as part "
            "of the phenomenon being explained, but 44.9% of the clear-status literature uses AI "
            "methodologically or combines methodological and substantive uses. These groups cannot "
            "be treated as one homogeneous research object."
        ),
        (
            "AI role. AI is most frequently assigned an instrumental role. Among 1,414 papers with "
            "a substantive role, 41.7% frame AI as a tool and 24.5% as a research method. The remaining "
            "literature assigns broader theoretical roles: 17.8% position AI as context, 10.0% as a "
            "firm capability, 4.7% as infrastructure, and 1.3% as an actor or agent. The AI construct "
            "in entrepreneurship therefore extends beyond a tool interpretation, but capability, "
            "infrastructural, and agentic formulations remain comparatively uncommon."
        ),
        (
            "Technical type. Only 886 papers, 54.3% of Combined entrepreneurship, identify a "
            "substantive technical form. Machine learning accounts for 47.0% of these papers, "
            "followed by generative AI at 14.0% and analytics at 8.5%. The remaining papers are "
            "distributed across natural language processing, automation, general AI, large language "
            "models, predictive AI, deep learning, recommender systems, and computer vision. The "
            "generic label AI consequently combines several technically distinct phenomena, while "
            "almost half of the entrepreneurship population does not identify a technical form in "
            "the available evidence."
        ),
        (
            "Mechanism. An observable mechanism is present in 838 papers, or 51.3% of the population. "
            "Prediction is the leading mechanism at 27.0%, followed closely by learning at 23.0% and "
            "uncertainty reduction at 11.5%. Judgment, stakeholder interaction, resource access, "
            "experimentation, search, and automated decision-making form smaller but theoretically "
            "important groups. The literature therefore concentrates on informational mechanisms, "
            "but almost half of the papers do not state what AI changes or enables within the title, "
            "abstract, or author keywords."
        ),
        (
            "Level of analysis. Level is the most consistently specified dimension: 1,590 papers, or "
            "97.4%, identify a substantive level. The firm dominates at 45.7%, followed by the "
            "individual entrepreneur at 20.5% and industry at 9.7%. Team-level work remains limited. "
            "Entrepreneurship research consequently places AI primarily within organizational settings, "
            "even where the theoretical consequences concern entrepreneurial judgment or action."
        ),
        (
            "Process stage. A process stage is observable in 1,065 papers, or 65.3%. Innovation accounts "
            "for 40.1% of the specified stages, followed by resource acquisition at 12.5% and static-input "
            "treatment at 8.3%. Because process stage is an exploratory dimension, these results indicate "
            "where AI is positioned in the entrepreneurial process but should not be treated as a fully "
            "stable classification."
        ),
        (
            "Scope conditions. Among the 1,245 papers stating a scope condition, 52.7% are sector-specific "
            "and 26.3% country-specific. Established firms, small and medium-sized enterprises, digital "
            "platforms, particular AI forms, early-stage ventures, ecosystems, and high-technology startups "
            "account for much smaller shares. The construct is therefore heavily embedded in sectoral and "
            "national settings rather than presented as universally applicable."
        ),
        (
            "Definition clarity. Only 464 papers, 28.4% of the Combined entrepreneurship population, "
            "display a definitional signal in the available evidence. Of these, 79.5% provide a partial "
            "definition, 13.1% define AI by example, and 7.3% provide an explicit definition aligned with "
            "the claim. This exploratory result concerns what is visible in titles, abstracts, and author "
            "keywords; it is not a judgment that the complete papers fail to define AI."
        ),
        (
            "Taken together, the observed construct is phenomenon-oriented but instrumentally framed, "
            "predominantly firm-level, technically centred on machine learning, and most often connected "
            "to prediction and learning. Its specification is nevertheless uneven: study status, role, and "
            "level are usually observable, whereas technical form, mechanism, and especially definitional "
            "signals are considerably less visible. This marginal portrait establishes the baseline for "
            "the nested analysis that follows."
        ),
    ]
    last = figure_caption
    for text in detail_texts:
        last = insert_paragraph_after(last, text, style=opening.style.name)

    robustness = find_paragraph(document, "The Gemini re-estimation retains")
    last._p.addnext(robustness._p)


def rewrite_horizontal_contrasting(document: Document) -> None:
    heading = find_paragraph(document, "4.2 Horizontal contrasting across business domains")
    intro = insert_paragraph_after(
        heading,
        "Horizontal contrasting applies the same observed-category rule to every domain and "
        "compares each within-domain share with the complete 22,345-paper corpus baseline. "
        "The baseline includes papers outside the nine selected domain rows, and papers can "
        "belong to more than one domain because assignments follow the source journal's official "
        "Scopus classifications. Figure 5 and Table 9 therefore show differences in construct "
        "specification, not mutually exclusive market shares.",
    )

    summary = find_paragraph(document, "The horizontal comparison shows specialization")
    replace_paragraph(
        summary,
        "AI role. Management Science and Operations Research provides the clearest instrumental "
        "profile: 77.8% of its 4,160 papers with an observed role treat AI as a tool, compared with "
        "62.3% across the full-corpus observed base. Entrepreneurship departs from this pattern in "
        "two ways. In Leading entrepreneurship journals, 32.3% of observed roles position AI as a "
        "research method, 15.1 percentage points above the full-corpus share. In Additional "
        "entrepreneurship, 12.9% position AI as a firm capability, 8.1 points above the baseline. "
        "The domain contrast is therefore not simply more or less AI; it concerns the theoretical "
        "work assigned to the construct.",
    )
    summary.style = intro.style
    last = summary
    last = insert_paragraph_after(
        last,
        "Technical type. Finance remains strongly machine-learning oriented: machine learning "
        "accounts for 60.7% of its 805 papers naming a technical form, compared with 48.4% in the "
        "full corpus. Management of Technology and Innovation instead gives greater weight to "
        "generative AI, which represents 12.3% of its 1,753 named technical forms compared with "
        "7.0% across the baseline. A technical label's prevalence therefore varies with the domain "
        "conversation in which AI is embedded.",
        style=summary.style.name,
    )
    last = insert_paragraph_after(
        last,
        "Mechanism. The strongest differences occur in what AI is claimed to change. Prediction "
        "accounts for 70.4% of observed mechanisms in environmental and sustainability research, "
        "27.6 percentage points above the full-corpus baseline. Marketing emphasizes transformed "
        "stakeholder interaction, at 27.6% compared with 8.8% across the corpus, while organization "
        "studies gives greater weight to altered judgment, at 15.9% compared with 7.4%. The same "
        "general AI label consequently supports different causal accounts across domains.",
        style=summary.style.name,
    )
    last = insert_paragraph_after(
        last,
        "Taken together, the horizontal contrasts show specialization rather than one management-wide "
        "AI construct. Operations-oriented research foregrounds tools, environmental research prediction, "
        "marketing interaction, organization studies judgment, finance machine learning, and technology-and-"
        "innovation research generative AI. Entrepreneurship combines methodological, contextual, capability, "
        "learning, judgment, and generative formulations. Domain context therefore changes both what AI is "
        "taken to be and the mechanism through which it is theorized to matter.",
        style=summary.style.name,
    )

    ft50 = find_paragraph(document, "The FT50 restriction")
    replace_paragraph(
        ft50,
        "The complete full-corpus distributions and FT50-restricted matrices are reported through the "
        "methodological platform, while Supplementary Appendix A11 documents the domain aggregation and "
        "comparison baseline. FT50 is retained as a boundary and robustness restriction rather than presented "
        "as a second descriptive results narrative because several domain-specific denominators become small."
    )


def standardize_leading_results_labels(document: Document) -> None:
    """Use the manuscript's public population label within the revised Results."""

    replacements = {
        "Table 8. Selected Core versus Additional entrepreneurship contrasts (observed view)": (
            "Table 8. Selected Leading versus Additional entrepreneurship contrasts (observed view)"
        ),
        "The within-entrepreneurship comparison is not flat. Core entrepreneurship": (
            "The within-entrepreneurship comparison is not flat. Leading entrepreneurship journals"
        ),
        "The nested comparison shows that these differences are concentrated": None,
        "The Core-Additional boundary result is the most stable": None,
    }
    caption = find_paragraph(
        document,
        "Table 8. Selected Core versus Additional entrepreneurship contrasts",
    )
    replace_paragraph(caption, replacements[caption.text.strip()])

    comparison = find_paragraph(
        document,
        "The within-entrepreneurship comparison is not flat. Core entrepreneurship",
    )
    replace_paragraph(
        comparison,
        comparison.text.replace(
            "Core entrepreneurship is",
            "Leading entrepreneurship journals are",
        ),
    )

    nested = find_paragraph(
        document,
        "The nested comparison shows that these differences are concentrated",
    )
    replace_paragraph(
        nested,
        nested.text.replace("in Core", "in Leading entrepreneurship journals")
        .replace("; Core gives", "; Leading journals give"),
    )

    robustness = find_paragraph(
        document,
        "The Core-Additional boundary result is the most stable",
    )
    replace_paragraph(
        robustness,
        robustness.text.replace(
            "The Core-Additional boundary result",
            "The Leading-Additional boundary result",
        ).replace(
            "Core entrepreneurship is",
            "Leading entrepreneurship journals are",
        ),
    )

    table8 = next(
        table
        for table in document.tables
        if table.rows
        and table.rows[0].cells[0].text.strip() == "Dimension"
        and len(table.rows[0].cells) == 5
        and table.rows[0].cells[2].text.strip() == "Core"
    )
    set_cell(table8.rows[0].cells[2], "Leading", bold=True)
    set_cell(table8.rows[0].cells[4], "Leading minus Additional", bold=True)


def normalize_revised_layout(document: Document) -> None:
    """Match the researcher-edited body style and prevent awkward table splits."""

    body_reference = find_paragraph(document, "The type-by-role matrix provides")
    body_style = body_reference.style
    body_font = next(
        (run.font.name for run in body_reference.runs if run.font.name),
        "Times New Roman",
    )
    revised_prefixes = (
        "The construct-specification results begin",
        "Study status. Among the 1,497 papers",
        "AI role. AI is most frequently assigned",
        "Technical type. Only 886 papers",
        "Mechanism. An observable mechanism",
        "Level of analysis. Level is the most consistently",
        "Process stage. A process stage is observable",
        "Scope conditions. Among the 1,245 papers",
        "Definition clarity. Only 464 papers",
        "Taken together, the observed construct",
        "Horizontal contrasting applies the same observed-category",
        "AI role. Management Science and Operations Research",
        "Technical type. Finance remains strongly machine-learning oriented",
        "Mechanism. The strongest differences occur",
        "Taken together, the horizontal contrasts",
        "The complete full-corpus distributions and FT50-restricted matrices",
    )
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith(revised_prefixes):
            paragraph.style = body_style
            for run in paragraph.runs:
                run.font.name = body_font
                run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), body_font)

    results = find_paragraph(document, "Results")
    results.paragraph_format.page_break_before = True

    observed_note = find_paragraph(document, "Note. Percentages in the final column")
    for run in observed_note.runs:
        run.font.size = Pt(9)

    horizontal_caption = find_paragraph(
        document, "Table 9. Selected theoretically meaningful horizontal contrasts"
    )
    horizontal_caption.paragraph_format.page_break_before = True
    horizontal_caption.paragraph_format.keep_with_next = True


def validate(document: Document) -> None:
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    required = [
        "4.1 Construct specification within entrepreneurship",
        "Table 6. Observed construct specification in Combined entrepreneurship",
        "Study status. Among the 1,497 papers",
        "4.2 Horizontal contrasting across business domains",
        "AI role. Management Science and Operations Research",
        "Technical type. Finance remains strongly machine-learning oriented",
        "Mechanism. The strongest differences occur",
        "4.5 Entrepreneurship interpretation",
        "Table 8. Selected Leading versus Additional entrepreneurship contrasts",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Missing revised manuscript content: {missing}")
    forbidden = [
        "4.1 Growth and analytical population",
        "4.2 Construct specification within entrepreneurship",
        "Table 6. Leading nested specification result",
        "Table 8. Selected theoretically meaningful horizontal contrasts",
        "Table 8. Selected Core versus Additional entrepreneurship contrasts",
    ]
    present = [item for item in forbidden if item in text]
    if present:
        raise RuntimeError(f"Stale results content remains: {present}")
    table6 = [
        table
        for table in document.tables
        if table.rows and table.rows[0].cells[0].text.strip() == "Dimension"
        and len(table.rows[0].cells) == 3
        and table.rows[0].cells[1].text.strip().startswith("Papers specifying")
    ]
    if len(table6) != 1 or len(table6[0].rows) != 9:
        raise RuntimeError("Observed construct Table 6 was not created correctly")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(MANUSCRIPT, BACKUP)

    document = Document(MANUSCRIPT)
    if any(
        paragraph.text.strip().startswith(
            "Table 6. Observed construct specification in Combined entrepreneurship"
        )
        for paragraph in document.paragraphs
    ):
        if any(
            paragraph.text.strip().startswith(
                "Table 8. Selected Core versus Additional entrepreneurship contrasts"
            )
            for paragraph in document.paragraphs
        ):
            standardize_leading_results_labels(document)
        normalize_revised_layout(document)
        validate(document)
        temporary = MANUSCRIPT.with_suffix(".results-restructure.tmp.docx")
        document.save(temporary)
        Document(temporary)
        temporary.replace(MANUSCRIPT)
        print(f"Verified and normalized: {MANUSCRIPT}")
        return

    growth = find_paragraph(document, "4.1 Growth and analytical population")
    construct = find_paragraph(document, "4.2 Construct specification within entrepreneurship")
    remove_from_start_until(growth, construct)
    renumber_existing_results(document)

    frame = pd.read_csv(SPECIFICATION)
    expand_construct_specification(document, frame)
    rewrite_horizontal_contrasting(document)
    standardize_leading_results_labels(document)
    normalize_revised_layout(document)
    validate(document)

    temporary = MANUSCRIPT.with_suffix(".results-restructure.tmp.docx")
    document.save(temporary)
    reopened = Document(temporary)
    validate(reopened)
    temporary.replace(MANUSCRIPT)
    print(f"Updated {MANUSCRIPT}")
    print(f"Backup  {BACKUP}")


if __name__ == "__main__":
    main()
